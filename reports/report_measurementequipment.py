from docx.enum.text import WD_BREAK
from docx.enum.text import WD_COLOR_INDEX, WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Cm
from docx.shared import Pt
import psycopg2

def set_col_width_dev(table):
    widths = (Cm ( 4.4 ) ,Cm ( 6 ))
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width=width
def q_run(connD, querry):
	username = connD[0]
	password = connD[1]
	host = connD[2]
	kport = "5432"
	kdb = "postgres"
	# cs = ' host="localhost",database="postgres", user= "postgres" , password="info" '
	cs = "dbname=%s user=%s password=%s host=%s port=%s" % (kdb, username, password, host, kport)
	conn = None
	conn = psycopg2.connect(str(cs))
	cur = conn.cursor()
	cur.execute(querry)
	try:
		result = cur.fetchall()
		return result
	except:
		pass
	conn.commit()
	cur.close()

def MarVibENG(document,connD,rn):

	summary = document.add_paragraph()
	r=summary.add_run('Measurement equipment:')
	summary.alignment=WD_ALIGN_PARAGRAPH.CENTER
	r.bold=True
	r.font.name = 'Calibri'


	shiptable = document.add_table(rows=6, cols=2)  # trzeba usunąć enter przed
	shiptable.style = 'Table Grid'
	shiptable.cell(0, 0).merge(shiptable.cell(0, 1))
	shiptable.alignment=WD_TABLE_ALIGNMENT.CENTER

	##SHIPTABLE [0,0]
	ht = shiptable.cell(0, 0).paragraphs[0]
	r0 = ht.add_run()
	r0.text = 'Technical data'
	r0.font.name = 'Calibri'
	ht.alignment=WD_ALIGN_PARAGRAPH.CENTER

	##SHIPTABLE [1,0]
	ht = shiptable.cell(1, 0).paragraphs[0]
	r0 = ht.add_run()
	r0.text = 'Maker:'
	r0.font.name = 'Calibri'
	##SHIPTABLE [2,0]
	ht = shiptable.cell(2, 0).paragraphs[0]
	r0 = ht.add_run()
	r0.text = 'Type:'
	r0.font.name = 'Calibri'
	##SHIPTABLE [3,0]
	ht = shiptable.cell(3, 0).paragraphs[0]
	r0 = ht.add_run()
	r0.text = 'Serial number:'
	r0.font.name = 'Calibri'
	##SHIPTABLE [4,0]
	ht = shiptable.cell(4, 0).paragraphs[0]
	r0 = ht.add_run()
	r0.text = 'Measuring range:'
	r0.font.name = 'Calibri'
	##SHIPTABLE [5,0]
	ht = shiptable.cell(5, 0).paragraphs[0]
	r0 = ht.add_run()
	r0.text = 'Indication error:'
	r0.font.name = 'Calibri'

	querry = "select device, serialno, caldue from equipment where lp = (select equipment from shipsdata where shipid = (select shipid from harmonogram where report_number = '" + str(rn) + "' limit 1) limit 1) limit 1"
	try:eqdata = list(q_run(connD, querry))[0]
	except: eqdata=list(['ERROR','ERROR','ERROR'])
	##SHIPTABLE [1,1]
	ht = shiptable.cell(1, 1).paragraphs[0]
	r0 = ht.add_run()
	r0.text = str('Info Marine')
	r0.font.name = 'Calibri'
	##SHIPTABLE [2,1]
	ht = shiptable.cell(2, 1).paragraphs[0]
	r0 = ht.add_run()
	r0.text = str(eqdata[0])
	r0.font.name = 'Calibri'
	##SHIPTABLE [3,1]
	ht = shiptable.cell(3, 1).paragraphs[0]
	r0 = ht.add_run()
	r0.text = str(eqdata[1])
	r0.font.name = 'Calibri'
	# ##SHIPTABLE [4,1]
	# ht = shiptable.cell(4, 1).paragraphs[0]
	# r0 = ht.add_run()
	# r0.text = str(eqdata[2])

	##SHIPTABLE [4,1]
	ht = shiptable.cell(4, 1).paragraphs[0]
	r0 = ht.add_run()
	r0.text = str('2Hz-30kHz / RPM = 60-20000')
	r0.font.name = 'Calibri'
	##SHIPTABLE [5,1]
	ht = shiptable.cell(5, 1).paragraphs[0]
	r0 = ht.add_run()
	r0.text = str('±0,5%')
	r0.font.name = 'Calibri'

	ht = document.add_paragraph()
	r0 = ht.add_run()
	r0.text = str('MarVib is calibrated, certificate for verification - if required.')
	r0.font.name = 'Calibri'
	set_col_width_dev ( shiptable )