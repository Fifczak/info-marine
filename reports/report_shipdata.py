import psycopg2
from docx.enum.text import WD_ALIGN_PARAGRAPH


def getshipsdata ( parent ):
    # przemyśleć czy jest sens rozbijać na poszczególne zmienne, czy lepiej result[indeks] wbijać od razu do tabeli
    kport = 5432
    kdb = 'postgres'
    kport = "5432"
    username = 'postgres'
    password = 'info'
    host = 'localhost'
    cs = "dbname=%s user=%s password=%s host=%s port=%s" % (kdb ,username ,password ,host ,kport)
    conn = None
    conn = psycopg2.connect ( str ( cs ) )
    cur = conn.cursor ()
    # bez limit 1, zeby pobierać od razu wszystkie rekordy
    print ( parent[ 0 ][ 0 ] )
    querry = "select * from shipsdata where shipid = '" + str ( parent[ 0 ][ 0 ] ) + "'"

    cur.execute ( querry )
    result = cur.fetchall ()
    conn.commit ()
    cur.close ()
    # tu probuje dojsc o co chodzi z indeksowaniem.. jezeli mamy w bazie np 12 rekordów, to trzeba zapisac 12x [0] przy results? ;x
    print ( 'Cały res:' + str ( result ) )
    print ( 'Res[0] - usuwa jedynie kwadratowy nawias:' + str ( result[ 0 ] ) )
    print ( 'Res[1] - 0 w indeksach dla tabeli 2W:' + str ( result[ 0 ][ 0 ] ) )
    global ship_type ,Lpp ,Breadth  # dałem jako zmienną globalną
    ship_type = str ( result[ 0 ][ 1 ] )  # tu jest 'ship type'
    Breadth = str ( round ( result[ 0 ][ 3 ] ,2 ) )  # tu jest Lbp
    Lpp = str ( round ( result[ 0 ][ 2 ] ,2 ) )  # tu jest B
    print ( 'Ship type:' + ship_type )
    print ( 'Długość:' + Lpp )
    print ( 'Szerokosc:' + Breadth )


    return result

def shipdata ( document ) :

    # sdata-nazwa dla tabeli z danymi statku
    sdata = document.add_table ( rows=3 ,cols=2 )
    # konwencja: c - nr komórki, 1 liczba oznacza indeks wiersza 2 liczba indeks kolumny
    c00 = sdata.cell ( 0 ,0 ).paragraphs[ 0 ].add_run ( 'Ship Type:' )
    c00.bold = True
    sdata.style = 'Table Grid'
    c00.add_break ()
    sdata.cell ( 0 ,0 ).paragraphs[ 0 ].add_run ( ship_type )
    c10 = sdata.cell ( 1 ,0 ).paragraphs[ 0 ].add_run ( 'Sea depth:' )
    c10.bold = True
    c10.add_break ()
    sdata.cell ( 1 ,0 ).paragraphs[ 0 ].add_run ( 'Least two times greater than Vessel draught.' )
    c01 = sdata.cell ( 0 ,1 ).paragraphs[ 0 ].add_run ( 'Main dimensions: ' )
    c01.bold = True
    # sdata.cell(0,1).paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
    sdata.cell ( 0 ,1 ).paragraphs[ 0 ].runs[ 0 ].alignment = WD_ALIGN_PARAGRAPH.CENTER
    sdata.cell ( 0 ,1 ).paragraphs[ 0 ].add_run ().add_break ()
    sdata.cell ( 0 ,1 ).paragraphs[ 0 ].add_run (
        'Length (b.p.) ...................................' + '%.2f'%Lpp.replace ( "." ,"," ) + 'm' ).add_break ()
    sdata.cell ( 0 ,1 ).paragraphs[ 0 ].add_run (
        'Breadth (B) ...................................' + Breadth.replace ( "." ,"," ) + 'm' )
    sdata.cell ( 1 ,1 ).paragraphs[ 0 ].add_run ( 'Weather conditions: ' ).bold = True
    weather = sdata.cell ( 1 ,1 ).add_paragraph ( '-' ).alignment = WD_ALIGN_PARAGRAPH.CENTER
    sdata.cell ( 1 ,1 ).add_paragraph ()
    sdata.cell ( 2 ,0 ).merge ( sdata.cell ( 2 ,1 ) )
    sdata.cell ( 2 ,0 ).paragraphs[ 0 ].add_run ( 'Measurement method: ' )
    isoexpl = sdata.cell ( 2 ,0 ).add_paragraph ()
    isoexplr = isoexpl.add_run ( 'According to standard ISO 10816 : - procedure No. 2 Measurement report' )
    isoexplr.bold = True
