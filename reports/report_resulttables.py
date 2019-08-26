import time

import chart_scripts
from report_styles import loadstyles
import psycopg2
from chart_scripts import *
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT
from docx.enum.text import WD_COLOR_INDEX, WD_ALIGN_PARAGRAPH
from docx.shared import Cm,Pt
from docx.shared import RGBColor
from report_database import *
from report_headtables import *
from docx.shared import Inches
from docx.shared import RGBColor
from tqdm import tqdm
import os

from PIL import ImageTk, Image



def q_run(connD, querry):
	username = connD[0]
	password = connD[1]
	host = connD[2]
	kport = "5432"
	kdb = "postgres"

	cs = "dbname=%s user=%s password=%s host=%s port=%s" % (kdb, username, password, host, kport)
	conn = None
	conn = psycopg2.connect(str(cs))
	cur = conn.cursor()
	cur.execute(querry)
	try:
		result = cur.fetchall()
		return result
	except:
		pass
	conn.commit()
	cur.close()



def set_cols_width_trend ( table ):
	widths = (Cm ( 1.5 ) ,Cm ( 15 ))
	for row in table.rows:
		for idx ,width in enumerate ( widths ):
			row.cells[ idx ].width = width



def set_col_width_GSR( table ):  # funkcja do stałej szerokości komórek w raportach GSR
	widths = (Cm( 7 ) ,Cm( 1.5 ) ,Cm( 1.5 ) ,Cm( 1.5 ) ,Cm( 1.5 ) ,Cm( 1.8 ) ,Cm( 5 ))
	for row in table.rows:
		for idx ,width in enumerate( widths ):
			row.cells[ idx ].width = width

def prepare_IM( connD ,report_number):  # RETURN MEASLIST
	measlist = list()
	class meas( object ):
		def __init__( self ):
			self.id = ''
			self.name = ''
			self.maxval = ''
			self.maxval2 = ''
			self.limit = ''
			self.maxenv = '-'
			self.trend = list()
			self.standard = ''
			self.date = ''
			self.drivenby = ''
			self.sort = ''
			self.sort2 = ''
			self.pms = ''
			self.Dlimit = ''
			self.BID = ''
			self.VSG = False
	def loadData( connD ):
		def countLimit( standard ,value ):
			for limNo in limits:
				if str( limNo[ 0 ] ) == str( standard ):
					if value <= float( limNo[ 3 ] ):  # IF LIM1
						limSrt = str( limNo[ 2 ] )
						break
					else:
						if value <= float( limNo[ 5 ] ):  # IF LIM2
							limSrt = str( limNo[ 4 ] )
							break
						else:
							if value <= float( limNo[ 7 ] ):  # IF LIM3
								limSrt = str( limNo[ 6 ] )
								break
							else:
								limSrt = str( limNo[ 8 ] )
								break
			try:return limSrt
			except:return ('Limit count error')
		def getTrend( self ):
			trendlist = list()
			for line in reportresults:
				if str( line[ 0 ] ) == str( self.id ):
					if str( line[ 1 ] ) != str( report_number ):
						newdate = time.strptime( str( line[ 6 ] ) ,"%Y-%m-%d" )
						mydate = time.strptime( str( self.date ) ,"%Y-%m-%d" )
						if newdate < mydate:
							trendlist.append( str( line[ 3 ] ) )  # VAL
							trendlist.append( str( line[ 6 ] ) )  # DATE
							change = abs( (float( line[ 3 ] ) - float( self.maxval )) / float( line[ 3 ] ) )
							#print(self.name, change)
							if change <= 0.05:
								TREND = 'C'
							elif float( line[ 3 ] ) < float( self.maxval ):
								TREND = 'U'
							elif float( line[ 3 ] ) > float( self.maxval ):
								TREND = 'D'
							trendlist.append( TREND )  # TREND=> U-UP, D-DOWN, C-CONST
							break
			self.trend = trendlist
			return measlist
		querry = "Select parent,max(date) from measurements_low where raport_number = '" + str(report_number) + "' group by parent limit 1"
		parent = list(q_run( connD ,querry ))[0][0]
		maxdate = list(q_run(connD, querry))[0][1]
		querry = """select standard,
							limit_1_value,limit_1_name,
							limit_2_value,limit_2_name, 
							limit_3_value,limit_3_name,
							limit_4_value,limit_4_name,
						envflag
					from standards"""
		limits = q_run( connD ,querry )
		querry = """select 
					 ml.id, ml.raport_number, dev.name,  max(ml.value) as RMS, ml2.max as Envelope, dev.norm ,max(ml.date), dev.drivenby,dss.sort,dev.pms,sta.limit_4_value,mlVSG.max,bid.bourbon, pts.visible
					from measurements_low as ml
					left join (select 
								 ml.id, ml.raport_number,  max(ml.value)
								 from measurements_low as ml
							     left join points pts on ml.id = pts.id and ml.point = pts.point
								 where ml.value <> -1 and type = 'envelope P-K' and parent = {} and pts.visible = True
								 group by ml.id, ml.raport_number order by raport_number DESC) as ml2 on ml.id = ml2.id and ml.raport_number = ml2.raport_number
					 left join (select ml.id, ml.raport_number,max(value) 
									from measurements_low  ml
									left join points pts on ml.id = pts.id and ml.point = pts.point
									where unit = 'VSG' and pts.visible = True
									group by ml.id,ml.raport_number) mlVSG on ml.id = mlVSG.id and ml.raport_number = mlVSG.raport_number 
					 left join devices as dev on ml.id = dev.id
					 left join(select cast (id as integer), sort from ds_structure where id ~E'^\\\d+$' ) as dss on ml.id = dss.id
					 left join standards sta on dev.norm = sta.standard
					 left join bourbonid bid on ml.id = bid.devid
					 left join points pts on ml.id = pts.id and ml.point = pts.point
					 left join harmonogram har on ml.raport_number = har.report_number
					 where ml.date <= cast('{}' as date) and ml.date >= cast('{}' as date) - interval '2 years'  and ml.value <> -1 and ml.type = 'RMS' and ml.parent = {} and dss.sort is distinct from null and pts.visible = True and har.trendchart_visible_flag = True
					 group by ml.id, ml.raport_number, ml2.max,dev.name, dev.norm,dev.drivenby,dss.sort,dev.pms,sta.limit_4_value,mlVSG.max,bid.bourbon, pts.visible
					  order by raport_number DESC""".format(str(parent),maxdate,maxdate,str(parent))
		reportresults = q_run( connD ,querry )

		querry = "select id, drivenby from devices where drivenby IN" \
							 " (select id from measurements_low where raport_number = '{}' group by id)".format(report_number)
		drivenbyresult = q_run( connD ,querry )
		_idlist = list()
		_idlist.clear()
		drivingdict = dict()
		for line in drivenbyresult:
			drivingdict[line[1]] = line[0]
		for line in reportresults:
			if str(line[1]) == report_number:
				_idlist.append(str(line[0]).strip())


		for line in reportresults:
			x = meas()
			if str( line[ 1 ] ) == str( report_number ):
				x.id = line[ 0 ]

				if x.id in drivingdict:
					drivid = str(drivingdict[x.id]).strip()

					if drivid  not in _idlist:
						y = meas()
						y.id = drivid
						y.maxval = '-'
						y.maxenv = '-'
						querry = """select 
					ml.id, ml.raport_number, dev.name,  max(ml.value) as RMS, ml2.max as Envelope, dev.norm ,max(ml.date), dev.drivenby,dss.sort,dev.pms,sta.limit_4_value,mlVSG.max,bid.bourbon, pts.visible
					from measurements_low as ml
					left join (select 
								 ml.id, ml.raport_number,  max(ml.value)
								 from measurements_low as ml
							     left join points pts on ml.id = pts.id and ml.point = pts.point
								 where ml.value <> -1 and type = 'envelope P-K' and ml.id = {} and pts.visible = True
								 group by ml.id, ml.raport_number order by raport_number DESC) as ml2 on ml.id = ml2.id and ml.raport_number = ml2.raport_number
					 left join (select ml.id, ml.raport_number,max(value) 
									from measurements_low  ml
									left join points pts on ml.id = pts.id and ml.point = pts.point
									where unit = 'VSG' and pts.visible = True
									group by ml.id,ml.raport_number) mlVSG on ml.id = mlVSG.id and ml.raport_number = mlVSG.raport_number 
					 left join devices as dev on ml.id = dev.id
					 left join(select cast (id as integer), sort from ds_structure where id ~E'^\\\d+$' ) as dss on ml.id = dss.id
					 left join standards sta on dev.norm = sta.standard
					 left join bourbonid bid on ml.id = bid.devid
					 left join points pts on ml.id = pts.id and ml.point = pts.point
					 left join harmonogram har on ml.raport_number = har.report_number
					 where ml.date <= cast('{}' as date) and ml.date >= cast('{}' as date) - interval '2 years'  and ml.value <> -1 and ml.type = 'RMS' and ml.parent = {} and dss.sort is distinct from null and pts.visible = True and har.trendchart_visible_flag = True
				 group by ml.id, ml.raport_number, ml2.max,dev.name, dev.norm,ml.date,dev.drivenby,dss.sort,dev.pms,sta.limit_4_value,mlVSG.max,bid.bourbon, pts.visible
				  order by raport_number DESC""".format(drivid,maxdate,maxdate,drivid)
						try:
							drivline = list(q_run(connD, querry))[0]

							y.name = drivline[2]
							y.limit = '-'
							y.drivenby = x.id
							y.sort0 = drivline[8]
							y.sort = drivline[8][:5]
							y.sort2 = drivline[8][:2]
							y.pms = drivline[9]
							y.Dlimit = drivline[10]
							measlist.append(y)
						except:
							print('driven by error(Driven side):',x.id, '/' , y.id)
				x.name = line[ 2 ]
				if str(line[11]) == 'None':
					x.maxval = line[ 3 ]
					x.maxenv = line[ 4 ]
				else:
					x.maxval = line[11]
					x.VSG = True
				x.standard = line[ 5 ]
				x.date = line[ 6 ]
				getTrend( x )
				x.limit = countLimit( x.standard ,x.maxval )
				x.drivenby = line[ 7 ]

				x.sort0 = line[8]
				x.sort = line[ 8 ][ :5 ]
				x.sort2 = line[ 8 ][ :2 ]
				x.pms = line [9]
				x.Dlimit = line [10]
				x.BID = line[12]
				measlist.append( x )

		for x in measlist:

			if str(x.drivenby) in str(_idlist) or x.drivenby == None or x.drivenby == 0 :
				pass
			else:
				y = meas()
				y.id = x.drivenby
				_idlist.append(str(y.id))
				y.maxval = '-'
				y.maxenv = '-'
				querry = """select 
					ml.id, ml.raport_number, dev.name,  max(ml.value) as RMS, ml2.max as Envelope, dev.norm ,max(ml.date), dev.drivenby,dss.sort,dev.pms,sta.limit_4_value,mlVSG.max,bid.bourbon, pts.visible
					from measurements_low as ml
					left join (select 
								 ml.id, ml.raport_number,  max(ml.value)
								 from measurements_low as ml
								 left join points pts on ml.id = pts.id and ml.point = pts.point
								 where ml.value <> -1 and type = 'envelope P-K' and ml.id = {} and pts.visible = True
								 group by ml.id, ml.raport_number order by raport_number DESC) as ml2 on ml.id = ml2.id and ml.raport_number = ml2.raport_number
					 left join (select ml.id, ml.raport_number,max(value) 
									from measurements_low  ml
									left join points pts on ml.id = pts.id and ml.point = pts.point
									where unit = 'VSG' and pts.visible = True
									group by ml.id,ml.raport_number) mlVSG on ml.id = mlVSG.id and ml.raport_number = mlVSG.raport_number 
					 left join devices as dev on ml.id = dev.id
					 left join(select cast (id as integer), sort from ds_structure where id ~E'^\\\d+$' ) as dss on ml.id = dss.id
					 left join standards sta on dev.norm = sta.standard
					 left join bourbonid bid on ml.id = bid.devid
					 left join points pts on ml.id = pts.id and ml.point = pts.point
					 left join harmonogram har on ml.raport_number = har.report_number
					 where ml.date <= cast('{}' as date) and ml.date >= cast('{}' as date) - interval '2 years'  and ml.value <> -1 and ml.type = 'RMS' and ml.parent = {} and dss.sort is distinct from null and pts.visible = True and har.trendchart_visible_flag = True
				 group by ml.id, ml.raport_number, ml2.max,dev.name, dev.norm,ml.date,dev.drivenby,dss.sort,dev.pms,sta.limit_4_value,mlVSG.max,bid.bourbon, pts.visible
				  order by raport_number DESC""".format(y.id,maxdate,maxdate, y.id)

				try:
					drivline = list(q_run(connD, querry))[0]

					y.name = drivline[2]
					y.limit = '-'
					y.drivenby = x.id
					y.sort0 = drivline[8]
					y.sort = drivline[8][:5]
					y.sort2 = drivline[8][:2]
					y.pms = drivline[9]
					y.Dlimit = drivline[10]
					measlist.append(y)
				except:
					print('driven by error(Driving side) :',x.id, '/' , y.id)

	loadData( connD )

	return measlist


def drawtable_IM(document, measlist, connD, report_number,reporttype):
	def set_cols_width(table,widths):
		r = -1
		for row in tqdm(table.rows):
			r +=1
			for idx, width in enumerate(widths):
				row.cells[idx].width = width
	def colorlimit(par, limit):
		if limit == 'Cl. A' or limit == 'In Limit' or limit == 'V. I' or limit == 'Cl. A/B' or limit == 'In limit' or limit == 'Cl. B':
			par.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
		elif limit == 'Cl. B':
			par.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
		elif limit == 'Cl. C' or limit == 'V. II':
			par.font.highlight_color = WD_COLOR_INDEX.YELLOW

	os.chdir(os.path.dirname(os.path.realpath(__file__)))

	respar = document.add_paragraph('Results', 'IM HEAD')

	respar.add_run().add_break()
	resrun = respar.add_run(
		"In table are presented only readings with max. RMS results for each device equipment:")
	resrun.font.size = Pt(11)
	resrun.font.name = 'Calibri'
	resrun.font.color.rgb = RGBColor(0, 0, 0)
	if reporttype != 1 and reporttype != 2 and reporttype != 7 and reporttype != 8:
		root = tk.Tk()
		root.withdraw()
		MsgBox = messagebox.askquestion('Chart controll', 'Do you want to check te charts data?', )
		if MsgBox == 'yes':
			GUI = True
		elif MsgBox == 'no':
			GUI = False
		root.destroy()

	# querry = "Select parent from measurements_low where raport_number = '" + str(report_number) + "' limit 1"
	# parent = list(q_run(connD, querry))[0][0]
	querry = "Select parent,max(date) from measurements_low where raport_number = '" + str(
		report_number) + "' group by parent limit 1"
	parent = list(q_run(connD, querry))[0][0]
	maxdate = list(q_run(connD, querry))[0][1]

	querry = "Select sort, id from ds_structure where parent = '" + str(parent) + "' order by sort"
	sortlistQ = q_run(connD, querry)
	trueMeasList = list()
	activeIdList = list()
	activeSortList = list()
	activeSortPList = list()
	idlist = list()
	drivenByList = list()

	for sort in sortlistQ:
		for meas in measlist:

			if (sort[1]).isdigit == False:
				trueMeasList.append('header')
				drivenByList.append(997)
				idlist.append(997)
			if str(sort[1]) == str(meas.id):
				trueMeasList.append(meas)
				activeIdList.append(str(meas.id))
				activeSortList.append(str(meas.sort))
				activeSortPList.append(str(meas.sort2))
				drivenByList.append(meas.drivenby)
				idlist.append(meas.id)

	counter = 0
	counterCharts = 0
	xcord = 0
	i = -1



	for measStrip in sortlistQ:
		i += 1




	if reporttype == 1:
		col_BID = None
		col_PMS = 0
		col_name = 1
		col_val = 2
		col_limit = None
		col_class = 3
		col_env = 4
		col_trend = 5
		col_remark = 6
		widths = (Cm(2 ) ,Cm( 6.5 ) ,Cm( 1.3) ,Cm( 1.3) ,Cm( 1.3 ),Cm( 1.8 ) ,Cm( 7.8 ))

	if reporttype == 2:
		col_BID = None
		col_PMS = None
		col_name = 0
		col_val = 1
		col_limit = None
		col_class = 2
		col_env = 3
		col_trend = 4
		col_remark = 5
		widths = (Cm( 6.5 ) ,Cm( 1.3) ,Cm( 1.3) ,Cm( 1.3 ),Cm( 1.8 ) ,Cm( 7.8 ))

	if reporttype == 3:
		col_BID = None
		col_PMS = 0
		col_name = 1
		col_val = 2
		col_limit = None
		col_class = 3
		col_env = 4
		col_trend = None
		col_remark = 5
		widths = (Cm(2 ) ,Cm( 7 ) ,Cm( 1.4) ,Cm( 1.4) ,Cm( 1.8 ) ,Cm( 8 ))

	if reporttype == 4:
		col_BID = None
		col_PMS = None
		col_name = 0
		col_val = 1
		col_limit = None
		col_class = 2
		col_env = 3
		col_trend = None
		col_remark = 4
		widths = (Cm( 7 ) ,Cm( 1.4) ,Cm( 1.4) ,Cm( 1.8 ) ,Cm( 8 ))

	if reporttype == 5:
		col_BID = None
		col_PMS = 0
		col_name = 1
		col_val = 2
		col_limit = 3
		col_class = 4
		col_env = 5
		col_trend = None
		col_remark = 6
		widths = (Cm(2 ) ,Cm( 6 ) ,Cm( 1.4),Cm( 1.1) ,Cm( 1.4) ,Cm( 1.8) ,Cm( 8 ))

	if reporttype == 6:
		col_BID = None
		col_PMS = None
		col_name = 0
		col_val = 1
		col_limit = 2
		col_class = 3
		col_env = 4
		col_trend = None
		col_remark = 5
		widths = (Cm( 6 ) ,Cm( 1.4),Cm( 1.1) ,Cm( 1.4) ,Cm( 1.8) ,Cm( 8 ))

	if reporttype == 7:
		col_BID = 0
		col_PMS = 1
		col_name = 2
		col_val = 3
		col_limit = None
		col_class = 4
		col_env = 5
		col_trend = 6
		col_remark = 7
		widths = (Cm(1 ),Cm(2 ) ,Cm( 6.5 ) ,Cm( 1.3) ,Cm( 1.3) ,Cm( 1.3 ),Cm( 1.8 ) ,Cm( 6.8 ))

	if reporttype == 8:
		col_BID = 0
		col_PMS = None
		col_name = 1
		col_val = 2
		col_limit = None
		col_class = 3
		col_env = 4
		col_trend = 5
		col_remark = 6
		widths = (Cm(2 ) ,Cm( 6.5 ) ,Cm( 1.3) ,Cm( 1.3) ,Cm( 1.3 ),Cm( 1.8 ) ,Cm( 6.8 ))

	colsn = len(widths)
	resulttable = document.add_table(rows=1, cols=colsn)
	resulttable.style = 'Table Grid'




	if col_BID != None:
		ht = resulttable.cell(0, col_BID).paragraphs[0]
		ht.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
		resulttable.cell(0, col_BID).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
		r0 = ht.add_run('Bourbon ID').bold = True

	if col_PMS != None:
		ht = resulttable.cell(0, col_PMS).paragraphs[0]
		ht.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
		resulttable.cell(0, col_PMS).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
		r0 = ht.add_run('PMS').bold = True

	if col_name != None:
		ht = resulttable.cell(0, col_name).paragraphs[0]
		ht.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
		resulttable.cell(0, col_name).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
		r0 = ht.add_run('Machine name').bold = True

	if col_val != None:
		ht = resulttable.cell(0, col_val).paragraphs[0]
		ht.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
		resulttable.cell(0, col_val).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
		r0 = ht.add_run('Velocity' + chr(10) + 'RMS' + chr(10) + '(mm/s)' + chr(10) + 'Max').bold = True
	if col_limit != None:
		ht = resulttable.cell( 0 ,col_limit ).paragraphs[ 0 ]
		ht.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
		resulttable.cell(0,col_limit).vertical_alignment=WD_ALIGN_VERTICAL.CENTER
		r0 = ht.add_run( 'ISO limit' ).bold=True

	if col_class != None:
		ht = resulttable.cell(0, col_class).paragraphs[0]
		ht.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
		resulttable.cell(0, col_class).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
		r0 = ht.add_run('ISO standard').bold = True

	if col_env != None:
		ht = resulttable.cell(0, col_env).paragraphs[0]
		ht.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
		resulttable.cell(0, col_env).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
		resulttable.cell(0, col_env).paragraphs[0].add_run().add_break()
		if str(parent) == '150':
			unit = '(gE)'
			envtype = '0-Peak'
		elif str(parent) == '234':
			unit = '(m/s2)'
			envtype = 'RMS'
		else:
			unit = '(m/s2)'
			envtype = '0-Peak'

		r0 = ht.add_run('Bearing Envelope' + chr(10) + (envtype) + chr(10) + (unit) + chr(10) +  'Max').bold = True

	if col_trend != None:
		ht = resulttable.cell( 0 ,col_trend ).paragraphs[ 0 ]
		ht.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
		resulttable.cell(0,col_trend).vertical_alignment=WD_ALIGN_VERTICAL.CENTER
		r0 = ht.add_run( 'Trend Velocity RMS (mm/s) Max' ).bold=True

	if col_remark != None:
		ht = resulttable.cell(0, col_remark).paragraphs[0]
		ht.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
		resulttable.cell(0, col_remark).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
		r0 = ht.add_run('Remarks and suggestions').bold = True

	set_cols_width(resulttable,widths)

	xcord = 0
	i = -1
	ids = list()


	for measStrip in tqdm(sortlistQ):

		i += 1
		if (measStrip[1]).isdigit() == False:  ######## NAGŁÓWKI
			#try:

				if measStrip[0][-5:] == '00.00' and sortlistQ[i + 1][0][-5:] != '00.00' and \
					sortlistQ[i + 1][0][-3:] == '.00':
					if str(measStrip[0][:2]) in activeSortPList:

						resulttable.add_row()
						ht = resulttable.cell(xcord + 1, 0).paragraphs[0]
						resulttable.cell(xcord + 1, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

						xcord += 1
						ht.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

						r0 = ht.add_run(measStrip[1])
						r0.font.color.rgb = RGBColor(0, 0, 255)
						r0.bold = True
						resulttable.cell(xcord, 0).merge(resulttable.cell(xcord, len(widths)-1))
						resulttable.rows[xcord].height = WD_ROW_HEIGHT.EXACTLY
						resulttable.rows[xcord].height = Cm(0.5)

						continue
				elif measStrip[0][-5:] != '00.00' and measStrip[0][-3:] == '.00':
					if str(measStrip[0][:5]) in activeSortList:
						resulttable.add_row()
						ht = resulttable.cell(xcord + 1, 0).paragraphs[0]
						resulttable.cell(xcord + 1, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

						xcord += 1

						r0 = ht.add_run(measStrip[1])
						r0.font.color.rgb = RGBColor(128, 0, 128)
						r0.bold = True
						resulttable.cell(xcord, 0).merge(resulttable.cell(xcord, len(widths)-1))
						resulttable.rows[xcord].height = WD_ROW_HEIGHT.EXACTLY
						resulttable.rows[xcord].height = Cm(0.5)
						continue
			#except:
				#pass
		else:  ######## POMIARY
			p = -1

			for xx in trueMeasList:

				p += 1
				if str(measStrip[1]) == str(xx.id):
					#try:
						resulttable.add_row()
						ids.append(xx.id)
						if col_BID != None:
							ht = resulttable.cell(xcord + 1, col_BID).paragraphs[0]
							ht.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
							resulttable.cell(xcord + 1, col_BID).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
							r0 = ht.add_run(str(xx.BID))

						if col_PMS != None:
							ht = resulttable.cell(xcord + 1, col_PMS).paragraphs[0]
							ht.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
							resulttable.cell(xcord + 1, col_PMS).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
							if str(xx.pms) == '0' or str(xx.pms) == '': xx.pms = '-'
							r0 = ht.add_run(str(xx.pms))

						if col_name != None:
							ht = resulttable.cell(xcord + 1, col_name).paragraphs[0]
							ht.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
							resulttable.cell(xcord + 1, col_name).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
							r0 = ht.add_run(xx.name)

						if col_val != None:
							ht = resulttable.cell(xcord + 1, col_val).paragraphs[0]
							ht.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
							resulttable.cell(xcord + 1, col_val).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
							if xx.VSG == False:
								try:

									#r0 = ht.add_run('.%3f' % txt_result.replace(".", ","))
									r0 = ht.add_run((str('%.3f' % xx.maxval).replace('.', ',')))


								except:
									pass
									#print(xx.maxval)
							else:
								txt_result = str(round(xx.maxval, 1))
								r0 = ht.add_run(str(txt_result.replace(".", ",")) + '(VSG)')

						if col_limit != None:
							ht = resulttable.cell( xcord + 1 ,col_limit ).paragraphs[ 0 ]
							ht.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
							resulttable.cell(xcord + 1, col_limit).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
							txt_result = str( round( xx.Dlimit ,1 ) )
							r0 = ht.add_run( txt_result.replace( "." ,"," ) )

						if col_class != None:
							ht = resulttable.cell(xcord + 1, col_class).paragraphs[0]
							ht.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
							resulttable.cell(xcord + 1, col_class).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
							r0 = ht.add_run(str(xx.limit))
							colorlimit(r0, str(xx.limit))

						if col_env != None:
							ht = resulttable.cell(xcord + 1, col_env).paragraphs[0]
							ht.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
							resulttable.cell(xcord + 1, col_env).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
							try:
								xx.maxenv = (str('%.3f' % xx.maxenv).replace('.', ','))

							except:
								xx.maxenv = xx.maxenv
							r0 = ht.add_run(str(xx.maxenv))

						if col_trend != None:


							if str( xx.limit ) == 'Cl. D':  # TEGO DLA CZYTELNOSCI LEPIEJ ZROBIC FUNKCJE
								ht = resulttable.cell(xcord + 1, col_trend).add_paragraph()
								ht.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
								resulttable.cell(xcord + 1, col_trend).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
								try:
									t1 = resulttable.cell(xcord + 1, col_trend).paragraphs[0].add_run('_')
									t1.font.size = Pt(1)
									t1.font.color.rgb = RGBColor(255, 255, 255)
									trendtemp1 = float(xx.trend[ 0 ])
									trendtemp2 = xx.trend[ 1 ]
									trendtemp3 = xx.trend[ 2 ]
									if str( trendtemp3 ) == 'U':
										r0 = ht.add_run()
										r0.add_picture( 'up.gif' )
									elif str( trendtemp3 ) == 'D':
										r0 = ht.add_run()
										r0.add_picture( 'down.gif' )
									elif str( trendtemp3 ) == 'C':
										r0 = ht.add_run()
										r0.add_picture( 'none.gif' )
									text_trendtemp1 = str ( '%.3f'%trendtemp1 )
									p0 = ht.add_run( '\n' + 'Last value:' )
									font = p0.font
									font.size = Pt(8)
									font.name = 'Arial'
									p0 = ht.add_run( '\n' + str( trendtemp2 ) )
									font = p0.font
									font.size = Pt(8)
									font.name = 'Arial'
									p0 = ht.add_run( '\n' +  text_trendtemp1.replace('.',',')  )
									font = p0.font
									font.size = Pt(8)
									font.name = 'Arial'
									font.bold = True
									t1 = resulttable.cell(xcord + 1, col_trend).add_paragraph().add_run('_')
									t1.font.size = Pt(1)
									t1.font.color.rgb = RGBColor(255,255,255)
								except:
									pass

						#REMARKS
						ht = resulttable.cell(xcord + 1, col_remark).paragraphs[0]
						ht.add_run(' ')
						resulttable.cell(xcord + 1, col_remark).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
						ht.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


						#TRENDCHART

						#if xcord >= rowscount: break

						if reporttype != 1 and reporttype != 2 and reporttype != 7 and reporttype != 8:
							#try:

								try:
									nextid = str(drivenByList[p + 1])
								except:
									nextid = -1
								if nextid == str(xx.id):


									xcord += 1
									break
								else:
									resulttable.add_row()
									xcord += 1
									resulttable.cell(xcord + 1, 0).merge(resulttable.cell(xcord + 1, len(widths)-1))
									#dd = resulttable.cell(xcord + 1, 0).add_paragraph()

									resulttable.cell(xcord + 1, 0).paragraphs[0].add_run(' ')


									ht = resulttable.cell(xcord + 1, 0).add_paragraph()

									image = (trendchart(ids, connD, GUI, xx.VSG,maxdate).giveimage())

									p0 = ht.add_run()
									p0.add_picture(image, width=Inches(6.5))
									resulttable.cell(xcord + 1, 0).add_paragraph().add_run(' ')
									ids.clear()
									xcord += 1
								break
							#except:  # OSTATNIA LINIA TABELI
								#print('last line error')
							# 	try:
							# 		xcord += 1
							# 		resulttable.cell(xcord + 1, 0).merge(resulttable.cell(xcord + 1, len(widths)-1))
							# 		resulttable.cell(xcord + 1, 0).add_paragraph()
							# 		ht = resulttable.cell(xcord + 1, 0).paragraphs[1]
							# 		image = (trendchart(ids, connD, GUI, xx.VSG).giveimage())
							# 		p0 = ht.add_run()
							# 		p0.add_picture(image, width=Inches(6.5))
							# 		resulttable.cell(xcord + 1, 0).add_paragraph()
							# 		ids.clear()
							# 	except:
							# 		pass
						else:
							xcord += 1
					# except:
					# 	print('oor')
	ht = document.add_paragraph()
	r0 = ht.add_run()

	r = 0

	for row in tqdm(resulttable.rows):

		c = 0
		for cell in row.cells:

			paragraphs = cell.paragraphs
			for paragraph in paragraphs:
				for run in paragraph.runs:
					#try:
						cl = resulttable.cell(r, c)
						#print('row:' ,r,' column',c,cl.text,' merged: ',  cl._tc.right - cl._tc.left)
						#print(c,len(row.cells))
						if c == col_remark  and  cl._tc.right - cl._tc.left == 1 and r > 0:
							#print('last')
							#if cl._tc.right - cl._tc.left == 1:
								#print('no merge')
								font = run.font
								font.size = Pt(9)
								font.name = 'Times New Roman'

						elif cl._tc.right - cl._tc.left == len(row.cells) and font.bold != True :

							font = run.font

							font.size = Pt(1)
						elif c == col_trend and r > 0 :
							pass
						elif cl._tc.right - cl._tc.left > 1:
							font  = run.font
							font.size = Pt(9)
							font.name = 'Calibri'
						else:
							font = run.font
							font.size = Pt(8)
							font.name = 'Arial'
					# except:
					# 	pass
			c += 1
		r += 1
	if reporttype != 1 and reporttype != 2 and reporttype != 7 and reporttype != 8:
		pass
	else:
		trendresults(document)

def trendresults ( document ) :
	trendpar = document.add_paragraph ( 'Trend results:' )
	trendpar.runs[ 0 ].bold = True
	trendpar.runs[ 0 ].underline = True
	trendpar.runs[ 0 ].font.size = Pt ( 8 )
	trendpar.runs[ 0 ].font.name = 'Arial'
	trendpar.add_run()
	trendsres = document.add_table ( rows=3 ,cols=2 )
	up = trendsres.cell ( 0 ,0 ).paragraphs[ 0 ].add_run ()
	up.add_picture ( 'C:\\overmind\\Data\\up.gif' )
	zero = trendsres.cell ( 1 ,0 ).paragraphs[ 0 ].add_run ()
	zero.add_picture ( 'C:\\overmind\\Data\\none.gif' )
	down = trendsres.cell ( 2 ,0 ).paragraphs[ 0 ].add_run ()
	down.add_picture ( 'C:\\overmind\\Data\\down.gif' )
	trendsres.cell ( 0 ,1 ).text = 'Whenever new results are increased more than 5% of previous measurements'
	trendsres.cell(0,1).paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.LEFT
	trendsres.cell ( 1 ,1 ).text = 'Whenever new results are in range plus / minus 5% of previous measurements'
	trendsres.cell(0,1).paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.LEFT
	trendsres.cell ( 2 ,1 ).text = 'Whenever new results are reduced more than 5% of previous measurements'
	trendsres.cell(0,1).paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.LEFT
	trendsres.cell ( 0 ,1 ).paragraphs[ 0 ].runs[ 0 ].font.size = Pt ( 8 )
	trendsres.cell ( 0 ,1 ).paragraphs[ 0 ].runs[ 0 ].font.name = 'Arial'
	trendsres.cell ( 1 ,1 ).paragraphs[ 0 ].runs[ 0 ].font.size = Pt ( 8 )
	trendsres.cell ( 1 ,1 ).paragraphs[ 0 ].runs[ 0 ].font.name = 'Arial'
	trendsres.cell ( 2 ,1 ).paragraphs[ 0 ].runs[ 0 ].font.size = Pt ( 8 )
	trendsres.cell ( 2 ,1 ).paragraphs[ 0 ].runs[ 0 ].font.name = 'Arial'
	trendsres.cell(0, 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
	trendsres.cell(0, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
	trendsres.cell(1, 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
	trendsres.cell(1, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
	trendsres.cell(2, 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
	trendsres.cell(2, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
	set_cols_width_trend ( trendsres )