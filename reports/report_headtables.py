from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from report_database import *
from docx.enum.text import WD_BREAK
from docx.enum.text import WD_LINE_SPACING

# używana wszędzie w raportach IM
def standard_info_table ( connD ,document ,rn_ ):
    def getdaterangestr():
        querry = "select min(date),max(date) from measurements_low where raport_number = '" + str(rn_) + "'"
        datestr = list(q_run( connD ,querry ))[0]
        mindate = datestr[0]
        maxdate = datestr[1]
        if str(mindate) == str(maxdate): outputdate = mindate
        else: outputdate = str(mindate) + ' - ' + str(maxdate)
        return outputdate
    username = connD[ 0 ]
    password = connD[ 1 ]
    host = connD[ 2 ]
    querry = "select main.name, shd.imo, main2.name from main right join shipsdata as shd on main.id = shd.shipid right join main as main2 on main.parent = main2.id where main.id = (select shipid from harmonogram where report_number = '" + str(rn_) + "')"
    result = q_run( connD ,querry )


    shipstr = result[ 0 ][ 0 ]

    headtable = document.add_table( rows=2 ,cols=3 )  # trzeba usunąć enter przed
    headtable.style = 'Table Grid'

    headtable.cell( 0 ,0 ).merge( headtable.cell( 0 ,1 ).merge( headtable.cell( 0 ,2 ) ) )
    ht = headtable.cell( 0 ,0 ).paragraphs[  0]

    ht.alignment = WD_ALIGN_PARAGRAPH.CENTER
    headtable.cell(0,0).vertical_alignment=WD_ALIGN_VERTICAL.CENTER
    r0 = ht.add_run('Vibration diagnostic report ')
    #r0.text =
    r0.font.name = 'Calibri'
    r0.font.size = Pt( 12 )
    r1 = ht.add_run( str(rn_) )
    r1.bold = True
    r1.font.name = 'Calibri'

    ht = headtable.cell( 1 ,0 ).paragraphs[ 0 ]
    ht.alignment = WD_ALIGN_PARAGRAPH.LEFT

    r0 = ht.add_run( 'Project: ' )
    r0.font.name = 'Calibri'
    r1 = ht.add_run( shipstr )
    r1.bold = True
    r1.font.name = 'Calibri'

    imostr = str( result[ 0 ][ 1 ] )
    ht = headtable.cell( 1 ,0 ).add_paragraph()
    r0 = ht.add_run( 'IMO no: ' )
    r0.font.name = 'Calibri'
    r1 = ht.add_run( imostr )
    r1.font.name = 'Calibri'
    r1.bold = True

    ownerstr = str( result[ 0 ][ 2 ] )
    ht = headtable.cell( 1 ,0 ).add_paragraph()
    r0 = ht.add_run( 'Ordered by: ' )
    r0.font.name = 'Calibri'
    r1 = ht.add_run( ownerstr )
    r1.bold = True
    r1.font.name = 'Calibri'
    ht = headtable.cell( 1 ,1 ).paragraphs[ 0 ]
    r0 = ht.add_run( 'Date of measurement: ' )
    r0.add_break()

    r0.font.name = 'Calibri'
    r0.add_break ( WD_BREAK . LINE_CLEAR_LEFT)
    ht.alignment = WD_ALIGN_PARAGRAPH.CENTER
    headtable.cell(0,0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    ht.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ht.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
    datestr = getdaterangestr()
    r0 = ht.add_run(str(datestr))
    r0.bold = True
    r0.font.name = 'Calibri'
    headtable.cell ( 1 ,1 ).add_paragraph ()
    ht = headtable.cell( 1 ,2 ).paragraphs[ 0 ]
    ht.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r0 = ht.add_run( 'Place of measurement:' )
    r0.font.name = 'Calibri'

    headtable.cell( 1 ,2 ).add_paragraph()
    ht = headtable.cell( 1 ,2 ).add_paragraph()
    ht.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r0 = ht.add_run( 'During normal operation' )
    r0.bold = True
    r0.font.name = 'Calibri'
    section = document.sections[ 0 ]
    # footer = section.footer
    # footer.add_paragraph( shipstr + ' ' + rn_ )
    # footer.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # footer.vertical_alignment = WD_ALIGN_VERTICAL.TOP



# tabela początkowa kamtro
def standard_Kamtro_table ( document ):
    headtable = document.add_table( rows=3 ,cols=2 )
    headtable.cell( 0 ,0 ).add_paragraph( 'KAMTRO KAMTRO KAMTRO: ' ,'texthead' )
    headtable.cell( 0 ,0 ).alignment = WD_ALIGN_PARAGRAPH.LEFT


def standard_GSR_table ( document ):
    # htstyle.style[0]=WD_STYLE_TYPE.TABLE
    # assert styles[0].type == WD_STYLE_TYPE.PARAGRAPH

    headtable = document.add_table( rows=2 ,cols=3 )
    headtable.style = 'Table Grid'
    headtable.cell( 0 ,0 ).text = 'Raport z pomiaru wibracji'
    proj = headtable.cell( 1 ,0 ).paragraphs[ 0 ].add_run( 'Projekt:' )
    proj.add_break()
    headtable.cell( 1 ,0 ).paragraphs[ 0 ].add_run( 'Kompresory powietrza' ).bold = True
    headtable.cell( 1 ,0 ).paragraphs[ 0 ].alignment = WD_ALIGN_PARAGRAPH.CENTER

    headtable.cell( 0 ,1 ).paragraphs[ 0 ].add_run( rntxt ).bold = True
    headtable.cell( 0 ,1 ).paragraphs[ 0 ].alignment = WD_ALIGN_PARAGRAPH.CENTER

    zlec = headtable.cell( 1 ,1 ).paragraphs[ 0 ].add_run()
    zlec.text = 'Zleceniodawca:'
    zlec.add_break()
    zlec_gsr = headtable.cell( 1 ,1 ).paragraphs[ 0 ].add_run()
    zlec_gsr.text = 'Stocznia Remontowa Gdańsk'
    zlec_gsr.bold = True
    headtable.cell( 1 ,1 ).paragraphs[ 0 ].alignment = WD_ALIGN_PARAGRAPH.CENTER

    headtable.cell( 0 ,2 ).paragraphs[ 0 ].add_run( 'Data pomiaru:' )
    headtable.cell( 0 ,2 ).paragraphs[ 0 ].add_run().add_break()

    headtable.cell( 0 ,2 ).paragraphs[ 0 ].alignment = WD_ALIGN_PARAGRAPH.CENTER
    date = headtable.cell( 0 ,2 ).paragraphs[ 0 ].add_run( 'Tu wstaw date pomiaru' )
    date.bold = True

    headtable.cell( 1 ,2 ).paragraphs[ 0 ].add_run( 'Miejsce pomiaru:' )
    headtable.cell( 1 ,2 ).paragraphs[ 0 ].add_run().add_break()
    place = headtable.cell( 1 ,2 ).paragraphs[ 0 ].add_run( 'Gdańsk' )
    place.bold = True
    headtable.cell( 1 ,2 ).paragraphs[ 0 ].alignment = WD_ALIGN_PARAGRAPH.CENTER


    print( 'Generowanie tabeli HT' )
