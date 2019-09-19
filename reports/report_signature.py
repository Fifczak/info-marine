import docx,psycopg2
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
def q_run(connD, querry):
	username = connD[0]
	password = connD[1]
	host = connD[2]
	kport = "5432"
	kdb = "postgres"
	# cs = ' host="localhost",database="postgres", user= "postgres" , password="info" '
	cs = "dbname=%s user=%s password=%s host=%s port=%s" % (kdb, username, password, host, kport)
	conn = None
	conn = psycopg2.connect(str(cs))
	cur = conn.cursor()
	cur.execute(querry)
	try:
		result = cur.fetchall()
		return result
	except:
		pass
	conn.commit()
	cur.close()

def getname(connD):
	querry = "select full_name from users where login = '" + str(connD[0]).strip() + "' limit 1"
	try:
		name = list(q_run(connD,querry))[0][0]
		return name
	except:
		return('Get name from database error')

def signatureENG(document,connD):


	H0=document.add_paragraph()
	H0L=H0.add_run("Prepared by:                                                                                                                          Approved by:")
	H0L.alignment=WD_PARAGRAPH_ALIGNMENT.LEFT

	H0L.font.name = 'Calibri'
	signpar = document.add_paragraph ()
	seCAP = signpar.add_run (
		'Service Engineer' )
	seCAP.alignment = WD_ALIGN_PARAGRAPH.LEFT
	seCAP.font.name = 'Calibri'

	name = str ( getname (connD) )
	signpar = document.add_paragraph ()
	nn = signpar.add_run( str ( name ))
	nn.font.name = 'Calibri'



