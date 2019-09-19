from docx.enum.text import WD_BREAK
from docx.enum.text import WD_COLOR_INDEX, WD_ALIGN_PARAGRAPH
from docx.shared import Cm
from docx.shared import Pt
import psycopg2

def q_run(connD, querry):
	username = connD[0]
	password = connD[1]
	host = connD[2]
	kport = "5432"
	kdb = "postgres"
	# cs = ' host="localhost",database="postgres", user= "postgres" , password="info" '
	cs = "dbname=%s user=%s password=%s host=%s port=%s" % (kdb, username, password, host, kport)
	conn = None
	conn = psycopg2.connect(str(cs))
	cur = conn.cursor()
	cur.execute(querry)
	try:
		result = cur.fetchall()
		return result
	except:
		pass
	conn.commit()
	cur.close()

def IMVibEng(document,connD,rn):
	document.add_paragraph()
	#TABELKA Z PARAMETRAMI STATKU
	shiptable = document.add_table(rows=3, cols=2)  # trzeba usunąć enter przed
	shiptable.style = 'Table Grid'
	shiptable.cell(2, 0).merge(shiptable.cell(2, 1))

	querry = "select shiptype,lenght,bradth from shipsdata where shipid = (select shipid from harmonogram where report_number = '" + str(rn) + "' limit 1)"

	shipdata = list(q_run( connD ,querry ))[0]

	##SHIPTABLE [0,0]
	ht = shiptable.cell(0, 0).paragraphs[0]
	r0 = ht.add_run()
	r0.text = 'Ship type:'
	r0.font.name = 'Calibri'
	r0.font.bold = True
	ht = shiptable.cell(0, 0).add_paragraph()
	r0 = ht.add_run()
	r0.text = str(shipdata[0])
	r0.font.name = 'Calibri'

	##SHIPTABLE [0,1]

	ht = shiptable.cell(0, 1).paragraphs[0]
	r0 = ht.add_run()
	r0.text = 'Main dimensions:'
	r0.font.name = 'Calibri'
	r0.font.bold = True
	ht = shiptable.cell(0, 1).add_paragraph()
	r0 = ht.add_run()
	r0.text = 'Length(b.p).......................................' + str("%.2f" % shipdata[1]).replace('.',',') + ' m'
	r0.font.name = 'Calibri'
	ht = shiptable.cell(0, 1).add_paragraph()
	r0 = ht.add_run()
	r0.text = 'Breadth(B.)........................................' + str("%.2f" %shipdata[2]).replace('.',',') + ' m'
	r0.font.name = 'Calibri'

	##SHIPTABLE [1,0]
	ht = shiptable.cell(1, 0).paragraphs[0]
	r0 = ht.add_run()
	r0.text = 'Sea depth:'
	r0.font.name = 'Calibri'
	r0.font.bold = True
	ht = shiptable.cell(1, 0).add_paragraph()
	r0 = ht.add_run()
	r0.text = 'Least twice times greater than Vessel draught'
	r0.font.name = 'Calibri'

	##SHIPTABLE [2,0]
	ht = shiptable.cell(2, 0).paragraphs[0]
	r0 = ht.add_run()
	r0.text = 'Measurement method:'
	r0.font.bold = True
	r0.font.name = 'Calibri'
	ht = shiptable.cell(2, 0).add_paragraph()
	r0 = ht.add_run()
	r0.text = 'According to standard ISO 10816 : - procedure No. 2 Measurement report'
	r0.font.name = 'Calibri'

	document.add_paragraph()
	summary = document.add_paragraph()
	s1 = summary.add_run('Summary:')
	s1.bold = True

	s1.font.name = 'Calibri'
	summary.runs[0].add_break()
	s1 = summary.add_run(
		'Next measurements should be done in three month period to obtain trend value for each equipment, in some cases even one month period is preferable.')
	s1.font.name = 'Calibri'
	s1.add_break()
	s1.add_break()
	s1 = summary.add_run(
		'This report is prepared in good faith based on measurement diagnostic done on available running rotary machine and documentation submitted.')
	s1.font.size = Pt(
		12)
	s1.font.name = 'Calibri'
	document.add_paragraph()