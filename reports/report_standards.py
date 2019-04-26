from docx.enum.text import WD_BREAK
from docx.enum.text import WD_COLOR_INDEX, WD_ALIGN_PARAGRAPH
from docx.shared import Cm
from docx.shared import Pt



def set_col_width(table):
    widths = (Cm ( 1.4 ) ,Cm ( 3 ) ,Cm ( 1 ) ,Cm ( 9 ))
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width=width


def set_col_width_lim ( table ):
    widths = (Cm ( 2 ) ,Cm ( 6 ))
    for row in table.rows:
        for idx ,width in enumerate ( widths ):
            row.cells[ idx ].width = width
def set_col_width_GSR(table):
    widths=(Cm(5.5),Cm(1.5),Cm(1.5),Cm(2),Cm(2),Cm(2),Cm(3.5))
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width=width


def set_col_width_std ( table ):
    widths = (Cm ( 3 ) ,Cm ( 15 ))
    for row in table.rows:
        for idx ,width in enumerate ( widths ):
            row.cells[ idx ].width = width

#limity, trzeba przemyśleć w jaki sposób uzupełniać tabelę
#pomysł: pobierac listę urządzeń z tabeli measurements i wtedy pętla z wyszukiwaniem limitów, break na kolejnym urządzeniu z tym samym limitem

def standards(document):
    H0 = document.add_paragraph('Vibration standards')
    H0.style=document.styles['IM HEAD']
    p0 = document.add_paragraph('Following standards may applied for assessment:').style=document.styles['IM TEXT']
    t_std=document.add_table(6,2)
    std=document.tables[1]


    # norma 8528-9
    ISO_8528=t_std.cell(0,0).text='ISO 8528:9'
    ISO_8528_desc=t_std.cell(0,1).text='Reciprocating internal combustion engine driven alternating current generating sets- Measurement and evaluation of mechanical vibrations.'


    # formatowanie
    # nazwa
    t_std.cell(0, 0).paragraphs[0].runs[0].font.name = 'Calibri'
    t_std.cell(0, 0).paragraphs[0].runs[0].font.size = Pt(11)
    t_std.cell(0, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    t_std.cell(0, 0).paragraphs[0].runs[0].bold = True


    #opis
    t_std.cell(0, 1).paragraphs[0].runs[0].font.name = 'Calibri'
    t_std.cell(0, 1).paragraphs[0].runs[0].font.size = Pt(11)
    t_std.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    t_std.cell(0, 1).paragraphs[0].runs[0].bold = False


    #norma 10816
    ISO_10816 = t_std.cell(1, 0).text = 'ISO 10816'
    ISO_10816_desc = t_std.cell(1, 1).text = 'Evaluation of machine vibration by measurements on non-rotating parts'
    #nazwa

    t_std.cell(1, 0).paragraphs[0].runs[0].font.name = 'Calibri'
    t_std.cell(1, 0).paragraphs[0].runs[0].font.size = Pt(11)
    t_std.cell(1, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    t_std.cell(1, 0).paragraphs[0].runs[0].bold = True

    #opis
    t_std.cell(1, 1).paragraphs[0].runs[0].font.name = 'Calibri'
    t_std.cell(1, 1).paragraphs[0].runs[0].font.size = Pt(11)
    t_std.cell(1, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    t_std.cell(1, 1).paragraphs[0].runs[0].bold = False

    # norma 20816
    ISO_20816=t_std.cell(2,0).text = 'ISO 20816'
    ISO_20816_desc=t_std.cell(2,1).text='Evaluation of machine vibration by measurements on non-rotating parts'

    #formatowanie
    #nazwa
    t_std.cell(2, 0).paragraphs[0].runs[0].font.name = 'Calibri'
    t_std.cell(2, 0).paragraphs[0].runs[0].font.size = Pt(11)
    t_std.cell(2, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    t_std.cell(2, 0).paragraphs[0].runs[0].bold = True

    #opis
    t_std.cell(2, 1).paragraphs[0].runs[0].font.name = 'Calibri'
    t_std.cell(2, 1).paragraphs[0].runs[0].font.size = Pt(11)
    t_std.cell(2, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    t_std.cell(2, 1).paragraphs[0].runs[0].bold = False

    #norma 14694
    ISO_14694=t_std.cell(3,0).text='ISO 14694'
    ISO_14694_desc=t_std.cell(3,1).text='Industrial fans - Specifications for balance quality and vibration levels'


    #formatowanie
    #nazwa
    t_std.cell(3,0).paragraphs[0].runs[0].font.name='Calibri'
    t_std.cell(3,0).paragraphs[0].runs[0].font.size=Pt(11)
    t_std.cell(3, 0).paragraphs[0].runs[0].font.size = Pt(11)
    t_std.cell(3, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    t_std.cell(3, 0).paragraphs[0].runs[0].bold = True

    #opis

    t_std.cell(3, 1).paragraphs[0].runs[0].font.name = 'Calibri'
    t_std.cell(3, 1).paragraphs[0].runs[0].font.size = Pt(11)
    t_std.cell(3, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    t_std.cell(3, 1).paragraphs[0].runs[0].bold = False

    #norma 20283

    ISO_20283=t_std.cell(4,0).text='ISO 20283'
    ISO_20283_desc=t_std.cell(4,1).text='Mechanical vibration- measurement of vibration on ships- parts 1-4'

    #formatowanie
    #nazwa

    t_std.cell(4,0).paragraphs[0].runs[0].font.name='Calibri'
    t_std.cell(4,0).paragraphs[0].runs[0].font.size=Pt(11)
    t_std.cell(4, 0).paragraphs[0].runs[0].font.size = Pt(11)
    t_std.cell(4, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    t_std.cell(4, 0).paragraphs[0].runs[0].bold = True


    #opis

    t_std.cell(4, 1).paragraphs[0].runs[0].font.name = 'Calibri'
    t_std.cell(4, 1).paragraphs[0].runs[0].font.size = Pt(11)
    t_std.cell(4, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    t_std.cell(4, 1).paragraphs[0].runs[0].bold = False

    #norma VDI 3839

    VDI_3839=t_std.cell(5,0).text='VDI 3839'
    VDI_3839_desc=t_std.cell(5,1).text='Addition to ISO 10816-3: Measurement and evaluation of mechanical vibration of screw-type compressors and Roots-blowers'

    #formatowanie

    #nazwa
    t_std.cell(5,0).paragraphs[0].runs[0].font.name='Calibri'
    t_std.cell(5,0).paragraphs[0].runs[0].font.size=Pt(11)
    t_std.cell(5, 0).paragraphs[0].runs[0].font.size = Pt(11)
    t_std.cell(5, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    t_std.cell(5, 0).paragraphs[0].runs[0].bold = True

    #opis

    t_std.cell(5, 1).paragraphs[0].runs[0].font.name = 'Calibri'
    t_std.cell(5, 1).paragraphs[0].runs[0].font.size = Pt(11)
    t_std.cell(5, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    t_std.cell(5, 1).paragraphs[0].runs[0].bold = False
    set_col_width_std ( t_std )

#legenda dla IM 'legend accordin to vib class'
def legend_IM(document):

    document.add_paragraph()
    H0=document.add_paragraph('Legend according to vibration class')
    H0.style=document.styles['IM HEAD']
    H0.add_run()
    legend = document.add_table(5, 2)
    cl_a = legend.cell(0,0).text='Cl. A'
    cl_b = legend.cell(1,0).text='Cl. B'
    cl_c = legend.cell(2,0).text='Cl. C'
    cl_d = legend.cell(3,0).text='Cl. D'
    cl_d_b = legend.cell(4,0).text='Cl. D'


    cl_a_desc=legend.cell(0,1).text='Newly commissioned'
    cl_b_desc=legend.cell(1,1).text='Unrestricted'
    cl_c_desc=legend.cell(2,1).text='Restricted long-term operation'
    cl_d_desc=legend.cell(3,1).text='High probability of damage, action required'
    cl_db_desc=legend.cell(4,1).text='Vibrations over the limits but actions are not required.'

    cl_a_def_format=legend.cell(0,1).paragraphs[0].runs[0]
    cl_a_def_format.font.name='Calibri'
    cl_a_def_format.font.size=Pt(11)

    cl_b_def_format=legend.cell(1,1).paragraphs[0].runs[0]
    cl_b_def_format.font.name='Calibri'
    cl_b_def_format.font.size=Pt(11)

    cl_c_desc_format=legend.cell(2,1).paragraphs[0].runs[0]
    cl_c_desc_format.font.name='Calibri'
    cl_c_desc_format.font.size=Pt(11)


    cl_d_desc_format=legend.cell(3,1).paragraphs[0].runs[0]
    cl_d_desc_format.font.name='Calibri'
    cl_d_desc_format.font.size=Pt(11)

    cl_db_desc_format=legend.cell(4,1).paragraphs[0].runs[0]
    cl_db_desc_format.font.name='Calibri'
    cl_db_desc_format.font.size=Pt(11)

    #kopiemy do runa zeby ustawic highlights
    #trzeba przemyśleć jak nadać jasny zielony, BRIGHTGREEN nie trybi
    cla_format=legend.cell(0,0).paragraphs[0].runs[0]
    clb_format=legend.cell(1,0).paragraphs[0].runs[0]
    clc_format=legend.cell(2,0).paragraphs[0].runs[0]
    cld_format=legend.cell(3,0).paragraphs[0].runs[0]
    cldb_format=legend.cell(4,0).paragraphs[0].runs[0]

    cla_format.font.highlight_color=WD_COLOR_INDEX.GREEN
    cla_format.font.name='Calibri'
    cla_format.font.size=Pt(11)


    clb_format.font.highlight_color=WD_COLOR_INDEX.GREEN
    clb_format.font.name = 'Calibri'
    clb_format.font.size = Pt(11)

    clc_format.font.highlight_color=WD_COLOR_INDEX.YELLOW
    clc_format.font.name='Calibri'
    clc_format.font.size = Pt(11)

    cld_format.font.highlight_color=WD_COLOR_INDEX.RED
    cld_format.font.name='Calibri'
    cld_format.font.size = Pt(11)

    cldb_format.font.name='Calibri'
    cldb_format.font.bold=True
    cldb_format.font.size = Pt(11)


    end=document.add_paragraph()
    er=end.add_run()
    er.add_break(WD_BREAK.PAGE)
    set_col_width_lim ( legend )


#legenda dla kamtro - do dokończenia [odwolanie-if w report.py]
def legend_KAMTRO(document):
    H0 = document.add_paragraph('Legend according to vibration class').style = document.styles['IM HEAD']

def standards_GSR(document):

    H0=document.add_paragraph()
    description=H0.add_run('Zgodnie z ISO 10816-3 silniki elektryczne o mocy przekraczającej 15 kW, o prędkościach obrotowych z zakresu 120- 15 000 obr/min, posiadają następujące limity wartości drgań:')
    H0.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY

    #konwencja:
        #definicja klasy [ClA/B/C/D]
        #jednostka[mm/s]
        #opis

    limittable=document.add_table(4,4)

    Cla=limittable.cell(0,0).text='Cl. A'
    limittable.cell(0,1).text='0 - 3,5'
    limittable.cell(0,2).text='mm/s'
    limittable.cell(0,3).text='nowo oddane'

    Cla_format=limittable.cell(0,0).paragraphs[0].runs[0]
    Cla_format.font.highlight_color=WD_COLOR_INDEX.GREEN

    Clb = limittable.cell(1, 0).text = 'Cl. B'
    limittable.cell(1, 1).text = '3,5 - 7,1'
    limittable.cell(1, 2).text = 'mm/s'
    limittable.cell(1, 3).text = 'bez zastrzeżeń'
    Clb_format=limittable.cell(1,0).paragraphs[0].runs[0]
    Clb_format.font.highlight_color=WD_COLOR_INDEX.GREEN

    Clc = limittable.cell(2, 0).text = 'Cl. C'
    limittable.cell(2,1).text='7,1 - 11,0'
    limittable.cell(2, 2).text = 'mm/s'
    limittable.cell(2, 3).text = 'ograniczone długotrwałe operacje'

    Clc_format=limittable.cell(2,0).paragraphs[0].runs[0]
    Clc_format.font.highlight_color=WD_COLOR_INDEX.YELLOW


    Cld = limittable.cell(3, 0).text = 'Cl. D'
    limittable.cell(3,1).text='11,0 - więcej'
    limittable.cell(3, 2).text = 'mm/s'
    limittable.cell(3, 3).text = 'wysokie prawdopodobieństwo uszkodzenia'

    Cld_format=limittable.cell(3,0).paragraphs[0].runs[0]
    Cld_format.font.highlight_color=WD_COLOR_INDEX.RED

    document.add_paragraph('')
    H1=document.add_paragraph('Zgodnie z dodatkiem VDI 3839 do normy DIN ISO 10816-3 sprężarki śrubowe bez kół synchronizujących o mocy powyżej 37 kW wraz z przekładniami osadzone na sztywnym fundamencie, podlegają ocenie według wartości prędkości drgań zgodnie z podziałem:')
    H2=document.add_paragraph('Dla zakresu A: 10-1000Hz')
    limittable_rangeA=document.add_table(rows=3,cols=4)
    limittable_rangeA.cell(0,0).text='Zone I'
    limittable_rangeA.cell(0,0).paragraphs[0].runs[0].font.highlight_color=WD_COLOR_INDEX.GREEN
    limittable_rangeA.cell(0,1).text='0 - 8'
    limittable_rangeA.cell(0,2).text='mm/s'
    limittable_rangeA.cell(0,3).text='poziom akceptowalny'

    limittable_rangeA.cell(1,0).text='Zone II'
    limittable_rangeA.cell(1 , 0).paragraphs[0].runs[0].font.highlight_color=WD_COLOR_INDEX.YELLOW
    limittable_rangeA.cell(1,1).text='8 - 12'
    limittable_rangeA.cell(1, 2).text = 'mm/s'
    limittable_rangeA.cell(1, 3).text = 'alarm'

    limittable_rangeA.cell(2,0).text='Zone III'
    limittable_rangeA.cell(2 , 0).paragraphs[0].runs[0].font.highlight_color=WD_COLOR_INDEX.RED
    limittable_rangeA.cell(2,1).text=' > 12'
    limittable_rangeA.cell(2, 2).text = 'mm/s'
    limittable_rangeA.cell(2, 3).text = 'ograniczone długotrwałe operacje'

    H3=document.add_paragraph('Dla zakresu B: 10Hz - 2,2 x prędkość obrotowa')

    limittable_rangeB=document.add_table(rows=3,cols=4)
    limittable_rangeB.width=Cm(10)

    limittable_rangeB.cell(0,0).text='Zone I'
    limittable_rangeB.cell(0,0).paragraphs[0].runs[0].font.highlight_color=WD_COLOR_INDEX.GREEN
    limittable_rangeB.cell(0,1).text='0 - 3'
    limittable_rangeB.cell(0,2).text='mm/s'
    limittable_rangeB.cell(0,3).text='poziom akceptowalny'

    limittable_rangeB.cell(1,0).text='Zone II'
    limittable_rangeB.cell(1,0).paragraphs[0].runs[0].font.highlight_color=WD_COLOR_INDEX.YELLOW
    limittable_rangeB.cell(1,1).text='3 - 4,5'
    limittable_rangeB.cell(1,2).text='mm/s'
    limittable_rangeB.cell(1,3).text='alarm'

    limittable_rangeB.cell(2,0).text='Zone III'
    limittable_rangeB.cell(2,0).paragraphs[0].runs[0].font.highlight_color=WD_COLOR_INDEX.RED
    limittable_rangeB.cell(2,1).text=' > 4,5'
    limittable_rangeB.cell(2,2).text='mm/s'
    limittable_rangeB.cell(2,3).text='ograniczone długotrwałe operacje'



    h4=document.add_paragraph()
    run=h4.add_run()
    run.add_break(WD_BREAK.PAGE)
    set_col_width(limittable)
    set_col_width(limittable_rangeA)
    set_col_width(limittable_rangeB)
