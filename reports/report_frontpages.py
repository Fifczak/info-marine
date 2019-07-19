from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

#standard IM report (PMS)
def standard_with_PMS(document):
	#enter po tabeli
	document.add_paragraph()
	# header nr 1
	H1 = document.add_paragraph('Measurement condition')
	H1.style = document.styles['IM HEAD']
	H1.alignment = WD_ALIGN_PARAGRAPH.LEFT
	T1 = document.add_paragraph('Measurements were taken during normal operating condition.')
	T1.style = document.styles['IM TEXT']
	T1.alignment = WD_ALIGN_PARAGRAPH.LEFT

	document.add_paragraph()

	#header nr2
	H2=document.add_paragraph('Results presentation')
	H2.style = document.styles['IM HEAD']

	h3=document.add_paragraph()
	h3.style = document.styles['IM TEXT']
	h3.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
	r1=h3.add_run('Measured values are presented in the table below. Each machine if applicable is separated for driver (el. motor, diesel engine, etc.) and driven unit (pump, compressor, etc.) ')
	r2=h3.add_run('First and second columns')
	r2.italic=True
	r2.underline=True
	r2 = h3.add_run ( ' of the table consist PMS number and name of the equipment. ' )
	r3=h3.add_run('Third column')
	r3.italic=True
	r3.underline=True
	r4=h3.add_run(' contains the highest value of vibration velocity measured on the equipment in all measurement points. ')
	r5=h3.add_run('Fourth column')
	r5.italic=True
	r5.underline=True
	r6=h3.add_run(' contains ISO classification limit. ')
	r7=h3.add_run('Fifth column')
	r7.underline=True
	r7.italic=True
	r8=h3.add_run(' contains additional readings of enveloped value of acceleration, which is helpful in detection of early stage of bearing wear.  ')
	r9=h3.add_run('Sixth column')
	r9.italic=True
	r9.underline=True
	r10=h3.add_run('contains vibration trend values if previous results are available from the same source. ')
	r11=h3.add_run('Seventh column')
	r11.underline=True
	r11.italic=True
	r12=h3.add_run('contains remarks and suggestions based on the analysis of vibration signal. This column can be taken as the final conclusion about machine condition. If cell is empty, it means that there is no existing problem or defect shown in vibration signal.')
#standard IM report (non PMS)
def standard_without_PMS(document):
	document.add_paragraph()
	H1 = document.add_paragraph('Measurement condition')
	H1.style = document.styles['IM HEAD']
	H1.alignment = WD_ALIGN_PARAGRAPH.LEFT
	T1 = document.add_paragraph('Measurements were taken during normal operating condition.')
	T1.style = document.styles['IM TEXT']
	T1.alignment = WD_ALIGN_PARAGRAPH.LEFT
	document.add_paragraph()
	H2 = document.add_paragraph('Results presentation')
	H2.style = document.styles['IM HEAD']
	H2.alignment = WD_ALIGN_PARAGRAPH.LEFT

	T1 = document.add_paragraph()
	r0 = T1.add_run(
			'Measured values are presented in the table below. Each machine if applicable is separated for driver (el. motor, diesel engine, etc.) and driven unit (pump, compressor, etc.). ')
	r1 = T1.add_run('First column')
	r1.italic = True
	r1.underline = True
	r2 = T1.add_run(' of the table consist name of the equipment. ')
	r1 = T1.add_run('Second column')
	r1.italic = True
	r1.underline = True
	r2 = T1.add_run(
			' contains the highest value of vibration velocity measured on the equipment in all measurement points. ')
	r1 = T1.add_run('Third column')
	r1.italic = True
	r1.underline = True
	r2 = T1.add_run(
			' contains classification of the vibration class according to proper ISO standard and other normative documents. Classification depends on highest reading of measured equipment only. ')
	r1 = T1.add_run('Fourth column')
	r1.italic = True
	r1.underline = True
	r2 = T1.add_run(
			' contains additional readings of enveloped value of acceleration, which is helpful in detection of early stage of bearing wear. ')
	r1 = T1.add_run('Fifth column')
	r1.italic = True
	r1.underline = True
	r2 = T1.add_run(' contains vibration trend values if previous results are available from the same source. ')
	r1 = T1.add_run('Sixth column')
	r1.italic = True
	r1.underline = True
	r2 = T1.add_run(
			' contains remarks and suggestions based on the analysis of vibration signal. This column can be taken as the final conclusion about machine condition. If cell is empty, it means that there is no existing problem or defect shown in vibration signal.')
	T1.style = document.styles['IM TEXT']
	T1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
#donnelly with PMS
def donnelly_with_PMS(document):
	document.add_paragraph ()
	H1 = document.add_paragraph ( 'Measurement condition' )
	H1.style = document.styles [ 'IM HEAD' ]
	H1.alignment = WD_ALIGN_PARAGRAPH.LEFT
	T1 = document.add_paragraph ( 'Measurements were taken during normal operating condition.' )
	T1.style = document.styles [ 'IM TEXT' ]
	T1.alignment = WD_ALIGN_PARAGRAPH.LEFT
	document.add_paragraph ()
	H2 = document.add_paragraph ( 'Results presentation' )
	H2.style = document.styles [ 'IM HEAD' ]
	H2.alignment = WD_ALIGN_PARAGRAPH.LEFT

	T1 = document.add_paragraph ()
	r0 = T1.add_run (
		'Measured values are presented in the table below. Each machine if applicable is separated for driver (el. motor, diesel engine, etc.) and driven unit (pump, compressor, etc.). ' )
	r1 = T1.add_run ( 'First and second columns' )
	r1.italic = True
	r1.underline = True
	r2 = T1.add_run ( ' of the table consist PMS number and name of the equipment. ' )
	r1 = T1.add_run ( 'Third column' )
	r1.italic = True
	r1.underline = True
	r2 = T1.add_run (
		' contains the highest value of vibration velocity measured on the equipment in all measurement points. ' )
	r1 = T1.add_run ( 'Fourth column' )
	r1.italic = True
	r1.underline = True
	r2 = T1.add_run (
		' contains classification of the vibration class according to proper ISO standard and other normative documents. Classification depends on highest reading of measured equipment only. ' )
	r1 = T1.add_run ( 'Fifth column' )
	r1.italic = True
	r1.underline = True
	r2 = T1.add_run (
		' contains additional readings of enveloped value of acceleration, which is helpful in detection of early stage of bearing wear. ' )
	r1 = T1.add_run ( 'Sixth column' )
	r1.italic = True
	r1.underline = True
	r2 = T1.add_run ( ' contains remarks and suggestions based on the analysis of vibration signal. This column can be taken as the final conclusion about machine condition. If cell is empty, it means that there is no existing problem or defect shown in vibration signal.' )

def donnelly_without_PMS(document):
	document.add_paragraph()
	H1 = document.add_paragraph('Measurement condition')
	H1.style = document.styles['IM HEAD']
	H1.alignment = WD_ALIGN_PARAGRAPH.LEFT
	T1 = document.add_paragraph('Measurements were taken during normal operating condition.')
	T1.style = document.styles['IM TEXT']
	T1.alignment = WD_ALIGN_PARAGRAPH.LEFT
	document.add_paragraph()
	H2 = document.add_paragraph('Results presentation')
	H2.style = document.styles['IM HEAD']
	H2.alignment = WD_ALIGN_PARAGRAPH.LEFT

	T1 = document.add_paragraph()
	r0 = T1.add_run(
			'Measured values are presented in the table below. Each machine if applicable is separated for driver (el. motor, diesel engine, etc.) and driven unit (pump, compressor, etc.). ')
	r1 = T1.add_run('First column')
	r1.italic = True
	r1.underline = True
	r2 = T1.add_run(' of the table consist name of the equipment. ')
	r1 = T1.add_run('Second column')
	r1.italic = True
	r1.underline = True
	r2 = T1.add_run(
			' contains the highest value of vibration velocity measured on the equipment in all measurement points. ')
	r1 = T1.add_run('Third column')
	r1.italic = True
	r1.underline = True
	r2 = T1.add_run(
			' contains classification of the vibration class according to proper ISO standard and other normative documents. Classification depends on highest reading of measured equipment only. ')
	r1 = T1.add_run('Fourth column')
	r1.italic = True
	r1.underline = True
	r2 = T1.add_run(
			' contains additional readings of enveloped value of acceleration, which is helpful in detection of early stage of bearing wear. ')
	r1 = T1.add_run('Fifth column')
	r1.italic = True
	r1.underline = True
	r2 = T1.add_run(
			' contains remarks and suggestions based on the analysis of vibration signal. This column can be taken as the final conclusion about machine condition. If cell is empty, it means that there is no existing problem or defect shown in vibration signal.')
	r2 = T1.add_run()
	T1.style = document.styles['IM TEXT']
	T1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

#siem z limitem i PMS + wykresy (impro)
def siem_PMS_limit(document):
	document.add_paragraph()
	H0=document.add_paragraph("Measurement condition")
	H0.style=document.styles['IM HEAD']
	document.add_paragraph('Measurements were taken during normal condition.')
	document.add_paragraph()
	H1=document.add_paragraph('Results presentation')
	H1.style=document.styles['IM HEAD']
	document.add_paragraph('Measured values are presented in the table below. Each machine if applicable is separated for driver (el. motor, diesel engine, etc.) and driven unit (pump, compressor, etc.). ')
	p1=document.add_paragraph()
	r1=p1.add_run('First and second column')
	r1.italic=True
	r1.underline=True
	r2=p1.add_run(' of the table consist PMS number and name of the equipment. ')
	r3=p1.add_run('Third column')
	r3.underline=True
	r3.italic=True
	r4=p1.add_run(' contains the highest value of vibration velocity measured on the equipment in all measurement points. ')
	r5=p1.add_run('Fourth column')
	r5.underline=True
	r5.italic=True
	r6=p1.add_run(' contains ISO classification limit. ')
	r7=p1.add_run('Fifth column')
	r7.underline=True
	r7.italic=True
	r8=p1.add_run(' contains classification of vibration class according to proper ISO standard and other normative documents. Classification depends on highest reading of measured equipment only. ')
	r9=p1.add_run('Sixth column')
	r9.italic=True
	r9.underline=True
	r10=p1.add_run(' contains additional readings of enveloped value of acceleration, which is helpful in detection of early stage of bearing wear. ')
	r11=p1.add_run('Seventh column')
	r11.underline=True
	r11.italic=True
	r12=p1.add_run(' contains remarks and suggestions based on the analysis of vibration signal. This column can be taken as the final conclusion about machine condition. If cell is empty, it means that there is no existing problem or defect shown in vibration signal. ')
	r14=p1.add_run()
	p1.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
#siem z limitem bez PMS

def standard_non_PMS_limit(document):   ## TUTAJ TRZEBA POPRAWIC NA OFPOWIEDNIA FORMATKE. USUNĄŁEM PMS ale widze ze niema np max limitu ? do przejzenia
	document.add_paragraph()
	H0=document.add_paragraph('Measurement condition')
	H0.style=document.styles['IM HEAD']
	document.add_paragraph('Measurements were taken during normal condition.')
	document.add_paragraph()
	H1=document.add_paragraph('Results presentation')
	H1.style=document.styles['IM HEAD']
	p1=document.add_paragraph()
	r0=p1.add_run("Measured values are presented in the table below. Each machine if applicable is separated for driver (el. motor, diesel engine, etc.) and driven unit (pump, compressor, etc.). ")
	r1=p1.add_run('First column')
	r1.underline=True
	r1.italic=True
	r2=p1.add_run(' of the table consist name of the equipment. ')
	r3=p1.add_run('Second column')
	r3.underline=True
	r3.italic=True
	r4=p1.add_run(' contains the highest value of vibration velocity measured on the equipment in all measurement points. ')
	r5=p1.add_run('Third column')
	r5.italic=True
	r5.underline=True
	r6=p1.add_run(' contains ISO classification limit. ')
	r7=p1.add_run('Fourth column')
	r7.underline=True
	r7.italic=True
	r8=p1.add_run(' contains classification of the vibration class according to proper ISO standard and other normative documents. Classification depends on highest reading of measured equipment only. ')
	r9=p1.add_run('Fifth column')
	r9.italic=True
	r9.underline=True
	r10=p1.add_run(' contains additional readings of enveloped value of acceleration, which is helpful in detection of early stage of bearing wear. ')
	r11=p1.add_run('Sixth column')
	r11.underline=True
	r11.italic=True
	r12=p1.add_run(' contains remarks and suggestions based on the analysis of vibration signal. This column can be taken as the final conclusion about machine condition. If cell is empty, it means that there is no existing problem or defect shown in vibration signal.')
	r12=p1.add_run()
	r12.add_break()
	p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
#siem z PMS i LIMITEM
#poniżej improwizacja - nie znalazłem raportu dla jakiegokolwiek siema z PMS



def standard_GSR(document):
	document.add_paragraph()
	H0=document.add_paragraph('Warunki pomiaru wibracji')
	H0.runs[0].add_break()
	H0.add_run('Wibracje mierzone w Stoczni Remontowej w Gdańsku. Przedmiotem pomiaru są kompresory powietrza firmy Boge.')
	H0.runs[1].add_break()
	H0.add_run('Mierzone punkty są podzielone na źródła mocy oraz elementy wykonawcze.')
	document.add_picture('C:\Overmind\data\screw_comp.png')
	H1=document.add_paragraph()
	pict_desc=H1.add_run('(Typowa maszyna z rozmieszczonymi punktami pomiarowymi)')
	pict_desc.font.size=Pt(10)
def charts_and_limits(document):
	document.add_paragraph()
	H0=document.add_paragraph("Measurement condition")
	H0.style=document.styles['IM HEAD']
	document.add_paragraph('Measurements were taken during normal condition.')
	document.add_paragraph()
	H1=document.add_paragraph('Results presentation')
	H1.style=document.styles['IM HEAD']
	document.add_paragraph('Measured values are presented in the table below. Each machine if applicable is separated for driver (el. motor, diesel engine, etc.) and driven unit (pump, compressor, etc.). ')
	p1=document.add_paragraph()
	r1=p1.add_run('First and second column')
	r1.italic=True
	r1.underline=True
	r2=p1.add_run(' of the table consist PMS number and name of the equipment. ')
	r3=p1.add_run('Third column')
	r3.underline=True
	r3.italic=True
	r4=p1.add_run(' contains ISO classification limit. ')
	r5=p1.add_run('Fourth column')
	r5.underline=True
	r5.italic=True
	r6=p1.add_run(' contains the highest value of vibration velocity measured on the equipment in all measurement points. ')
	r7=p1.add_run('Fifth column')
	r7.underline=True
	r7.italic=True
	r8=p1.add_run(' contains classification of vibration class according to proper ISO standard and other normative documents. Classification depends on highest reading of measured equipment only. ')
	r9=p1.add_run('Sixth column')
	r9.italic=True
	r9.underline=True
	r10=p1.add_run(' contains additional readings of enveloped value of acceleration, which is helpful in detection of early stage of bearing wear. ')
	r11=p1.add_run('Seventh column')
	r11.underline=True
	r11.italic=True
	r12=p1.add_run(' contains vibration trend values if previous results are available from the same source. ')
	r13=p1.add_run('Eighth column')
	r13.underline=True
	r13.italic=True
	r14=p1.add_run(' contains remarks and suggestions based on the analysis of vibration signal. This column can be taken as the final conclusion about machine condition. If cell is empty, it means that there is no existing problem or defect shown in vibration signal. ')
	p1.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
def charts_no_limits(document):
	document.add_paragraph()
	H0=document.add_paragraph("Measurement condition")
	H0.style=document.styles['IM HEAD']
	document.add_paragraph('Measurements were taken during normal condition.')
	document.add_paragraph()
	H1=document.add_paragraph('Results presentation')
	H1.style=document.styles['IM HEAD']
	document.add_paragraph('Measured values are presented in the table below. Each machine if applicable is separated for driver (el. motor, diesel engine, etc.) and driven unit (pump, compressor, etc.). ')
	p1=document.add_paragraph()
	r1=p1.add_run('First and second column')
	r1.italic=True
	r1.underline=True
	r2=p1.add_run(' of the table consist PMS number and name of the equipment. ')
	r3=p1.add_run('Third column')
	r3.underline=True
	r3.italic=True
	r4=p1.add_run(' contains ISO classification limit. ')
	r5=p1.add_run('Fourth column')
	r5.underline=True
	r5.italic=True
	r6=p1.add_run(' contains the highest value of vibration velocity measured on the equipment in all measurement points. ')
	r7=p1.add_run('Fifth column')
	r7.underline=True
	r7.italic=True
	r8=p1.add_run(' contains classification of vibration class according to proper ISO standard and other normative documents. Classification depends on highest reading of measured equipment only. ')
	r9=p1.add_run('Sixth column')
	r9.italic=True
	r9.underline=True
	r10=p1.add_run(' contains additional readings of enveloped value of acceleration, which is helpful in detection of early stage of bearing wear. ')
	r11=p1.add_run('Seventh column')
	r11.underline=True
	r11.italic=True
	r12=p1.add_run(' contains vibration trend values if previous results are available from the same source. ')
	r13=p1.add_run('Eighth column')
	r13.underline=True
	r13.italic=True
	r14=p1.add_run(' contains remarks and suggestions based on the analysis of vibration signal. This column can be taken as the final conclusion about machine condition. If cell is empty, it means that there is no existing problem or defect shown in vibration signal. ')
	p1.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY