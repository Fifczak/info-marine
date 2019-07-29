import sys
import datetime
import os
import os.path
import time
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_COLOR_INDEX
import matplotlib.pyplot as plt

from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.shared import RGBColor

from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_BREAK
from docx.enum.text import WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_TAB_ALIGNMENT ,WD_TAB_LEADER

from tkinter import messagebox
from tkinter import *
from tkinter.ttk import *

from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

import numpy as np

import psycopg2

tk = Tk ()
document = Document ( 'C:\\overmind\\Data\\base.docx' )
def set_col_widths(table):
    widths = (Inches(0.6), Inches(5))
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width

def q_run ( connD ,querry ):
    username = connD [ 0 ]
    password = connD [ 1 ]
    host = connD [ 2 ]
    kport = "5432"
    kdb = "postgres"
    # cs = ' host="localhost",database="postgres", user= "postgres" , password="info" '
    cs = "dbname=%s user=%s password=%s host=%s port=%s"%(kdb ,username ,password ,host ,kport)
    conn = None
    conn = psycopg2.connect ( str ( cs ) )
    cur = conn.cursor ()
    cur.execute ( querry )
    try:
        result = cur.fetchall ()
        return result
    except:
        pass
    conn.commit ()
    cur.close ()


######################################
#################STYLE################
######################################

obj_styles = document.styles
obj_charstyle = obj_styles.add_style ( 'Head' ,WD_STYLE_TYPE.PARAGRAPH )
style = document.styles [ 'Head' ]
font = style.font
font.name = 'Calibri'
font.size = Pt ( 16 )
font.bold = True
font.color.rgb = RGBColor ( 68 ,114 ,196 )

obj_charstyle = obj_styles.add_style ( 'listlvl1' ,WD_STYLE_TYPE.PARAGRAPH )
style = document.styles [ 'listlvl1' ]
font = style.font
font.name = 'Calibri Light'
font.size = Pt ( 16 )
font.bold = True
font.color.rgb = RGBColor ( 47 ,84 ,150 )

obj_charstyle = obj_styles.add_style ( 'listlvl2' ,WD_STYLE_TYPE.PARAGRAPH )
style = document.styles [ 'listlvl2' ]
font = style.font
font.name = 'Calibri Light'
font.size = Pt ( 13 )
font.bold = True
font.color.rgb = RGBColor ( 47 ,84 ,150 )

obj_charstyle = obj_styles.add_style ( 'texthead' ,WD_STYLE_TYPE.PARAGRAPH )
style = document.styles [ 'texthead' ]
font = style.font
font.name = 'Calibri Light'
font.size = Pt ( 11 )
font.bold = True

obj_charstyle = obj_styles.add_style ( 'textheadU' ,WD_STYLE_TYPE.PARAGRAPH )
style = document.styles [ 'textheadU' ]
font2 = style.font
font2.name = 'Calibri Light'
font2.size = Pt ( 11 )
font2.bold = True
font2.underline = True

obj_charstyle = obj_styles.add_style ( 'textt1' ,WD_STYLE_TYPE.PARAGRAPH )
style = document.styles [ 'textt1' ]
font = style.font
font.name = 'Calibri'
font.size = Pt ( 11 )

obj_charstyle = obj_styles.add_style ( 'texBtt1' ,WD_STYLE_TYPE.PARAGRAPH )
style = document.styles [ 'texBtt1' ]
font = style.font
font.name = 'Calibri'
font.size = Pt ( 11 )
font.bold = True

obj_charstyle = obj_styles.add_style ( 'warning' ,WD_STYLE_TYPE.PARAGRAPH )
style = document.styles [ 'warning' ]
font = style.font
font.name = 'Calibri Light'
font.size = Pt ( 13 )

obj_charstyle = obj_styles.add_style ( 'description' ,WD_STYLE_TYPE.PARAGRAPH )
style = document.styles [ 'description' ]
font = style.font
font.name = 'Calibri'
font.size = Pt ( 9 )
font.italic = True


######################################


######################################
#######DEFINICJE FUNKCJI#############
######################################
def createchart ( host ,username ,password ,id ,rn ):
    connD = [ username ,password ,host ]

    querry = "select point from measurements_low where id = " + str ( id ) + " and raport_number = '" + str (
        rn ) + "' order by value desc limit 1;"
    maxPoint = q_run ( connD ,querry ) [ 0 ] [ 0 ]

    X = [ ]
    Y = [ ]
    x = [ ]
    try:
        querry = "select chart from meascharts where id = '" + str ( id ) + "' and report_number = '" + str (
            rn ) + "' and point = '" + str ( maxPoint ) + "' and domain = 'FFT' and type = 'Vel' ;"
        chartstr = q_run ( connD ,querry ) [ 0 ] [ 0 ]
        chartstrlist = chartstr.split ( ";" )
        for line in chartstrlist:
            x.append ( float ( line ) )

    except:
        if host == 'localhost':
            querry = "select lo_export(measurements.chart, 'C:\\Overmind\\temp\\tempchart.csv') from measurements where id = " + str (
                id ) + " and report_number = '" + str ( rn ) + "';"
        if host == '192.168.8.125':
            servpath = r'/home/filip/Public/tempchart.csv'
            querry = "select lo_export(measurements.chart,'" + servpath + "') from measurements where id = '" + str (
                id ) + "' and report_number = '" + str ( rn ) + "' and point = '" + str ( result [ 0 ] [ 0 ] ) + "';"

        q_run ( connD ,querry )
        # os.remove(spath_)

        if host == 'localhost':
            while not os.path.exists ( r'C:\\Overmind\\temp\\tempchart.csv' ):
                time.sleep ( 1 )
            try:
                x = np.loadtxt ( r'C:\\Overmind\\temp\\tempchart.csv' )
            except:
                messagebox.showinfo ( "Title" ,str ( querry ) )
        if host == '192.168.8.125':
            servpath = r'\\192.168.8.125\Public\tempchart.csv'
            while not os.path.exists ( servpath ):
                time.sleep ( 1 )
            try:
                x = np.loadtxt ( servpath )
            except:
                messagebox.showinfo ( "Title" ,str ( querry ) )

    xlen = int ( len ( x ) - 2 )
    tp = x [ 0 ]
    dt = (x [ 1 ]/xlen)
    x = x [ 2: ]
    it = 0
    t = tp
    maxf = 0

    if id == 17159 or id == 17161 or id == 17162 or id == 17163 or id == 17164 or id == 17165 or id == 18137 or id == 18127 or id == 18129 or id == 18131 or id == 18133 or id == 18135:
        maxf = 1000  # motor
        minf = 2
    if id == 17160 or id == 17166 or id == 17167 or id == 17168 or id == 17169 or id == 17170 or id == 18138 or id == 18128 or id == 18130 or id == 18132 or id == 18134 or id == 18136:
        maxf = 1000  # gear
        minf = 5
    for i in x:
        if t < minf:
            x [ it ] = 0
        if t > maxf:
            x [ it ] = 0
        it += 1
        t += dt
        Y.append ( t )

    fig = plt.figure ( figsize=(15 ,6) ,dpi=80 )
    plt.cla ()
    plt.plot ( Y ,x ,linewidth=0.3 )
    plt.xlabel ( 'Frequency [Hz]' )
    plt.ylabel ( 'Velocity[mm/s]' )
    axes = plt.gca ()

    maxval = np.amax ( x )
    maxcord = (np.argmax ( x ) + 2)*dt
    plt.plot ( [ 0 ,1000 ] ,[ maxval ,maxval ] ,lw=0.3 ,color='red' )
    plt.plot ( [ maxcord ,maxcord ] ,[ maxval ,maxval*1.1 ] ,lw=0.3 )
    axes.set_xlim ( [ 0 ,maxf ] )
    axes.set_ylim ( [ 0 ,maxval*1.1 ] )
    plt.savefig ( 'C:\\overmind\\temp\\tempchart.png' ,dpi=300 ,width=10 ,color='black' )

    try:
        if host == 'localhost':
            os.remove ( 'C:\\Overmind\\temp\\tempchart.csv' )
        if host == '192.168.8.125':
            servpath = r'\\192.168.8.125\Public\tempchart.csv'
            os.remove ( servpath )
    except:
        pass

    return maxPoint


def resulttable ( tn ,host ,username ,password ,tableno ,id1 ,id2 ,rn ):
    connD = [ username ,password ,host ]
    kport = "5432"
    kdb = "postgres"

    # cs = ' host="localhost",database="postgres", user= "postgres" , password="info" '
    cs = "dbname=%s user=%s password=%s host=%s port=%s"%(kdb ,username ,password ,host ,kport)
    conn = None
    conn = psycopg2.connect ( str ( cs ) )

    shading_elm_1 = parse_xml ( r'<w:shd {} w:fill="D3D3D3"/>'.format ( nsdecls ( 'w' ) ) )
    shading_elm_2 = parse_xml ( r'<w:shd {} w:fill="D3D3D3"/>'.format ( nsdecls ( 'w' ) ) )
    shading_elm_3 = parse_xml ( r'<w:shd {} w:fill="D3D3D3"/>'.format ( nsdecls ( 'w' ) ) )
    shading_elm_4 = parse_xml ( r'<w:shd {} w:fill="D3D3D3"/>'.format ( nsdecls ( 'w' ) ) )
    shading_elm_5 = parse_xml ( r'<w:shd {} w:fill="D3D3D3"/>'.format ( nsdecls ( 'w' ) ) )
    shading_elm_6 = parse_xml ( r'<w:shd {} w:fill="D3D3D3"/>'.format ( nsdecls ( 'w' ) ) )
    shading_elm_7 = parse_xml ( r'<w:shd {} w:fill="D3D3D3"/>'.format ( nsdecls ( 'w' ) ) )
    shading_elm_8 = parse_xml ( r'<w:shd {} w:fill="D3D3D3"/>'.format ( nsdecls ( 'w' ) ) )
    shading_elm_9 = parse_xml ( r'<w:shd {} w:fill="D3D3D3"/>'.format ( nsdecls ( 'w' ) ) )
    shading_elm_10 = parse_xml ( r'<w:shd {} w:fill="D3D3D3"/>'.format ( nsdecls ( 'w' ) ) )
    shading_elm_11 = parse_xml ( r'<w:shd {} w:fill="D3D3D3"/>'.format ( nsdecls ( 'w' ) ) )

    tableno.style = 'Table Grid'
    tableno.alignment = WD_TABLE_ALIGNMENT.CENTER
    tableno.cell ( 0 ,0 ).add_paragraph ( 'Thruster: ' + str ( tn + 1 ) ,'texBtt1' )
    tableno.cell ( 0 ,0 ).add_paragraph ( 'Load:' ,'texBtt1' )
    tableno.cell ( 0 ,0 ).add_paragraph ( 'RPM:' ,'texBtt1' )
    tableno.cell ( 0 ,1 ).add_paragraph ( 'Electric motor' ,'texBtt1' )
    tableno.cell ( 0 ,1 ).merge ( tableno.cell ( 0 ,2 ).merge (
        tableno.cell ( 0 ,3 ).merge ( tableno.cell ( 0 ,4 ).merge ( tableno.cell ( 0 ,5 ) ) ) ) )
    tableno.cell ( 0 ,6 ).add_paragraph ( 'Gear Part' ,'texBtt1' )
    tableno.cell ( 0 ,6 ).merge ( tableno.cell ( 0 ,7 ).merge ( tableno.cell ( 0 ,8 ) ) )
    tableno.cell ( 1 ,0 ).add_paragraph ( 'Position of accelerometers' ,'textt1' )
    tableno.rows [ 1 ].cells [ 0 ]._tc.get_or_add_tcPr ().append ( shading_elm_1 )

    tableno.cell ( 1 ,1 ).add_paragraph ( 'Mot Sup' ,'textt1' )
    tableno.cell ( 1 ,1 ).add_paragraph ( 'Proa' ,'textt1' )
    tableno.cell ( 1 ,1 ).add_paragraph ( 'H1' ,'textt1' )
    tableno.rows [ 1 ].cells [ 1 ]._tc.get_or_add_tcPr ().append ( shading_elm_2 )

    tableno.cell ( 1 ,2 ).add_paragraph ( 'Mot Sup' ,'textt1' )
    tableno.cell ( 1 ,2 ).add_paragraph ( 'Barbor' ,'textt1' )
    tableno.cell ( 1 ,2 ).add_paragraph ( 'HH1' ,'textt1' )
    tableno.rows [ 1 ].cells [ 2 ]._tc.get_or_add_tcPr ().append ( shading_elm_3 )

    tableno.cell ( 1 ,3 ).add_paragraph ( 'Mot Sup' ,'textt1' )
    tableno.cell ( 1 ,3 ).add_paragraph ( 'Proa' ,'textt1' )
    tableno.cell ( 1 ,3 ).add_paragraph ( 'H2' ,'textt1' )
    tableno.rows [ 1 ].cells [ 3 ]._tc.get_or_add_tcPr ().append ( shading_elm_4 )

    tableno.cell ( 1 ,4 ).add_paragraph ( 'Mot Sup' ,'textt1' )
    tableno.cell ( 1 ,4 ).add_paragraph ( 'Barbor' ,'textt1' )
    tableno.cell ( 1 ,4 ).add_paragraph ( 'HH2' ,'textt1' )
    tableno.rows [ 1 ].cells [ 4 ]._tc.get_or_add_tcPr ().append ( shading_elm_5 )

    tableno.cell ( 1 ,5 ).add_paragraph ( 'Mot Sup' ,'textt1' )
    tableno.cell ( 1 ,5 ).add_paragraph ( 'Ver' ,'textt1' )
    tableno.cell ( 1 ,5 ).add_paragraph ( 'V2' ,'textt1' )
    tableno.rows [ 1 ].cells [ 5 ]._tc.get_or_add_tcPr ().append ( shading_elm_6 )

    tableno.cell ( 1 ,6 ).add_paragraph ( 'Mot Sup' ,'textt1' )
    tableno.cell ( 1 ,6 ).add_paragraph ( 'Proa' ,'textt1' )
    tableno.cell ( 1 ,6 ).add_paragraph ( 'H3' ,'textt1' )
    tableno.rows [ 1 ].cells [ 6 ]._tc.get_or_add_tcPr ().append ( shading_elm_7 )

    tableno.cell ( 1 ,7 ).add_paragraph ( 'Mot Sup' ,'textt1' )
    tableno.cell ( 1 ,7 ).add_paragraph ( 'Proa' ,'textt1' )
    tableno.cell ( 1 ,7 ).add_paragraph ( 'HH3' ,'textt1' )
    tableno.rows [ 1 ].cells [ 7 ]._tc.get_or_add_tcPr ().append ( shading_elm_8 )

    tableno.cell ( 1 ,8 ).add_paragraph ( 'Mot Sup' ,'textt1' )
    tableno.cell ( 1 ,8 ).add_paragraph ( 'Ver' ,'textt1' )
    tableno.cell ( 1 ,8 ).add_paragraph ( 'V3' ,'textt1' )
    tableno.rows [ 1 ].cells [ 8 ]._tc.get_or_add_tcPr ().append ( shading_elm_9 )

    maxcord = 0
    maxval = 0
    result = [ ]

    querry = "select value from (select distinct ml.value, ml.point, pts.sort from measurements_low as ml RIGHT JOIN points as pts ON ml.id = pts.id and ml.point = pts.point where ml.id = " + str (
        id1 ) + " and ml.raport_number = '" + str (
        rn ) + "' and ml.point IN (select point from points as pts where id = " + str (
        id1 ) + " order by sort)) as al order by al.sort"
    # messagebox.showinfo("Title", str(querry))

    result = q_run ( connD ,querry )
    # messagebox.showinfo("Title", str(result))

    for j in range ( 5 ):
        if result [ j ] [ 0 ] > maxval:
            maxcord = j
            maxval = result [ j ] [ 0 ]

    for j in range ( 5 ):
        result_text = str ( '%.3f'%result [ j ] [ 0 ] )
        #result_text = str ( result [ j ] [ 0 ] )
        p = tableno.cell ( 2 ,1 + j ).add_paragraph ()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run ( result_text.replace ( "." ,"," ) )
        r.font.name = 'Calibri'
        r.font.size = Pt ( 11 )
        if maxcord == j:
            r.bold = True
        if result [ j ] [ 0 ] < 2.3:
            limittxt = 'Z. A'
        elif result [ j ] [ 0 ] > 2.3 and result [ j ] [ 0 ] < 4.5:
            limittxt = 'Z. B'
        elif result [ j ] [ 0 ] > 4.5 and result [ j ] [ 0 ] > 7.1:
            limittxt = 'Z. C'
        else:
            limittxt = 'Z. D'
        p = tableno.cell ( 3 ,1 + j ).add_paragraph ()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run ( limittxt )
        if limittxt == 'Z. A':
            r.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
        if limittxt == 'Z. B':
            r.font.highlight_color = WD_COLOR_INDEX.YELLOW
        if limittxt == 'Z. C' or limittxt == 'Z. D':
            r.font.highlight_color = WD_COLOR_INDEX.RED
        r.font.name = 'Calibri'
        r.font.size = Pt ( 11 )

    maxcord = 0
    maxval = 0
    result = [ ]

    cur = conn.cursor ()
    querry = "select value from (select distinct ml.value, ml.point, pts.sort from measurements_low as ml RIGHT JOIN points as pts ON ml.id = pts.id and ml.point = pts.point where ml.id = " + str (
        id2 ) + " and ml.raport_number = '" + str (
        rn ) + "' and ml.point IN (select point from points as pts where id = " + str (
        id2 ) + " order by sort)) as al order by al.sort"
    cur.execute ( querry )
    result = cur.fetchall ()
    conn.commit ()
    cur.close ()

    for j in range ( 3 ):
        if result [ j ] [ 0 ] > maxval:
            maxcord = j
            maxval = result [ j ] [ 0 ]

    for j in range ( 3 ):

        result_text = str ( '%.3f'%result [ j ] [ 0 ] )
        #result_text = str ( result [ j ] [ 0 ] )
        p = tableno.cell ( 2 ,6 + j ).add_paragraph ()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run ( result_text.replace ( "." ,"," ) )
        r.font.name = 'Calibri'
        r.font.size = Pt ( 11 )
        if maxcord == j:
            r.bold = True
        if result [ j ] [ 0 ] < 4.5:
            limittxt = 'Z. A'
        elif result [ j ] [ 0 ] > 4.5 and result [ j ] [ 0 ] < 11.2:
            limittxt = 'Z. B'
        else:
            limittxt = 'Out of limit'
        p = tableno.cell ( 3 ,6 + j ).add_paragraph ()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run ( limittxt )
        if limittxt == 'Z. A':
            r.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
        if limittxt == 'Z. B':
            r.font.highlight_color = WD_COLOR_INDEX.YELLOW
        if limittxt == 'Z. C':
            r.font.highlight_color = WD_COLOR_INDEX.RED
        r.font.name = 'Calibri'
        r.font.size = Pt ( 11 )

    # tableno.cell(2,6).add_paragraph(str(result[0][0]),'textt1')
    # if result[0][0] < 7:
    # tableno.cell(3,6).add_paragraph('In Limit','textt1')
    # else:
    # tableno.cell(3,6).add_paragraph('Out of limit','textt1')
    # tableno.cell(2,7).add_paragraph(str(result[1][0]),'textt1')
    # if result[1][0] < 7:
    # tableno.cell(3,7).add_paragraph('In Limit','textt1')
    # else:
    # tableno.cell(3,7).add_paragraph('Out of limit','textt1')
    # tableno.cell(2,8).add_paragraph(str(result[2][0]),'textt1')
    # if result[2][0] < 7:
    # tableno.cell(3,8).add_paragraph('In Limit','textt1')
    # else:
    # tableno.cell(3,8).add_paragraph('Out of limit','textt1')

    tableno.cell ( 2 ,0 ).add_paragraph ( 'Overall velocity RMS (mm/s)' ,'textt1' )
    tableno.rows [ 2 ].cells [ 0 ]._tc.get_or_add_tcPr ().append ( shading_elm_10 )

    tableno.cell ( 3 ,0 ).add_paragraph ( 'Vibration Class' ,'textt1' )
    tableno.rows [ 3 ].cells [ 0 ]._tc.get_or_add_tcPr ().append ( shading_elm_11 )


def getname ( host ,username ,password ):
    kport = "5432"
    kdb = "postgres"
    cs = "dbname=%s user=%s password=%s host=%s port=%s"%(kdb ,username ,password ,host ,kport)
    conn = None
    conn = psycopg2.connect ( str ( cs ) )
    cur = conn.cursor ()
    querry = "select full_name from users where login = '" + str ( username ) + "' limit 1;"
    cur.execute ( querry )
    result = cur.fetchall ()
    conn.commit ()
    cur.close ()
    return result


def getini ( host ,username ,password ):
    kport = "5432"
    kdb = "postgres"
    cs = "dbname=%s user=%s password=%s host=%s port=%s"%(kdb ,username ,password ,host ,kport)
    conn = None
    conn = psycopg2.connect ( str ( cs ) )
    cur = conn.cursor ()
    querry = "select ini from users where login = '" + str ( username ) + "' limit 1;"
    cur.execute ( querry )
    result = cur.fetchall ()
    conn.commit ()
    cur.close ()
    return result


def makereport ( username ,password ,host ,rn_ ,id_ ):  # AUTORUN!

    progress = Progressbar ( tk ,orient=HORIZONTAL ,length=200 ,mode='determinate' )
    progress [ 'value' ] = 50
    tk.update_idletasks ()
    progress.pack ()

    h1 = document.add_paragraph ( 'Diagnostic report' + ' ' + rn_ )
    h1.style = document.styles [ 'Head' ]
    h1.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph ()

    l1 = document.add_paragraph ( '1. General information' )
    l1.style = document.styles [ 'listlvl1' ]
    l1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    l1.paragraph_format.left_indent = Inches ( 0.25 )

    l11 = document.add_paragraph ( '1.1.	Diagnostic report' )
    l11.style = document.styles [ 'listlvl2' ]
    l11.alignment = WD_ALIGN_PARAGRAPH.LEFT
    l11.paragraph_format.left_indent = Inches ( 0.25 )

    document.add_paragraph ()

    table = document.add_table ( rows=3 ,cols=2 )
    # table.style = 'Table Grid'

    for cell in table.columns [ 0 ].cells:
        cell.width = Inches ( 1.7 )

    for cell in table.columns [ 1 ].cells:
        cell.width = Inches ( 1.7 )

    table.cell ( 0 ,0 ).add_paragraph ( 'Info Marine Case No.: ' ,'texthead' )
    table.cell ( 0 ,0 ).alignment = WD_ALIGN_PARAGRAPH.LEFT

    table.cell ( 0 ,1 ).add_paragraph ( 'INSERT CASE NUMBER' ,'textt1' )
    table.cell ( 0 ,1 ).alignment = WD_ALIGN_PARAGRAPH.LEFT

    table.cell ( 1 ,0 ).add_paragraph ( 'Report Number: ' ,'texthead' )
    table.cell ( 1 ,0 ).alignment = WD_ALIGN_PARAGRAPH.LEFT

    table.cell ( 1 ,1 ).add_paragraph ( rn_ ,'textt1' )
    table.cell ( 1 ,1 ).alignment = WD_ALIGN_PARAGRAPH.LEFT

    table.cell ( 2 ,0 ).add_paragraph ( 'Date of measurement:' ,'texthead' )
    table.cell ( 2 ,0 ).alignment = WD_ALIGN_PARAGRAPH.LEFT

    table.cell ( 2 ,1 ).add_paragraph ( 'DATE' ,'textt1' )
    table.cell ( 2 ,1 ).alignment = WD_ALIGN_PARAGRAPH.LEFT

    document.add_paragraph ()
    document.add_paragraph ()

    l12 = document.add_paragraph ( '1.2.	Platform Information' )
    l12.style = document.styles [ 'listlvl2' ]
    l12.alignment = WD_ALIGN_PARAGRAPH.LEFT
    l12.paragraph_format.left_indent = Inches ( 0.25 )

    document.add_paragraph ()

    table2 = document.add_table ( rows=8 ,cols=2 )
    # table.style = 'Table Grid'

    for cell in table2.columns [ 0 ].cells:
        cell.width = Inches ( 1.7 )

    for cell in table2.columns [ 1 ].cells:
        cell.width = Inches ( 1.7 )

    table2.cell ( 0 ,0 ).add_paragraph ( 'Name of platform' ,'texthead' )
    table2.cell ( 0 ,0 ).alignment = WD_ALIGN_PARAGRAPH.LEFT

    if id_ == '53':
        platname = 'Atlantis'
        imo_ = '9704128'
    elif id_ == '54':
        platname = 'Iolair'
        imo_ = '7816070'
    elif id_ == '55':
        platname = 'Neptuno'
        imo_ = '9662136'

    table2.cell ( 0 ,1 ).add_paragraph ( str ( platname ) ,'textt1' )
    table2.cell ( 0 ,1 ).alignment = WD_ALIGN_PARAGRAPH.LEFT

    table2.cell ( 1 ,0 ).add_paragraph ( 'The IMO Number:' ,'texthead' )
    table2.cell ( 1 ,0 ).alignment = WD_ALIGN_PARAGRAPH.LEFT

    table2.cell ( 1 ,1 ).add_paragraph ( str ( imo_ ) ,'textt1' )
    table2.cell ( 1 ,1 ).alignment = WD_ALIGN_PARAGRAPH.LEFT

    table2.cell ( 2 ,0 ).add_paragraph ( 'Owner of platform:' ,'texthead' )
    table2.cell ( 2 ,0 ).alignment = WD_ALIGN_PARAGRAPH.LEFT

    table2.cell ( 2 ,1 ).add_paragraph ( 'Cotemar S.A. de C.V.' ,'textt1' )
    table2.cell ( 2 ,1 ).alignment = WD_ALIGN_PARAGRAPH.LEFT

    table2.cell ( 3 ,0 ).add_paragraph ( 'Main dimensions:' ,'textheadU' )
    table2.cell ( 3 ,0 ).alignment = WD_ALIGN_PARAGRAPH.LEFT

    table2.cell ( 4 ,0 ).add_paragraph ( 'Length overall:' ,'texthead' )
    table2.cell ( 4 ,0 ).alignment = WD_ALIGN_PARAGRAPH.LEFT

    table2.cell ( 4 ,1 ).add_paragraph ( '95m' ,'textt1' )
    table2.cell ( 4 ,1 ).alignment = WD_ALIGN_PARAGRAPH.LEFT

    table2.cell ( 5 ,0 ).add_paragraph ( 'Breadth Extreme:' ,'texthead' )
    table2.cell ( 5 ,0 ).alignment = WD_ALIGN_PARAGRAPH.LEFT

    table2.cell ( 5 ,1 ).add_paragraph ( '67m' ,'textt1' )
    table2.cell ( 5 ,1 ).alignment = WD_ALIGN_PARAGRAPH.LEFT

    table2.cell ( 6 ,0 ).add_paragraph ( 'Deadweight:' ,'texthead' )
    table2.cell ( 7 ,0 ).alignment = WD_ALIGN_PARAGRAPH.LEFT

    table2.cell ( 6 ,1 ).add_paragraph ( '13280t' ,'textt1' )
    table2.cell ( 6 ,1 ).alignment = WD_ALIGN_PARAGRAPH.LEFT

    table2.cell ( 7 ,0 ).add_paragraph ( 'Propulsion: ' ,'texthead' )
    table2.cell ( 7 ,0 ).alignment = WD_ALIGN_PARAGRAPH.LEFT

    table2.cell ( 7 ,1 ).add_paragraph ( '6 Thrusters' ,'textt1' )
    table2.cell ( 7 ,1 ).alignment = WD_ALIGN_PARAGRAPH.LEFT

    document.add_page_break ()

    document.add_paragraph ()

    l2 = document.add_paragraph ( '2. Measurement' )
    l2.style = document.styles [ 'listlvl1' ]
    l2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    l2.paragraph_format.left_indent = Inches ( 0.25 )

    l21 = document.add_paragraph ( '2.1.	Specification of measuring position' )
    l21.style = document.styles [ 'listlvl2' ]
    l21.alignment = WD_ALIGN_PARAGRAPH.LEFT
    l21.paragraph_format.left_indent = Inches ( 0.25 )

    document.add_paragraph ()

    t1 = document.add_paragraph ( 'Vibration measurement were performed in three axes.' )
    t1.style = document.styles [ 'textt1' ]
    t1.alignment = WD_ALIGN_PARAGRAPH.LEFT

    document.add_paragraph ()

    t2 = document.add_paragraph ( 'V – Vertical axis' )
    t2.style = document.styles [ 'textt1' ]
    t2.alignment = WD_ALIGN_PARAGRAPH.LEFT

    document.add_paragraph ()

    t3 = document.add_paragraph ( 'H – Transverse axis' )
    t3.style = document.styles [ 'textt1' ]
    t3.alignment = WD_ALIGN_PARAGRAPH.LEFT

    document.add_paragraph ()

    t4 = document.add_paragraph ( 'HH – Longitudinal axis' )
    t4.style = document.styles [ 'textt1' ]
    t4.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # OBRAZEK KTORY WYKRZACZA WORDA
    p = document.add_paragraph ()
    r = p.add_run ()
    r.add_picture ( 'C:\\overmind\\Data\\Atlantis_thruster.jpg' )
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    des1 = document.add_paragraph ( 'Fig.1 Localization of measuring points' )
    des1.style = document.styles [ 'description' ]
    des1.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_page_break ()
    l22 = document.add_paragraph ( '2.2.	Measurement standard' )
    l22.style = document.styles [ 'listlvl2' ]
    l22.alignment = WD_ALIGN_PARAGRAPH.LEFT
    l22.paragraph_format.left_indent = Inches ( 0.25 )

    t5 = document.add_paragraph ( 'Following standards are applied for assessment:' )
    t5.style = document.styles [ 'textt1' ]
    t5.alignment = WD_ALIGN_PARAGRAPH.LEFT

    document.add_paragraph ()

    t6 = document.add_paragraph ( 'ISO 10816 – 1	General guidelines' )
    t6.style = document.styles [ 'textt1' ]
    t6.alignment = WD_ALIGN_PARAGRAPH.LEFT

    document.add_paragraph ()

    t7 = document.add_paragraph ( 'DNV Rules of classification of ships. Part 6 Chapter 15' )
    t7.style = document.styles [ 'textt1' ]
    t7.alignment = WD_ALIGN_PARAGRAPH.LEFT

    document.add_paragraph ()

    table3 = document.add_table ( rows=4 ,cols=2 )
    # table.style = 'Table Grid'

    for cell in table3.columns [ 0 ].cells:
        cell.width = Inches ( 1.7 )

    for cell in table3.columns [ 1 ].cells:
        cell.width = Inches ( 1.7 )

    table3.cell ( 0 ,0 ).add_paragraph ( 'The criteria(Motor):' ,'textheadU' )
    table3.cell ( 0 ,0 ).alignment = WD_ALIGN_PARAGRAPH.LEFT

    table3.cell ( 0 ,0 ).merge ( table3.cell ( 0 ,1 ) )

    table3.cell ( 1 ,0 ).add_paragraph ( 'Value:' ,'texthead' )
    table3.cell ( 1 ,0 ).alignment = WD_ALIGN_PARAGRAPH.LEFT
    table3.cell(2,0).add_paragraph('Range: ','texthead')
    table3.cell(2,1).add_paragraph('2-1000Hz','textt1')
    vel=table3.cell ( 1 ,1 ).add_paragraph ( 'Velocity' ,'textt1' )
    table3.cell ( 1 ,1 ).alignment = WD_ALIGN_PARAGRAPH.LEFT
    vel.add_run().add_break()
    vel.add_run('Zone A:        0-2,3 mm/s')
    vel.add_run ().add_break ()
    vel.add_run('Zone B:        2,3-4,5 mm/s')
    vel.add_run ().add_break ()
    vel.add_run('Zone C:        4,5-7,1 mm/s')
    vel.add_run ().add_break ()
    vel.add_run('Zone D:        above 7,1 mm/s')

    table3.cell ( 1 ,1 ).alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_col_widths ( table3 )
    document.add_paragraph ()

    table4 = document.add_table ( rows=4 ,cols=2 )
    # table.style = 'Table Grid'

    for cell in table4.columns [ 0 ].cells:
        cell.width = Inches ( 1.7 )

    for cell in table4.columns [ 1 ].cells:
        cell.width = Inches ( 1.7 )

    table4.cell ( 0 ,0 ).add_paragraph ( 'The criteria(Gear):' ,'textheadU' )
    table4.cell ( 0 ,0 ).alignment = WD_ALIGN_PARAGRAPH.LEFT

    table4.cell ( 0 ,0 ).merge ( table4.cell ( 0 ,1 ) )

    table4.cell ( 1 ,0 ).add_paragraph ( 'Value:' ,'texthead' )
    table4.cell ( 1 ,0 ).alignment = WD_ALIGN_PARAGRAPH.LEFT
    velgear=table4.cell(1,1).add_paragraph('Velocity: ','textt1')
    table4.cell(2,0).add_paragraph('Range: ','texthead')
    table4.cell(2,1).add_paragraph('5-1000Hz','textt1')
    velgear.add_run ().add_break ()
    velgear.add_run ( 'Zone A:        0-4,5 mm/s' )
    velgear.add_run ().add_break ()
    velgear.add_run ( 'Zone B:        4,5-11,2 mm/s' )
    velgear.add_run ().add_break ()
    velgear.add_run ( 'Zone C:        above 11,2 mm/s' )

    table4.cell ( 1 ,1 ).alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_col_widths(table4)
    document.add_paragraph('Legend according to vibration class','listlvl1')
    limittablelegend=document.add_table(rows=5,cols=2)
    limittablelegend.cell(0,0).paragraphs[0].add_run('Z. A').font.highlight_color=WD_COLOR_INDEX.BRIGHT_GREEN
    limittablelegend.cell(0,1).paragraphs[0].add_run('Newly commissioned')

    limittablelegend.cell ( 1 ,0 ).paragraphs [ 0 ].add_run ( 'Z. B' ).font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
    limittablelegend.cell(1,1).paragraphs[0].add_run('Unrestricted')

    limittablelegend.cell(2,0).paragraphs[0].add_run('Z. C').font.highlight_color=WD_COLOR_INDEX.YELLOW
    limittablelegend.cell(2,1).paragraphs[0].add_run('Restricted long-term operation')

    limittablelegend.cell(3,0).paragraphs[0].add_run('Z. D').font.highlight_color=WD_COLOR_INDEX.RED
    limittablelegend.cell(3,1).paragraphs[0].add_run('High probability of damage, action required')

    limittablelegend.cell(4,0).paragraphs[0].add_run('Z. D').font.bold=True
    limittablelegend.cell(4,1).paragraphs[0].add_run('Vibrations over the limits but actions are not required.')

    set_col_widths ( limittablelegend )

    document.add_page_break ()
    l23 = document.add_paragraph ( '2.3.	Measurement equipment' )
    l23.style = document.styles [ 'listlvl2' ]
    l23.alignment = WD_ALIGN_PARAGRAPH.LEFT
    l23.paragraph_format.left_indent = Inches ( 0.25 )

    document.add_paragraph ()

    t8 = document.add_paragraph ( 'WBK18  8-channel dynamic signal condition modules' )
    t8.style = document.styles [ 'textt1' ]
    t8.alignment = WD_ALIGN_PARAGRAPH.LEFT

    document.add_paragraph ()

    t10 = document.add_paragraph ( 'eZ – TOMAS (Total Online Monitoring and Analysis Software)' )
    t10.style = document.styles [ 'textt1' ]
    t10.alignment = WD_ALIGN_PARAGRAPH.LEFT

    document.add_paragraph ()

    table5 = document.add_table ( rows=3 ,cols=2 )
    # table.style = 'Table Grid'

    for cell in table5.columns [ 0 ].cells:
        cell.width = Inches ( 1.7 )

    for cell in table5.columns [ 1 ].cells:
        cell.width = Inches ( 1.7 )

    table5.cell ( 0 ,0 ).add_paragraph ( 'Window:' ,'textt1' )
    table5.cell ( 0 ,0 ).alignment = WD_ALIGN_PARAGRAPH.LEFT

    table5.cell ( 0 ,1 ).add_paragraph ( 'Hanning' ,'textt1' )
    table5.cell ( 0 ,1 ).alignment = WD_ALIGN_PARAGRAPH.LEFT

    table5.cell ( 1 ,0 ).add_paragraph ( 'Range:' ,'textt1' )
    table5.cell ( 1 ,0 ).alignment = WD_ALIGN_PARAGRAPH.LEFT

    table5.cell ( 1 ,1 ).add_paragraph ( '2Hz-1kHz' ,'textt1' )
    table5.cell ( 1 ,1 ).alignment = WD_ALIGN_PARAGRAPH.LEFT

    table5.cell ( 2 ,0 ).add_paragraph ( 'F resolution:' ,'textt1' )
    table5.cell ( 2 ,0 ).alignment = WD_ALIGN_PARAGRAPH.LEFT

    table5.cell ( 2 ,1 ).add_paragraph ( '0,16 Hz' ,'textt1' )
    table5.cell ( 2 ,1 ).alignment = WD_ALIGN_PARAGRAPH.LEFT

    document.add_paragraph ()

    l24 = document.add_paragraph ( '2.4.	Measurement condition' )
    l24.style = document.styles [ 'listlvl2' ]
    l24.alignment = WD_ALIGN_PARAGRAPH.LEFT
    l24.paragraph_format.left_indent = Inches ( 0.25 )

    document.add_paragraph ()

    t11 = document.add_paragraph ( 'Measurement were performed during normal operation of platform.' )
    t11.style = document.styles [ 'textt1' ]
    t11.alignment = WD_ALIGN_PARAGRAPH.LEFT

    document.add_paragraph ()

    l25 = document.add_paragraph ( '2.5.	Measurement results' )
    l25.style = document.styles [ 'listlvl2' ]
    l25.alignment = WD_ALIGN_PARAGRAPH.LEFT
    l25.paragraph_format.left_indent = Inches ( 0.25 )

    document.add_paragraph ()

    t12 = document.add_paragraph (
        'In below tables are presented RMS velocity values for all measured points of electric motor and gear part. Additionally, FFT graphs for highest values on electric motor and gear part are included.' )
    t12.style = document.styles [ 'textt1' ]
    t12.alignment = WD_ALIGN_PARAGRAPH.LEFT

    if id_ == '53':
        motors = [ 17159 ,17161 ,17162 ,17163 ,17164 ,17165 ]
        thrusters = [ 17160 ,17166 ,17167 ,17168 ,17169 ,17170 ]
    elif id_ == '54':
        platname = 'Iolar'
    elif id_ == '55':
        motors = [ 18137 ,18127 ,18129 ,18131 ,18133 ,18135 ]
        thrusters = [ 18138 ,18128 ,18130 ,18132 ,18134 ,18136 ]

    for j in range ( 6 ):
        maxP = 'MAXPOINT'
        progress [ 'value' ] = 50 + j*10
        tk.update_idletasks ()

        document.add_page_break ()
        # document.add_paragraph()
        table6 = document.add_table ( rows=4 ,cols=9 )
        resulttable ( j ,host ,username ,password ,table6 ,motors [ j ] ,thrusters [ j ] ,rn_ )
        chart = document.add_paragraph ()
        r = chart.add_run ()
        maxPoint = createchart ( host ,username ,password ,motors [ j ] ,rn_ )
        r.add_picture ( 'C:\\overmind\\temp\\tempchart.png' ,width=Inches ( 7.0 ) )
        os.remove ( 'C:\\overmind\\temp\\tempchart.png' )
        r = chart.add_run (
            ' \n Fig.' + str ( (j*2) + 2 ) + ' FFT velocity graph point ' + str ( maxPoint ) + ' on Thruster ' + str (
                j + 1 ) + ' el. Motor' )
        r.font.name = 'Calibri'
        r.font.size = Pt ( 9 )
        r.font.italic = True
        chart.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        chart = document.add_paragraph ()
        r = chart.add_run ()
        maxPoint = createchart ( host ,username ,password ,thrusters [ j ] ,rn_ )
        r.add_picture ( 'C:\\overmind\\temp\\tempchart.png' ,width=Inches ( 7.0 ) )
        os.remove ( 'C:\\overmind\\temp\\tempchart.png' )
        r = chart.add_run (
            ' \n Fig.' + str ( (j*2) + 3 ) + ' FFT velocity graph point ' + str ( maxPoint ) + ' on Thruster ' + str (
                j + 1 ) + ' gear' )
        r.font.name = 'Calibri'
        r.font.size = Pt ( 9 )
        r.font.italic = True
        chart.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_page_break ()
    l3 = document.add_paragraph ( '3. Summary' )
    l3.style = document.styles [ 'listlvl1' ]
    l3.alignment = WD_ALIGN_PARAGRAPH.LEFT
    l3.paragraph_format.left_indent = Inches ( 0.25 )
    document.add_paragraph ()
    l31 = document.add_paragraph ( '3.1.	Conclusion // SPRAWDZIC / UZUPEŁNIC PRZED WYSŁANIEM // ' )
    l31.style = document.styles [ 'listlvl2' ]
    l31.alignment = WD_ALIGN_PARAGRAPH.LEFT
    l31.paragraph_format.left_indent = Inches ( 0.25 )
    document.add_paragraph ()
    t13 = document.add_paragraph (
        'The vibration of 6 thruster (electric motors and gear parts) are on very low level. Analysis was performed only on vibration values. Oil analysis is not available.' )
    t13.style = document.styles [ 'textt1' ]
    t13.alignment = WD_ALIGN_PARAGRAPH.LEFT
    document.add_paragraph ()
    l32 = document.add_paragraph ( '3.2.	Recommendation // SPRAWDZIC / UZUPEŁNIC PRZED WYSŁANIEM // ' )
    l32.style = document.styles [ 'listlvl2' ]
    l32.alignment = WD_ALIGN_PARAGRAPH.LEFT
    l32.paragraph_format.left_indent = Inches ( 0.25 )
    document.add_paragraph ()
    t14 = document.add_paragraph (
        'Next measurements should be done within one month period or earlier if necessary.' )
    t14.style = document.styles [ 'textt1' ]
    t14.alignment = WD_ALIGN_PARAGRAPH.LEFT
    document.add_paragraph ()
    t15 = document.add_paragraph (
        'This report is prepared in good faith based on measurement diagnostic done on available object and documentation submitted.' )
    t15.style = document.styles [ 'textt1' ]
    t15.alignment = WD_ALIGN_PARAGRAPH.LEFT

    document.add_paragraph ()
    t15 = document.add_paragraph (
        'Report prepared by:                                                                                                                           Approved by:' )
    t15.style = document.styles [ 'texthead' ]
    t15.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fullname = getname ( host ,username ,password )
    t16 = document.add_paragraph ( str ( fullname [ 0 ] [ 0 ] ) )
    t16.style = document.styles [ 'textt1' ]
    t16.alignment = WD_ALIGN_PARAGRAPH.LEFT

    section = document.sections [ 0 ]
    footer = section.footer.paragraphs [ 0 ]
    footer.add_run ( 'Diagnostic report ' + str ( platname ) + ' Platform ' + str ( rn_ ) )
    footer.runs [ 0 ].font.name = 'Times New Roman'
    footer.runs [ 0 ].font.size = Pt ( 12 )
    footer.alignment = WD_ALIGN_PARAGRAPH.LEFT
    footer.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    footer_ = section.footer
    footer_.text = 'diagnostrep'
    mydate = datetime.datetime.now ()
    repdate = str ( mydate.year ) + '_' + str ( mydate.month ) + '_' + str ( mydate.day )
    ini = getini ( host ,username ,password )

    if id_ == '53':
        document.save ( 'C:\overmind\Reports\Diagnostic report Atlantis Platform ' + str ( rn_ ) + '_' + str (
            ini [ 0 ] [ 0 ] ) + '_' + str ( repdate ) + '.docx' )
    elif id_ == '54':
        document.save ( 'C:\overmind\Reports\Diagnostic report Iolar Platform ' + str ( rn_ ) + '_' + str (
            ini [ 0 ] [ 0 ] ) + '_' + str ( repdate ) + '.docx' )
    elif id_ == '55':
        document.save ( 'C:\overmind\Reports\Diagnostic report Neptuno Platform ' + str ( rn_ ) + '_' + str (
            ini [ 0 ] [ 0 ] ) + '_' + str ( repdate ) + '.docx' )

# makereport ( 'testuser' ,'info' ,'192.168.10.243' ,'2065-2019' ,'55' )
