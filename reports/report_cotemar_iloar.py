import sys
import datetime
import os
import os.path
import time

from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.shared import RGBColor

from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_BREAK
from docx.enum.text import WD_COLOR_INDEX
from docx.enum.text import WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT , WD_ALIGN_VERTICAL
from docx.enum.text import WD_TAB_ALIGNMENT , WD_TAB_LEADER

from tkinter import messagebox
from tkinter import *
from tkinter.ttk import *

from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

import numpy as np

import matplotlib.pyplot as plt

import psycopg2

tk = Tk ()
document = Document ( 'C:\\overmind\\Data\\base.docx' )

def set_col_widths(table):
    widths = (Inches(0.6), Inches(5))
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width
def q_run ( connD , querry ):
	username = connD [ 0 ]
	password = connD [ 1 ]
	host = connD [ 2 ]
	kport = "5432"
	kdb = "postgres"
	# cs = ' host="localhost",database="postgres", user= "postgres" , password="info" '
	cs = "dbname=%s user=%s password=%s host=%s port=%s"%(kdb , username , password , host , kport)
	conn = None
	conn = psycopg2.connect ( str ( cs ) )
	cur = conn.cursor ()
	cur.execute ( querry )
	try:
		result = cur.fetchall ()
		return result
	except:
		pass
	conn.commit ()
	cur.close ()


######################################
#################STYLE################
######################################

obj_styles = document.styles
obj_charstyle = obj_styles.add_style ( 'Head' , WD_STYLE_TYPE.PARAGRAPH )
style = document.styles [ 'Head' ]
font = style.font
font.name = 'Calibri'
font.size = Pt ( 16 )
font.bold = True
font.color.rgb = RGBColor ( 68 , 114 , 196 )

obj_charstyle = obj_styles.add_style ( 'listlvl1' , WD_STYLE_TYPE.PARAGRAPH )
style = document.styles [ 'listlvl1' ]
font = style.font
font.name = 'Calibri Light'
font.size = Pt ( 16 )
font.bold = True
font.color.rgb = RGBColor ( 47 , 84 , 150 )

obj_charstyle = obj_styles.add_style ( 'listlvl2' , WD_STYLE_TYPE.PARAGRAPH )
style = document.styles [ 'listlvl2' ]
font = style.font
font.name = 'Calibri Light'
font.size = Pt ( 13 )
font.bold = True
font.color.rgb = RGBColor ( 47 , 84 , 150 )

obj_charstyle = obj_styles.add_style ( 'texthead' , WD_STYLE_TYPE.PARAGRAPH )
style = document.styles [ 'texthead' ]
font = style.font
font.name = 'Calibri Light'
font.size = Pt ( 11 )
font.bold = True

obj_charstyle = obj_styles.add_style ( 'textheadU' , WD_STYLE_TYPE.PARAGRAPH )
style = document.styles [ 'textheadU' ]
font2 = style.font
font2.name = 'Calibri Light'
font2.size = Pt ( 11 )
font2.bold = True
font2.underline = True

obj_charstyle = obj_styles.add_style ( 'text' , WD_STYLE_TYPE.PARAGRAPH )
style = document.styles [ 'text' ]
font = style.font
font.name = 'Calibri'
font.size = Pt ( 11 )

obj_charstyle = obj_styles.add_style ( 'textB' , WD_STYLE_TYPE.PARAGRAPH )
style = document.styles [ 'textB' ]
font = style.font
font.name = 'Calibri'
font.size = Pt ( 11 )
font.bold = True

obj_charstyle = obj_styles.add_style ( 'textT' , WD_STYLE_TYPE.PARAGRAPH )
style = document.styles [ 'textT' ]
font = style.font
font.name = 'Calibri'
font.size = Pt ( 9 )

obj_charstyle = obj_styles.add_style ( 'textBT' , WD_STYLE_TYPE.PARAGRAPH )
style = document.styles [ 'textBT' ]
font = style.font
font.name = 'Calibri'
font.size = Pt ( 9 )
font.bold = True

obj_charstyle = obj_styles.add_style ( 'warning' , WD_STYLE_TYPE.PARAGRAPH )
style = document.styles [ 'warning' ]
font = style.font
font.name = 'Calibri Light'
font.size = Pt ( 13 )

obj_charstyle = obj_styles.add_style ( 'description' , WD_STYLE_TYPE.PARAGRAPH )
style = document.styles [ 'description' ]
font = style.font
font.name = 'Calibri'
font.size = Pt ( 9 )
font.italic = True


######################################


######################################
#######DEFINICJE FUNKCJI#############
######################################
def createchart ( host , username , password , id , rn , type ):
	connD = [ username , password , host ]
	querry = "select point from measurements_low where id = " + str ( id ) + " and raport_number = '" + str (
		rn ) + "' and type = '" + str ( type ) + "' order by value desc limit 1;"
	maxPoint = q_run ( connD , querry ) [ 0 ] [ 0 ]

	X = [ ]
	Y = [ ]
	x = [ ]

	if type == 'RMS':
		type = 'Vel'
	elif type == 'envelope P-K':
		type = 'Env'

	try:
		querry = "select chart from meascharts where id = '" + str ( id ) + "' and report_number = '" + str (
			rn ) + "' and point = '" + str ( maxPoint ) + "' and domain = 'FFT' and type = '" + str ( type ) + "' ;"
		chartstr = q_run ( connD , querry ) [ 0 ] [ 0 ]
		chartstrlist = chartstr.split ( ";" )
		for line in chartstrlist:
			x.append ( float ( line ) )
	except:
		if host == 'localhost':
			querry = "select lo_export(measurements.chart, 'C:\\Overmind\\temp\\tempchart.csv') from measurements where id = " + str (
				id ) + " and report_number = '" + str ( rn ) + "' and point = '" + str (
				maxPoint ) + "' and type = '" + str ( type ) + "' ;"
		if host == '192.168.8.125':
			servpath = r'/home/filip/Public/tempchart.csv'
			querry = "select lo_export(measurements.chart,'" + servpath + "') from measurements where id = '" + str (
				id ) + "' and report_number = '" + str ( rn ) + "' and point = '" + str (
				maxPoint ) + "' and type = '" + str ( type ) + "' ;"

		result = q_run ( connD , querry )
		# os.remove(spath_)
		if host == 'localhost':
			while not os.path.exists ( r'C:\\Overmind\\temp\\tempchart.csv' ):
				time.sleep ( 1 )
			try:
				x = np.loadtxt ( r'C:\\Overmind\\temp\\tempchart.csv' )
			except:
				messagebox.showinfo ( "Title" , str ( querry ) )
		if host == '192.168.8.125':
			servpath = r'\\192.168.8.125\Public\tempchart.csv'
			while not os.path.exists ( servpath ):
				time.sleep ( 1 )
			try:
				x = np.loadtxt ( servpath )
			except:
				messagebox.showinfo ( "Title" , str ( querry ) )

	xlen = int ( len ( x ) - 2 )
	tp = x [ 0 ]
	dt = (x [ 1 ]/xlen)
	x = x [ 2: ]
	it = 0
	t = tp
	maxf = 0

	if type == 'Env':
		maxf = 800
		minf = 2

	if type == 'Vel':
		if str ( id ) == '18100' or str ( id ) == '18102' or str ( id ) == '18104' or str ( id ) == '18106' or str (
				id ) == '18108' or str ( id ) == '18110' or str ( id ) == '18112' or str ( id ) == '18114' or str (
				id ) == '18116' or str ( id ) == '18118' or str ( id ) == '18120' or str ( id ) == '18122':
			maxf = 1000  # motor
			minf = 10
		if str ( id ) == '18101' or str ( id ) == '18103' or str ( id ) == '18105' or str ( id ) == '18107' or str (
				id ) == '18109' or str ( id ) == '18111' or str ( id ) == '18113' or str ( id ) == '18115' or str (
				id ) == '18117' or str ( id ) == '18119' or str ( id ) == '18121' or str ( id ) == '18123':
			maxf = 1000  # gear
			minf = 5

	for i in x:
		if t < minf:
			x [ it ] = 0
			itmin = it
		if t > (maxf):
			x [ it ] = 0
		it += 1
		t += dt
		Y.append ( t )

	fig = plt.figure ( figsize=(15 , 6) , dpi=80 )
	plt.cla ()
	plt.plot ( Y , x , linewidth=0.3 )
	plt.xlabel ( 'Frequency [Hz]' )
	if type=='Vel':
		plt.ylabel('Velocity [mm/s]')
	elif type=='Env':
		#plt.ylabel('Envelope')
		plt.ylabel('Acceleration [m/s2]')
	plt.xlabel ( 'Frequency [Hz]' )
	#plt.ylabel ( 'Velocity[mm/s]' )
	axes = plt.gca ()

	maxval = np.amax ( x )
	maxcord = (np.argmax ( x ) + 2)*dt
	plt.plot ( [ 0 , 800 ] , [ maxval , maxval ] , lw=0.3 , color='red' )
	plt.plot ( [ maxcord , maxcord ] , [ maxval , maxval*1.1 ] , lw=0.3 )
	if maxf == 200:
		maxf = 205
	axes.set_xlim ( [ 0 , maxf ] )
	axes.set_ylim ( [ 0 , maxval*1.1 ] )
	plt.savefig ( 'C:\\overmind\\temp\\tempchart.png' , dpi=300 , width=10 , color='black' )

	try:
		if host == 'localhost':
			os.remove ( 'C:\\Overmind\\temp\\tempchart.csv' )
		if host == '192.168.8.125':
			servpath = r'\\192.168.8.125\Public\tempchart.csv'
			os.remove ( servpath )
	except:
		pass
	return maxPoint
	plt.close ()


def resulttable_iloar ( tn , host , username , password , tableno , id1 , id2 , id3 , id4 , id5 , id6 , rn ):
	kport = "5432"
	kdb = "postgres"

	# cs = ' host="localhost",database="postgres", user= "postgres" , password="info" '
	cs = "dbname=%s user=%s password=%s host=%s port=%s"%(kdb , username , password , host , kport)
	conn = None
	conn = psycopg2.connect ( str ( cs ) )

	shading_elm_1 = parse_xml ( r'<w:shd {} w:fill="D3D3D3"/>'.format ( nsdecls ( 'w' ) ) )
	shading_elm_2 = parse_xml ( r'<w:shd {} w:fill="D3D3D3"/>'.format ( nsdecls ( 'w' ) ) )
	shading_elm_3 = parse_xml ( r'<w:shd {} w:fill="D3D3D3"/>'.format ( nsdecls ( 'w' ) ) )
	shading_elm_4 = parse_xml ( r'<w:shd {} w:fill="D3D3D3"/>'.format ( nsdecls ( 'w' ) ) )
	shading_elm_5 = parse_xml ( r'<w:shd {} w:fill="D3D3D3"/>'.format ( nsdecls ( 'w' ) ) )
	shading_elm_6 = parse_xml ( r'<w:shd {} w:fill="D3D3D3"/>'.format ( nsdecls ( 'w' ) ) )
	shading_elm_7 = parse_xml ( r'<w:shd {} w:fill="D3D3D3"/>'.format ( nsdecls ( 'w' ) ) )
	shading_elm_8 = parse_xml ( r'<w:shd {} w:fill="D3D3D3"/>'.format ( nsdecls ( 'w' ) ) )
	shading_elm_9 = parse_xml ( r'<w:shd {} w:fill="D3D3D3"/>'.format ( nsdecls ( 'w' ) ) )
	shading_elm_10 = parse_xml ( r'<w:shd {} w:fill="D3D3D3"/>'.format ( nsdecls ( 'w' ) ) )
	shading_elm_11 = parse_xml ( r'<w:shd {} w:fill="D3D3D3"/>'.format ( nsdecls ( 'w' ) ) )
	shading_elm_12 = parse_xml ( r'<w:shd {} w:fill="D3D3D3"/>'.format ( nsdecls ( 'w' ) ) )
	shading_elm_13 = parse_xml ( r'<w:shd {} w:fill="D3D3D3"/>'.format ( nsdecls ( 'w' ) ) )
	shading_elm_14 = parse_xml ( r'<w:shd {} w:fill="D3D3D3"/>'.format ( nsdecls ( 'w' ) ) )
	shading_elm_15 = parse_xml ( r'<w:shd {} w:fill="D3D3D3"/>'.format ( nsdecls ( 'w' ) ) )
	shading_elm_16 = parse_xml ( r'<w:shd {} w:fill="D3D3D3"/>'.format ( nsdecls ( 'w' ) ) )
	shading_elm_17 = parse_xml ( r'<w:shd {} w:fill="D3D3D3"/>'.format ( nsdecls ( 'w' ) ) )
	shading_elm_18 = parse_xml ( r'<w:shd {} w:fill="D3D3D3"/>'.format ( nsdecls ( 'w' ) ) )
	shading_elm_19 = parse_xml ( r'<w:shd {} w:fill="D3D3D3"/>'.format ( nsdecls ( 'w' ) ) )
	shading_elm_20 = parse_xml ( r'<w:shd {} w:fill="D3D3D3"/>'.format ( nsdecls ( 'w' ) ) )
	shading_elm_21 = parse_xml ( r'<w:shd {} w:fill="D3D3D3"/>'.format ( nsdecls ( 'w' ) ) )
	shading_elm_22 = parse_xml ( r'<w:shd {} w:fill="D3D3D3"/>'.format ( nsdecls ( 'w' ) ) )

	tableno.style = 'Table Grid'
	tableno.alignment = WD_TABLE_ALIGNMENT.CENTER
	if tn == 0:
		tableno.cell ( 0 , 0 ).paragraphs [ 0 ].add_run ( 'Thruster: SB AFT' )
		tableno.cell ( 0 , 0 ).paragraphs [ 0 ].style = document.styles [ 'textBT' ]
		tableno.cell ( 0 , 0 ).paragraphs [ 0 ].alignment = WD_ALIGN_PARAGRAPH.LEFT
		tableno.cell ( 0 , 0 ).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
	if tn == 1:
		# tableno.cell(0,0).add_paragraph('Thruster: SB FWD','textBT')
		tableno.cell ( 0 , 0 ).paragraphs [ 0 ].add_run ( 'Thruster: SB FWD' )
		tableno.cell ( 0 , 0 ).paragraphs [ 0 ].style = document.styles [ 'textBT' ]
		tableno.cell ( 0 , 0 ).paragraphs [ 0 ].alignment = WD_ALIGN_PARAGRAPH.LEFT
		tableno.cell ( 0 , 0 ).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
	if tn == 2:
		# tableno.cell(0,0).add_paragraph('Thruster: PS AFT','textBT')
		tableno.cell ( 0 , 0 ).paragraphs [ 0 ].add_run ( 'Thruster: PS AFT' )
		tableno.cell ( 0 , 0 ).paragraphs [ 0 ].style = document.styles [ 'textBT' ]
		tableno.cell ( 0 , 0 ).paragraphs [ 0 ].alignment = WD_ALIGN_PARAGRAPH.LEFT
		tableno.cell ( 0 , 0 ).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
	if tn == 3:
		tableno.cell ( 0 , 0 ).paragraphs [ 0 ].add_run ( 'Thruster: PS FWD' )
		tableno.cell ( 0 , 0 ).paragraphs [ 0 ].style = document.styles [ 'textBT' ]
		tableno.cell ( 0 , 0 ).paragraphs [ 0 ].alignment = WD_ALIGN_PARAGRAPH.LEFT
		tableno.cell ( 0 , 0 ).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

	# tableno.cell(0,0).paragraphs[0].add_run('Load:')
	# tableno.cell(0,0).paragraphs[0].add_run('RPM:')

	tableno.cell ( 0 , 1 ).paragraphs [ 0 ].style = document.styles [ 'textBT' ]
	tableno.cell ( 0 , 1 ).paragraphs [ 0 ].add_run ( 'Electric motor' )
	tableno.cell ( 0 , 1 ).paragraphs [ 0 ].runs [ 0 ].add_break ()
	tableno.cell ( 0 , 1 ).paragraphs [ 0 ].add_run ( 'RPM: [UZUPEŁNIĆ]' )
	tableno.cell ( 0 , 1 ).paragraphs [ 0 ].runs [ 1 ].add_break ()
	tableno.cell ( 0 , 1 ).paragraphs [ 0 ].add_run ( 'Running hours: [UZUPEŁNIĆ]' )
	tableno.cell ( 0 , 1 ).paragraphs [ 0 ].runs [ 2 ].add_break ()
	tableno.cell ( 0 , 1 ).paragraphs [ 0 ].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
	tableno.cell ( 0 , 1 ).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

	# tableno.cell(0,1).add_paragraph('RPM: [UZUPEŁNIĆ]','textBT').paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
	# tableno.cell(0,1).add_paragraph('Running hours: [UZUPEŁNIĆ]','textBT').paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
	tableno.cell ( 0 , 1 ).merge ( tableno.cell ( 0 , 2 ).merge (
		tableno.cell ( 0 , 3 ).merge ( tableno.cell ( 0 , 4 ).merge ( tableno.cell ( 0 , 5 ) ) ) ) )

	# tableno.cell(0,6).add_paragraph('Gear Part')
	tableno.cell ( 0 , 6 ).paragraphs [ 0 ].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
	tableno.cell ( 0 , 6 ).paragraphs [ 0 ].add_run ( 'Gear Part' )
	tableno.cell ( 0 , 6 ).paragraphs [ 0 ].runs [ 0 ].add_break ()
	tableno.cell ( 0 , 6 ).paragraphs [ 0 ].style = document.styles [ 'textBT' ]
	tableno.cell ( 0 , 6 ).paragraphs [ 0 ].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

	tableno.cell ( 0 , 6 ).paragraphs [ 0 ].add_run ( 'RPM: [UZUPEŁNIĆ]' )
	tableno.cell ( 0 , 6 ).paragraphs [ 0 ].runs [ 1 ].add_break ()
	tableno.cell ( 0 , 6 ).paragraphs [ 0 ].add_run ( 'Running hours: [UZUPEŁNIĆ]' )
	tableno.cell ( 0 , 6 ).paragraphs [ 0 ].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
	tableno.cell ( 0 , 6 ).merge ( tableno.cell ( 0 , 7 ).merge ( tableno.cell ( 0 , 8 ) ) )

	tableno.rows [ 0 ].cells [ 9 ]._tc.get_or_add_tcPr ().append ( shading_elm_19 )
	# tableno.rows[1].cells[9]._tc.get_or_add_tcPr().append(shading_elm_20)
	tableno.cell ( 0 , 9 ).merge ( tableno.cell ( 1 , 9 ) )
	# tableno.cell(0,9).add_paragraph('Direction','textBT')
	tableno.cell ( 0 , 9 ).paragraphs [ 0 ].style = document.styles [ 'textBT' ]
	tableno.cell ( 0 , 9 ).paragraphs [ 0 ].add_run ( 'Direction' )
	tableno.cell ( 0 , 9 ).paragraphs [ 0 ].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
	tableno.cell ( 0 , 9 ).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

	tableno.rows [ 0 ].cells [ 10 ]._tc.get_or_add_tcPr ().append ( shading_elm_21 )
	# tableno.rows[1].cells[10]._tc.get_or_add_tcPr().append(shading_elm_22)
	tableno.cell ( 0 , 10 ).merge ( tableno.cell ( 1 , 10 ) )
	# tableno.cell(0,10).add_paragraph('Current')
	tableno.cell ( 0 , 10 ).paragraphs [ 0 ].add_run ( 'Current' )
	tableno.cell ( 0 , 10 ).paragraphs [ 0 ].style = document.styles [ 'textBT' ]
	tableno.cell ( 0 , 10 ).paragraphs [ 0 ].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
	tableno.cell ( 0 , 10 ).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

	tableno.cell ( 1 , 0 ).paragraphs [ 0 ].add_run ( 'Position of accelerometers' )
	tableno.cell ( 1 , 0 ).paragraphs [ 0 ].style = document.styles [ 'textBT' ]
	tableno.rows [ 1 ].cells [ 0 ]._tc.get_or_add_tcPr ().append ( shading_elm_1 )

	tableno.cell ( 1 , 1 ).paragraphs [ 0 ].add_run ( 'H1' )
	tableno.cell ( 1 , 1 ).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
	tableno.cell ( 1 , 1 ).paragraphs [ 0 ].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
	tableno.cell ( 1 , 1 ).paragraphs [ 0 ].style = document.styles [ 'textT' ]
	tableno.rows [ 1 ].cells [ 1 ]._tc.get_or_add_tcPr ().append ( shading_elm_2 )

	tableno.cell ( 1 , 2 ).paragraphs [ 0 ].add_run ( 'V1' )
	tableno.cell ( 1 , 2 ).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
	tableno.cell ( 1 , 2 ).paragraphs [ 0 ].style = document.styles [ 'textT' ]
	tableno.cell ( 1 , 2 ).paragraphs [ 0 ].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
	tableno.rows [ 1 ].cells [ 2 ]._tc.get_or_add_tcPr ().append ( shading_elm_3 )

	tableno.cell ( 1 , 3 ).paragraphs [ 0 ].add_run ( 'H2' )
	tableno.cell ( 1 , 3 ).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
	tableno.cell ( 1 , 3 ).paragraphs [ 0 ].style = document.styles [ 'textT' ]
	tableno.cell ( 1 , 3 ).paragraphs [ 0 ].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
	tableno.rows [ 1 ].cells [ 3 ]._tc.get_or_add_tcPr ().append ( shading_elm_4 )

	tableno.cell ( 1 , 4 ).paragraphs [ 0 ].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
	tableno.cell ( 1 , 4 ).paragraphs [ 0 ].add_run ( 'V2' )
	tableno.cell ( 1 , 4 ).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
	tableno.cell ( 1 , 4 ).paragraphs [ 0 ].style = document.styles [ 'textT' ]
	tableno.rows [ 1 ].cells [ 4 ]._tc.get_or_add_tcPr ().append ( shading_elm_5 )

	tableno.cell ( 1 , 5 ).paragraphs [ 0 ].add_run ( 'A2' )
	tableno.cell ( 1 , 5 ).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
	tableno.cell ( 1 , 5 ).paragraphs [ 0 ].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
	tableno.cell ( 1 , 5 ).paragraphs [ 0 ].style = document.styles [ 'textT' ]
	tableno.rows [ 1 ].cells [ 5 ]._tc.get_or_add_tcPr ().append ( shading_elm_6 )

	tableno.cell ( 1 , 6 ).paragraphs [ 0 ].add_run ( 'H3' )
	tableno.cell ( 1 , 6 ).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
	tableno.cell ( 1 , 6 ).paragraphs [ 0 ].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
	tableno.rows [ 1 ].cells [ 6 ]._tc.get_or_add_tcPr ().append ( shading_elm_7 )
	tableno.cell ( 1 , 6 ).paragraphs [ 0 ].style = document.styles [ 'textT' ]

	tableno.cell ( 1 , 7 ).paragraphs [ 0 ].add_run ( 'V3' )
	tableno.cell ( 1 , 7 ).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
	tableno.cell ( 1 , 7 ).paragraphs [ 0 ].style = document.styles [ 'textT' ]
	tableno.cell ( 1 , 7 ).paragraphs [ 0 ].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
	tableno.rows [ 1 ].cells [ 7 ]._tc.get_or_add_tcPr ().append ( shading_elm_8 )

	tableno.cell ( 1 , 8 ).paragraphs [ 0 ].add_run ( 'A3' )
	tableno.cell ( 1 , 8 ).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
	tableno.cell ( 1 , 8 ).paragraphs [ 0 ].style = document.styles [ 'textT' ]
	tableno.cell ( 1 , 8 ).paragraphs [ 0 ].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
	tableno.rows [ 1 ].cells [ 8 ]._tc.get_or_add_tcPr ().append ( shading_elm_9 )

	maxcord = 0
	maxval = 0
	result = [ ]
	cur = conn.cursor ()
	querry = "select value from (select distinct ml.value, ml.point, pts.sort from measurements_low as ml RIGHT JOIN points as pts ON ml.id = pts.id and ml.point = pts.point where type = 'RMS' and ml.id = " + str (
		id1 ) + " and ml.raport_number = '" + str (
		rn ) + "' and ml.point IN (select point from points as pts where id = " + str (
		id1 ) + " order by sort)) as al order by al.sort"
	cur.execute ( querry )
	result = cur.fetchall ()
	conn.commit ()
	cur.close ()

	for j in range ( 5 ):
		if result [ j ] [ 0 ] > float(maxval):
			maxcord = j
			maxval = str ( float ( '%.3f'%result [ j ] [ 0 ] ))

	for j in range ( 5 ):
		result_text = str ( float ( '%.3f'%result [ j ] [ 0 ] ) )
		p = tableno.cell ( 2 , 1 + j ).paragraphs [ 0 ]
		tableno.cell ( 2 , 1 + j ).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
		p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
		r = p.add_run ( result_text.replace ( "." , "," ) )
		r.font.name = 'Calibri'
		r.font.size = Pt ( 9 )
		if maxcord == j:
			r.bold = True
		if result [ j ] [ 0 ] < 2.3:
			limittxt = 'Z. A'
		elif result [ j ] [ 0 ] > 2.3 and result [ j ] [ 0 ] < 4.5:
			limittxt = 'Z. B'
		elif result [ j ] [ 0 ] > 4.5 and result [ j ] [ 0 ] < 7.1:
			limittxt = 'Z. C'
		else:
			limittxt = 'Z. D'
		p = tableno.cell ( 3 , 1 + j ).paragraphs [ 0 ]
		p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
		tableno.cell ( 2 , 1 + j ).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
		r = p.add_run ( limittxt )
		if limittxt == 'Z. A':
			r.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
		elif limittxt == 'Z. B':
			r.font.highlight_color = WD_COLOR_INDEX.YELLOW
		elif limittxt == 'Z. C':
			r.font.highlight_color = WD_COLOR_INDEX.RED
		r.font.name = 'Calibri'
		r.font.size = Pt ( 9 )

	maxcord = 0
	maxval = 0
	result = [ ]

	cur = conn.cursor ()
	querry = "select value from (select distinct ml.value, ml.point, pts.sort from measurements_low as ml RIGHT JOIN points as pts ON ml.id = pts.id and ml.point = pts.point where type = 'RMS' and ml.id = " + str (
		id2 ) + " and ml.raport_number = '" + str (
		rn ) + "' and ml.point IN (select point from points as pts where id = " + str (
		id2 ) + " order by sort)) as al order by al.sort"
	cur.execute ( querry )
	result = cur.fetchall ()
	conn.commit ()
	cur.close ()

	for j in range ( 3 ):
		if result [ j ] [ 0 ] > maxval:
			maxcord = j
			maxval =  float ( '%.3f'%result [ j ] [ 0 ] )

	for j in range ( 3 ):

		result_text = str ( float ( '%.3f'%result [ j ] [ 0 ] ) )
		p = tableno.cell ( 2 , 6 + j ).paragraphs [ 0 ]
		tableno.cell ( 2 , 6 + j ).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
		p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
		r = p.add_run ( result_text.replace ( "." , "," ) )
		r.font.name = 'Calibri'
		r.font.size = Pt ( 9 )
		if maxcord == j:
			r.bold = True
		if result [ j ] [ 0 ] < 4.5:
			limittxt = 'Z. A'
		elif result [ j ] [ 0 ] > 4.5 and result [ j ] [ 0 ] < 11.2:
			limittxt = 'Z. B'
		else:
			limittxt = 'Z. C'
		p = tableno.cell ( 3 , 6 + j ).paragraphs [ 0 ]
		tableno.cell ( 3 , 6 + j ).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
		p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
		r = p.add_run ( limittxt )
		if limittxt == 'Z. A':
			r.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
		elif limittxt == 'Z. B':
			r.font.highlight_color = WD_COLOR_INDEX.YELLOW
		elif limittxt == 'Z. C':
			r.font.highlight_color = WD_COLOR_INDEX.RED
		r.font.name = 'Calibri'
		r.font.size = Pt ( 9 )

	maxcord = 0
	maxval = 0
	result = [ ]
	cur = conn.cursor ()
	querry = "select value from (select distinct ml.value, ml.point, pts.sort from measurements_low as ml RIGHT JOIN points as pts ON ml.id = pts.id and ml.point = pts.point where type = 'envelope P-K' and ml.id = " + str (
		id1 ) + " and ml.raport_number = '" + str (
		rn ) + "' and ml.point IN (select point from points as pts where id = " + str (
		id1 ) + " order by sort)) as al order by al.sort"
	cur.execute ( querry )
	result = cur.fetchall ()
	conn.commit ()
	cur.close ()

	for j in range ( 5 ):
		if result [ j ] [ 0 ] > float(maxval):
			maxcord = j
			maxval = str ( float ( '%.3f'%result [ j ] [ 0 ] ) )

	for j in range ( 5 ):
		result_text = str ( float ( '%.3f'%result [ j ] [ 0 ] )  )
		p = tableno.cell ( 4 , 1 + j ).paragraphs [ 0 ]
		p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
		tableno.cell ( 4 , 1 + j ).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
		r = p.add_run ( result_text.replace ( "." , "," ) )
		r.font.name = 'Calibri'
		r.font.size = Pt ( 9 )
		if maxcord == j:
			r.bold = True

	maxcord = 0
	maxval = 0
	result = [ ]

	cur = conn.cursor ()
	querry = "select value from (select distinct ml.value, ml.point, pts.sort from measurements_low as ml RIGHT JOIN points as pts ON ml.id = pts.id and ml.point = pts.point where type = 'envelope P-K' and ml.id = " + str (
		id2 ) + " and ml.raport_number = '" + str (
		rn ) + "' and ml.point IN (select point from points as pts where id = " + str (
		id2 ) + " order by sort)) as al order by al.sort"
	cur.execute ( querry )
	result = cur.fetchall ()
	conn.commit ()
	cur.close ()

	# messagebox.showinfo("Title", str(querry))
	for j in range ( 3 ):
		if result [ j ] [ 0 ] > maxval:
			maxcord = j
			maxval = float ( '%.3f'%result [ j ] [ 0 ] )

	for j in range ( 3 ):

		result_text = str (  '%.3f'%result [ j ] [ 0 ] )
		p = tableno.cell ( 4 , 6 + j ).paragraphs [ 0 ]
		tableno.cell ( 4 , 6 + j ).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
		p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
		r = p.add_run ( result_text.replace ( "." , "," ) )
		r.font.name = 'Calibri'
		r.font.size = Pt ( 9 )
		if maxcord == j:
			r.bold = True
	# !2!

	maxcord = 0
	maxval = 0
	result = [ ]
	cur = conn.cursor ()
	querry = "select value from (select distinct ml.value, ml.point, pts.sort from measurements_low as ml RIGHT JOIN points as pts ON ml.id = pts.id and ml.point = pts.point where type = 'RMS' and ml.id = " + str (
		id3 ) + " and ml.raport_number = '" + str (
		rn ) + "' and ml.point IN (select point from points as pts where id = " + str (
		id3 ) + " order by sort)) as al order by al.sort"
	cur.execute ( querry )
	result = cur.fetchall ()
	conn.commit ()
	cur.close ()

	for j in range ( 5 ):
		if result [ j ] [ 0 ] > float(maxval):
			maxcord = j
			maxval = str ( float ( '%.3f'%result [ j ] [ 0 ] ) )

	for j in range ( 5 ):
		result_text = str (  float ( '%.3f'%result [ j ] [ 0 ] )  )
		p = tableno.cell ( 5 , 1 + j ).paragraphs [ 0 ]
		p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
		tableno.cell ( 5 , 1 + j ).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
		r = p.add_run ( result_text.replace ( "." , "," ) )
		r.font.name = 'Calibri'
		r.font.size = Pt ( 9 )
		if maxcord == j:
			r.bold = True
		if result [ j ] [ 0 ] < 2.3:
			limittxt = 'Z. A'
		elif result [ j ] [ 0 ] > 2.3 and result [ j ] [ 0 ] < 4.5:
			limittxt = 'Z. B'
		elif result [ j ] [ 0 ] > 4.5 and result [ j ] [ 0 ] < 7.1:
			limittxt = 'Z. C'
		else:
			limittxt = 'Z. D'
		p = tableno.cell ( 6 , 1 + j ).paragraphs [ 0 ]
		p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
		tableno.cell ( 6 , 1 + j ).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
		r = p.add_run ( limittxt )
		if limittxt == 'Z. A':
			r.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
		elif limittxt == 'Z. B':
			r.font.highlight_color = WD_COLOR_INDEX.YELLOW
		elif limittxt == 'Z. C':
			r.font.highlight_color = WD_COLOR_INDEX.RED
		r.font.name = 'Calibri'
		r.font.size = Pt ( 9 )

	maxcord = 0
	maxval = 0
	result = [ ]

	cur = conn.cursor ()
	querry = "select value from (select distinct ml.value, ml.point, pts.sort from measurements_low as ml RIGHT JOIN points as pts ON ml.id = pts.id and ml.point = pts.point where type = 'RMS' and ml.id = " + str (
		id4 ) + " and ml.raport_number = '" + str (
		rn ) + "' and ml.point IN (select point from points as pts where id = " + str (
		id4 ) + " order by sort)) as al order by al.sort"
	cur.execute ( querry )
	result = cur.fetchall ()
	conn.commit ()
	cur.close ()

	# messagebox.showinfo("Title", str(querry))
	for j in range ( 3 ):
		if result [ j ] [ 0 ] > maxval:
			maxcord = j
			maxval = float ( '%.3f'%result [ j ] [ 0 ] )

	for j in range ( 3 ):

		result_text = str ( float ( '%.3f'%result [ j ] [ 0 ] ) )
		p = tableno.cell ( 5 , 6 + j ).paragraphs [ 0 ]
		p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
		tableno.cell ( 5 , 6 + j ).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
		r = p.add_run ( result_text.replace ( "." , "," ) )
		r.font.name = 'Calibri'
		r.font.size = Pt ( 9 )
		if maxcord == j:
			r.bold = True
		if result [ j ] [ 0 ] < 4.5:
			limittxt = 'Z. A'
		elif result [ j ] [ 0 ] > 4.5 and result [ j ] [ 0 ] < 11.2:
			limittxt = 'Z. B'
		else:
			limittxt = 'Z. C'
		p = tableno.cell ( 6 , 6 + j ).paragraphs [ 0 ]
		p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
		r = p.add_run ( limittxt )
		if limittxt == 'Z. A':
			r.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
		if limittxt == 'Z. B':
			r.font.highlight_color = WD_COLOR_INDEX.YELLOW
		if limittxt == 'Z. C':
			r.font.highlight_color = WD_COLOR_INDEX.RED
		r.font.name = 'Calibri'
		r.font.size = Pt ( 9 )

	maxcord = 0
	maxval = 0
	result = [ ]
	cur = conn.cursor ()
	querry = "select value from (select distinct ml.value, ml.point, pts.sort from measurements_low as ml RIGHT JOIN points as pts ON ml.id = pts.id and ml.point = pts.point where type = 'envelope P-K' and ml.id = " + str (
		id3 ) + " and ml.raport_number = '" + str (
		rn ) + "' and ml.point IN (select point from points as pts where id = " + str (
		id3 ) + " order by sort)) as al order by al.sort"
	cur.execute ( querry )
	result = cur.fetchall ()
	conn.commit ()
	cur.close ()

	for j in range ( 5 ):
		if result [ j ] [ 0 ] > float(maxval):
			maxcord = j
			maxval = str ( float ( '%.3f'%result [ j ] [ 0 ] ) )

	for j in range ( 5 ):
		result_text = str ( float ( '%.3f'%result [ j ] [ 0 ] )  )
		p = tableno.cell ( 7 , 1 + j ).paragraphs [ 0 ]
		tableno.cell ( 7 , 1 + j ).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
		p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
		r = p.add_run ( result_text.replace ( "." , "," ) )
		r.font.name = 'Calibri'
		r.font.size = Pt ( 9 )
		if maxcord == j:
			r.bold = True

	maxcord = 0
	maxval = 0
	result = [ ]

	cur = conn.cursor ()
	querry = "select value from (select distinct ml.value, ml.point, pts.sort from measurements_low as ml RIGHT JOIN points as pts ON ml.id = pts.id and ml.point = pts.point where type = 'envelope P-K' and ml.id = " + str (
		id4 ) + " and ml.raport_number = '" + str (
		rn ) + "' and ml.point IN (select point from points as pts where id = " + str (
		id4 ) + " order by sort)) as al order by al.sort"
	cur.execute ( querry )
	result = cur.fetchall ()
	conn.commit ()
	cur.close ()

	# messagebox.showinfo("Title", str(querry))
	for j in range ( 3 ):
		if result [ j ] [ 0 ] > maxval:
			maxcord = j
			maxval = float ( '%.3f'%result [ j ] [ 0 ] )

	for j in range ( 3 ):

		result_text = str (float ( '%.3f'%result [ j ] [ 0 ] )  )
		p = tableno.cell ( 7 , 6 + j ).paragraphs [ 0 ]
		tableno.cell ( 7 , 6 + j ).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
		p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
		r = p.add_run ( result_text.replace ( "." , "," ) )
		r.font.name = 'Calibri'
		r.font.size = Pt ( 9 )
		if maxcord == j:
			r.bold = True

	# !3!

	maxcord = 0
	maxval = 0
	result = [ ]
	cur = conn.cursor ()
	querry = "select value from (select distinct ml.value, ml.point, pts.sort from measurements_low as ml RIGHT JOIN points as pts ON ml.id = pts.id and ml.point = pts.point where type = 'RMS' and ml.id = " + str (
		id5 ) + " and ml.raport_number = '" + str (
		rn ) + "' and ml.point IN (select point from points as pts where id = " + str (
		id5 ) + " order by sort)) as al order by al.sort"
	cur.execute ( querry )
	result = cur.fetchall ()
	conn.commit ()
	cur.close ()

	for j in range ( 5 ):
		if result [ j ] [ 0 ] > maxval:
			maxcord = j
			maxval = round ( float ( '%.3f'%result [ j ] [ 0 ] ) , 3 )

	for j in range ( 5 ):
		result_text = str (  float ( '%.3f'%result [ j ] [ 0 ] ) )
		p = tableno.cell ( 8 , 1 + j ).paragraphs [ 0 ]
		tableno.cell ( 8 , 1 + j ).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
		p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
		r = p.add_run ( result_text.replace ( "." , "," ) )
		r.font.name = 'Calibri'
		r.font.size = Pt ( 9 )
		if maxcord == j:
			r.bold = True
		if result [ j ] [ 0 ] < 2.3:
			limittxt = 'Z. A'
		elif result [ j ] [ 0 ] > 2.3 and result [ j ] [ 0 ] < 4.5:
			limittxt = 'Z. B'
		elif result [ j ] [ 0 ] > 4.5 and result [ j ] [ 0 ] < 7.1:
			limittxt = 'Z, C'
		else:
			limittxt = 'Z. D'
		p = tableno.cell ( 9 , 1 + j ).paragraphs [ 0 ]
		tableno.cell ( 9 , 1 + j ).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
		p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
		r = p.add_run ( limittxt )
		if limittxt == 'Z. A':
			r.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
		elif limittxt == 'Z. B':
			r.font.highlight_color = WD_COLOR_INDEX.YELLOW
		elif limittxt == 'Z. C':
			r.font.highlight_color = WD_COLOR_INDEX.RED
		r.font.name = 'Calibri'
		r.font.size = Pt ( 9 )

	maxcord = 0
	maxval = 0
	result = [ ]

	cur = conn.cursor ()
	querry = "select value from (select distinct ml.value, ml.point, pts.sort from measurements_low as ml RIGHT JOIN points as pts ON ml.id = pts.id and ml.point = pts.point where type = 'RMS' and ml.id = " + str (
		id6 ) + " and ml.raport_number = '" + str (
		rn ) + "' and ml.point IN (select point from points as pts where id = " + str (
		id6 ) + " order by sort)) as al order by al.sort"
	cur.execute ( querry )
	result = cur.fetchall ()
	conn.commit ()
	cur.close ()

	# messagebox.showinfo("Title", str(querry))
	for j in range ( 3 ):
		if result [ j ] [ 0 ] > maxval:
			maxcord = j
			maxval =  float ( '%.3f'%result [ j ] [ 0 ] )

	for j in range ( 3 ):

		result_text = str ( float ( '%.3f'%result [ j ] [ 0 ] ) )
		p = tableno.cell ( 8 , 6 + j ).paragraphs [ 0 ]
		tableno.cell ( 8 , 6 + j ).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
		p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
		r = p.add_run ( result_text.replace ( "." , "," ) )
		r.font.name = 'Calibri'
		r.font.size = Pt ( 9 )
		if maxcord == j:
			r.bold = True
		if result [ j ] [ 0 ] < 4.5:
			limittxt = 'Z. A'
		elif result [ j ] [ 0 ] > 4.5 and result [ j ] [ 0 ] < 11.2:
			limittxt = 'Z. B'
		else:
			limittxt = 'Z. C'
		p = tableno.cell ( 9 , 6 + j ).paragraphs [ 0 ]
		tableno.cell ( 9 , 6 + j ).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
		p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
		r = p.add_run ( limittxt )
		if limittxt == 'Z. A':
			r.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
		elif limittxt == 'Z. B':
			r.font.highlight_color = WD_COLOR_INDEX.YELLOW
		elif limittxt == 'Z. C':
			r.font.highlight_color = WD_COLOR_INDEX.RED
		r.font.name = 'Calibri'
		r.font.size = Pt ( 9 )

	maxcord = 0
	maxval = 0
	result = [ ]
	cur = conn.cursor ()
	querry = "select value from (select distinct ml.value, ml.point, pts.sort from measurements_low as ml RIGHT JOIN points as pts ON ml.id = pts.id and ml.point = pts.point where type = 'envelope P-K' and ml.id = " + str (
		id5 ) + " and ml.raport_number = '" + str (
		rn ) + "' and ml.point IN (select point from points as pts where id = " + str (
		id5 ) + " order by sort)) as al order by al.sort"
	cur.execute ( querry )
	result = cur.fetchall ()
	conn.commit ()
	cur.close ()

	for j in range ( 5 ):
		if result [ j ] [ 0 ] > float(maxval):
			maxcord = j
			maxval = str( float ( '%.3f'%result [ j ] [ 0 ] )  )

	for j in range ( 5 ):
		result_text = str ( float ( '%.3f'%result [ j ] [ 0 ] )  )
		p = tableno.cell ( 10 , 1 + j ).paragraphs [ 0 ]
		tableno.cell ( 10 , 1 + j ).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
		p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
		r = p.add_run ( result_text.replace ( "." , "," ) )
		r.font.name = 'Calibri'
		r.font.size = Pt ( 9 )
		if maxcord == j:
			r.bold = True

	maxcord = 0
	maxval = 0
	result = [ ]

	cur = conn.cursor ()
	querry = "select value from (select distinct ml.value, ml.point, pts.sort from measurements_low as ml RIGHT JOIN points as pts ON ml.id = pts.id and ml.point = pts.point where type = 'envelope P-K' and ml.id = " + str (
		id6 ) + " and ml.raport_number = '" + str (
		rn ) + "' and ml.point IN (select point from points as pts where id = " + str (
		id6 ) + " order by sort)) as al order by al.sort"
	cur.execute ( querry )
	result = cur.fetchall ()
	conn.commit ()
	cur.close ()

	# messagebox.showinfo("Title", str(querry))
	for j in range ( 3 ):
		if result [ j ] [ 0 ] > maxval:
			maxcord = j
			maxval =float ( '%.3f'%result [ j ] [ 0 ] )

	for j in range ( 3 ):

		result_text = str (  float ( '%.3f'%result [ j ] [ 0 ] ) )
		p = tableno.cell ( 10 , 6 + j ).paragraphs [ 0 ]
		tableno.cell ( 10 , 6 + j ).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
		p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
		r = p.add_run ( result_text.replace ( "." , "," ) )
		r.font.name = 'Calibri'
		r.font.size = Pt ( 9 )
		if maxcord == j:
			r.bold = True

	RMSov = tableno.cell ( 2 , 0 ).paragraphs [ 0 ].add_run ( 'Overall velocity RMS (mm/s)' )
	RMSov = tableno.cell ( 2 , 0 ).paragraphs [ 0 ].style = document.styles [ 'textBT' ]

	tableno.rows [ 2 ].cells [ 0 ]._tc.get_or_add_tcPr ().append ( shading_elm_10 )
	tableno.cell ( 3 , 0 ).paragraphs [ 0 ].add_run ( 'Vibration Class' )
	tableno.cell ( 3 , 0 ).paragraphs [ 0 ].style = document.styles [ 'textBT' ]

	tableno.rows [ 3 ].cells [ 0 ]._tc.get_or_add_tcPr ().append ( shading_elm_11 )

	tableno.cell ( 4 , 0 ).paragraphs [ 0 ].add_run ( 'Envelope 0-P value(mm/s²)' )
	tableno.cell ( 4 , 0 ).paragraphs [ 0 ].style = document.styles [ 'textBT' ]

	tableno.rows [ 4 ].cells [ 0 ]._tc.get_or_add_tcPr ().append ( shading_elm_12 )
	tableno.cell ( 2 , 9 ).merge ( tableno.cell ( 3 , 9 ).merge ( tableno.cell ( 4 , 9 ) ) )
	tableno.cell ( 2 , 9 ).paragraphs [ 0 ].add_run ( 'SB' )
	tableno.cell ( 2 , 9 ).paragraphs [ 0 ].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
	tableno.cell ( 2 , 9 ).paragraphs [ 0 ].style = document.styles [ 'textBT' ]
	tableno.cell ( 2 , 9 ).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
	# tableno.cell(2,9).add_paragraph('[INSERT]','textBT').paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

	tableno.cell ( 2 , 10 ).merge ( tableno.cell ( 3 , 10 ).merge ( tableno.cell ( 4 , 10 ) ) )
	tableno.cell ( 2 , 10 ).paragraphs [ 0 ].add_run ( '[INSERT]' )
	tableno.cell ( 2 , 10 ).paragraphs [ 0 ].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
	tableno.cell ( 2 ,10 ).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
	tableno.cell ( 2 , 10 ).paragraphs [ 0 ].style = document.styles [ 'textBT' ]

	tableno.cell ( 5 , 0 ).paragraphs [ 0 ].add_run ( 'Overall velocity RMS (mm/s)' )
	tableno.cell ( 5 , 0 ).paragraphs [ 0 ].style = document.styles [ 'textBT' ]

	tableno.rows [ 5 ].cells [ 0 ]._tc.get_or_add_tcPr ().append ( shading_elm_13 )
	tableno.cell ( 6 , 0 ).paragraphs [ 0 ].add_run ( 'Vibration Class' )
	tableno.cell ( 6 , 0 ).paragraphs [ 0 ].style = document.styles [ 'textBT' ]
	tableno.rows [ 6 ].cells [ 0 ]._tc.get_or_add_tcPr ().append ( shading_elm_14 )

	tableno.cell ( 7 , 0 ).paragraphs [ 0 ].add_run ( 'Envelope 0-P value(mm/s²)' )
	tableno.cell ( 7 , 0 ).paragraphs [ 0 ].style = document.styles [ 'textBT' ]

	tableno.rows [ 7 ].cells [ 0 ]._tc.get_or_add_tcPr ().append ( shading_elm_15 )
	tableno.cell ( 5 , 9 ).merge ( tableno.cell ( 6 , 9 ).merge ( tableno.cell ( 7 , 9 ) ) )
	tableno.cell ( 5 , 9 ).paragraphs [ 0 ].add_run ( 'PS' )
	# tableno.cell(5,9)
	tableno.cell ( 5 , 9 ).paragraphs [ 0 ].style = document.styles [ 'textBT' ]
	tableno.cell ( 5 , 9 ).paragraphs [ 0 ].runs [ 0 ].add_break ()
	tableno.cell ( 5 , 9 ).paragraphs [ 0 ].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
	tableno.cell ( 5 , 9 ).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

	tableno.cell ( 5 , 10 ).merge ( tableno.cell ( 6 , 10 ).merge ( tableno.cell ( 7 , 10 ) ) )
	tableno.cell ( 5 , 10 ).paragraphs [ 0 ].add_run ( '[INSERT]' )
	tableno.cell ( 5 , 10 ).style = document.styles [ 'textBT' ]
	tableno.cell ( 5 , 10 ).paragraphs [ 0 ].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

	tableno.cell ( 8 , 0 ).paragraphs [ 0 ].add_run ( 'Overall velocity RMS (mm/s)' )
	tableno.cell ( 8 , 0 ).paragraphs [ 0 ].style = document.styles [ 'textBT' ]
	tableno.rows [ 8 ].cells [ 0 ]._tc.get_or_add_tcPr ().append ( shading_elm_16 )
	tableno.cell ( 9 , 0 ).paragraphs [ 0 ].add_run ( 'Vibration Class' )
	tableno.cell ( 9 , 0 ).paragraphs [ 0 ].style = document.styles [ 'textBT' ]
	tableno.rows [ 9 ].cells [ 0 ]._tc.get_or_add_tcPr ().append ( shading_elm_17 )
	tableno.cell ( 10 , 0 ).paragraphs [ 0 ].add_run ( 'Envelope 0-P value(mm/s²)' )
	tableno.cell ( 10 , 0 ).paragraphs [ 0 ].style = document.styles [ 'textBT' ]

	tableno.rows [ 10 ].cells [ 0 ]._tc.get_or_add_tcPr ().append ( shading_elm_18 )

	tableno.cell ( 8 , 9 ).merge ( tableno.cell ( 9 , 9 ).merge ( tableno.cell ( 10 , 9 ) ) )
	tableno.cell ( 8 , 9 ).paragraphs [ 0 ].add_run ( 'IDLE' )
	tableno.cell ( 8 , 9 ).paragraphs [ 0 ].runs [ 0 ].add_break ()
	tableno.cell ( 8 , 9 ).paragraphs [ 0 ].style = document.styles [ 'textBT' ]
	tableno.cell ( 8 , 9 ).paragraphs [ 0 ].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
	tableno.cell ( 8 , 9 ).paragraphs [ 0 ].add_run ( '0%' )
	tableno.cell ( 8 , 9 ).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

	tableno.cell ( 8 , 10 ).merge ( tableno.cell ( 9 , 10 ).merge ( tableno.cell ( 10 , 10 ) ) )
	tableno.cell ( 8 , 10 ).add_paragraph ( '[INSERT]' )
	tableno.cell(8,10).paragraphs[0].style=document.styles['textBT']
	tableno.cell ( 8 , 10 ).paragraphs [ 0 ].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER


def getname ( host , username , password ):
	kport = "5432"
	kdb = "postgres"
	cs = "dbname=%s user=%s password=%s host=%s port=%s"%(kdb , username , password , host , kport)
	conn = None
	conn = psycopg2.connect ( str ( cs ) )
	cur = conn.cursor ()
	querry = "select full_name from users where login = '" + str ( username ) + "' limit 1;"
	cur.execute ( querry )
	result = cur.fetchall ()
	conn.commit ()
	cur.close ()
	return result


def getini ( host , username , password ):
	kport = "5432"
	kdb = "postgres"
	cs = "dbname=%s user=%s password=%s host=%s port=%s"%(kdb , username , password , host , kport)
	conn = None
	conn = psycopg2.connect ( str ( cs ) )
	cur = conn.cursor ()
	querry = "select ini from users where login = '" + str ( username ) + "' limit 1;"
	cur.execute ( querry )
	result = cur.fetchall ()
	conn.commit ()
	cur.close ()
	return result


def makereport ( username , password , host , rn_ , id_ ):  # AUTORUN!

	progress = Progressbar ( tk , orient=HORIZONTAL , length=200 , mode='determinate' )
	progress [ 'value' ] = 50
	tk.update_idletasks ()
	progress.pack ()

	h1 = document.add_paragraph ( 'Diagnostic report' + ' ' + rn_ )
	h1.style = document.styles [ 'Head' ]
	h1.alignment = WD_ALIGN_PARAGRAPH.CENTER

	document.add_paragraph ()

	l1 = document.add_paragraph ( '1. General information' )
	l1.style = document.styles [ 'listlvl1' ]
	l1.alignment = WD_ALIGN_PARAGRAPH.LEFT
	l1.paragraph_format.left_indent = Inches ( 0.25 )

	l11 = document.add_paragraph ( '1.1.	Diagnostic report' )
	l11.style = document.styles [ 'listlvl2' ]
	l11.alignment = WD_ALIGN_PARAGRAPH.LEFT
	l11.paragraph_format.left_indent = Inches ( 0.25 )

	document.add_paragraph ()

	table = document.add_table ( rows=3 , cols=2 )
	# table.style = 'Table Grid'

	for cell in table.columns [ 0 ].cells:
		cell.width = Inches ( 1.7 )

	for cell in table.columns [ 1 ].cells:
		cell.width = Inches ( 1.7 )

	table.cell ( 0 , 0 ).add_paragraph ( 'Info Marine Case No.: ' , 'texthead' )
	table.cell ( 0 , 0 ).alignment = WD_ALIGN_PARAGRAPH.LEFT

	table.cell ( 0 , 1 ).add_paragraph ( 'INSERT CASE NUMBER' , 'text' )
	table.cell ( 0 , 1 ).alignment = WD_ALIGN_PARAGRAPH.LEFT

	table.cell ( 1 , 0 ).add_paragraph ( 'Report Number: ' , 'texthead' )
	table.cell ( 1 , 0 ).alignment = WD_ALIGN_PARAGRAPH.LEFT

	table.cell ( 1 , 1 ).add_paragraph ( rn_ , 'text' )
	table.cell ( 1 , 1 ).alignment = WD_ALIGN_PARAGRAPH.LEFT

	table.cell ( 2 , 0 ).add_paragraph ( 'Date of measurement:' , 'texthead' )
	table.cell ( 2 , 0 ).alignment = WD_ALIGN_PARAGRAPH.LEFT

	table.cell ( 2 , 1 ).add_paragraph ( 'DATE' , 'text' )
	table.cell ( 2 , 1 ).alignment = WD_ALIGN_PARAGRAPH.LEFT

	document.add_paragraph ()
	document.add_paragraph ()

	l12 = document.add_paragraph ( '1.2.	Platform Information' )
	l12.style = document.styles [ 'listlvl2' ]
	l12.alignment = WD_ALIGN_PARAGRAPH.LEFT
	l12.paragraph_format.left_indent = Inches ( 0.25 )

	document.add_paragraph ()

	table2 = document.add_table ( rows=8 , cols=2 )
	# table.style = 'Table Grid'

	for cell in table2.columns [ 0 ].cells:
		cell.width = Inches ( 1.7 )

	for cell in table2.columns [ 1 ].cells:
		cell.width = Inches ( 1.7 )

	table2.cell ( 0 , 0 ).add_paragraph ( 'Name of platform' , 'texthead' )
	table2.cell ( 0 , 0 ).alignment = WD_ALIGN_PARAGRAPH.LEFT

	if id_ == '53':
		platname = 'Atlantis'
	elif id_ == '54':
		platname = 'Iolair'
	elif id_ == '55':
		platname = 'Neptuno'

	table2.cell ( 0 , 1 ).add_paragraph ( str ( platname ) , 'text' )
	table2.cell ( 0 , 1 ).paragraphs[0].style=document.styles['text']
	table2.cell ( 0 , 1 ).alignment = WD_ALIGN_PARAGRAPH.LEFT

	table2.cell ( 1 , 0 ).add_paragraph ( 'The IMO Number:' , 'texthead' )
	table2.cell ( 1 , 0 ).alignment = WD_ALIGN_PARAGRAPH.LEFT

	table2.cell ( 1 , 1 ).add_paragraph ( '7816070' , 'text' )
	table2.cell ( 1 , 1 ).alignment = WD_ALIGN_PARAGRAPH.LEFT

	table2.cell ( 2 , 0 ).add_paragraph ( 'Owner of platform:' , 'texthead' )
	table2.cell ( 2 , 0 ).alignment = WD_ALIGN_PARAGRAPH.LEFT

	table2.cell ( 2 , 1 ).add_paragraph ( 'Cotemar S.A. de C.V.' , 'text' )
	table2.cell ( 2 , 1 ).alignment = WD_ALIGN_PARAGRAPH.LEFT

	table2.cell ( 3 , 0 ).add_paragraph ( 'Main dimensions:' , 'textheadU' )
	table2.cell ( 3 , 0 ).alignment = WD_ALIGN_PARAGRAPH.LEFT

	table2.cell ( 4 , 0 ).add_paragraph ( 'Length overall:' , 'texthead' )
	table2.cell ( 4 , 0 ).alignment = WD_ALIGN_PARAGRAPH.LEFT

	table2.cell ( 4 , 1 ).add_paragraph ( '102m' , 'text' )
	table2.cell ( 4 , 1 ).alignment = WD_ALIGN_PARAGRAPH.LEFT

	table2.cell ( 5 , 0 ).add_paragraph ( 'Breadth Extreme:' , 'texthead' )
	table2.cell ( 5 , 0 ).alignment = WD_ALIGN_PARAGRAPH.LEFT

	table2.cell ( 5 , 1 ).add_paragraph ( '76,5m' , 'text' )
	table2.cell ( 5 , 1 ).alignment = WD_ALIGN_PARAGRAPH.LEFT

	table2.cell ( 6 , 0 ).add_paragraph ( 'Deadweight:' , 'texthead' )
	table2.cell ( 7 , 0 ).alignment = WD_ALIGN_PARAGRAPH.LEFT

	table2.cell ( 6 , 1 ).add_paragraph ( '17000t' , 'text' )
	table2.cell ( 6 , 1 ).alignment = WD_ALIGN_PARAGRAPH.LEFT

	table2.cell ( 7 , 0 ).add_paragraph ( 'Propulsion: ' , 'texthead' )
	table2.cell ( 7 , 0 ).alignment = WD_ALIGN_PARAGRAPH.LEFT

	table2.cell ( 7 , 1 ).add_paragraph ( '4 Thrusters' , 'text' )
	table2.cell ( 7 , 1 ).alignment = WD_ALIGN_PARAGRAPH.LEFT

	document.add_page_break ()

	document.add_paragraph ()

	l2 = document.add_paragraph ( '2. Measurement' )
	l2.style = document.styles [ 'listlvl1' ]
	l2.alignment = WD_ALIGN_PARAGRAPH.LEFT
	l2.paragraph_format.left_indent = Inches ( 0.25 )

	l21 = document.add_paragraph ( '2.1.	Specification of measuring position' )
	l21.style = document.styles [ 'listlvl2' ]
	l21.alignment = WD_ALIGN_PARAGRAPH.LEFT
	l21.paragraph_format.left_indent = Inches ( 0.25 )

	document.add_paragraph ()

	t1 = document.add_paragraph ( 'Vibration measurement were performed in three axes.' )
	t1.style = document.styles [ 'text' ]
	t1.alignment = WD_ALIGN_PARAGRAPH.LEFT

	document.add_paragraph ()

	t2 = document.add_paragraph ( 'V – Vertical axis' )
	t2.style = document.styles [ 'text' ]
	t2.alignment = WD_ALIGN_PARAGRAPH.LEFT

	document.add_paragraph ()

	t3 = document.add_paragraph ( 'H – Transverse axis' )
	t3.style = document.styles [ 'text' ]
	t3.alignment = WD_ALIGN_PARAGRAPH.LEFT

	document.add_paragraph ()

	t4 = document.add_paragraph ( 'A – Longitudinal axis' )
	t4.style = document.styles [ 'text' ]
	t4.alignment = WD_ALIGN_PARAGRAPH.LEFT

	# OBRAZEK KTORY WYKRZACZA WORDA
	p = document.add_paragraph ()
	r = p.add_run ()
	r.add_picture ( 'C:\\overmind\\Data\\Iloar_thruster.jpg' )
	p.alignment = WD_ALIGN_PARAGRAPH.CENTER

	des1 = document.add_paragraph ( 'Fig.1 Localization of measuring points' )
	des1.style = document.styles [ 'description' ]
	des1.alignment = WD_ALIGN_PARAGRAPH.CENTER

	document.add_page_break ()
	l22 = document.add_paragraph ( '2.2.	Measurement standard' )
	l22.style = document.styles [ 'listlvl2' ]
	l22.alignment = WD_ALIGN_PARAGRAPH.LEFT
	l22.paragraph_format.left_indent = Inches ( 0.25 )

	t5 = document.add_paragraph ( 'Following standards are applied for assessment:' )
	t5.style = document.styles [ 'text' ]
	t5.alignment = WD_ALIGN_PARAGRAPH.LEFT

	document.add_paragraph ()

	t6 = document.add_paragraph ( 'ISO 10816 – 1	General guidelines' )
	t6.style = document.styles [ 'text' ]
	t6.alignment = WD_ALIGN_PARAGRAPH.LEFT

	document.add_paragraph ()

	t7 = document.add_paragraph ( 'DNV Rules of classification of ships. Part 6 Chapter 15' )
	t7.style = document.styles [ 'text' ]
	t7.alignment = WD_ALIGN_PARAGRAPH.LEFT

	document.add_paragraph ()

	table3 = document.add_table ( rows=4 , cols=2 )
	# table.style = 'Table Grid'

	for cell in table3.columns [ 0 ].cells:
		cell.width = Inches ( 1.7 )

	for cell in table3.columns [ 1 ].cells:
		cell.width = Inches ( 1.7 )

	table3.cell ( 0 , 0 ).add_paragraph ( 'The criteria(Motor):' , 'textheadU' )
	table3.cell ( 0 , 0 ).alignment = WD_ALIGN_PARAGRAPH.LEFT

	table3.cell ( 0 , 0 ).merge ( table3.cell ( 0 , 1 ) )

	table3.cell ( 1 , 0 ).add_paragraph ( 'Value:' , 'texthead' )
	table3.cell ( 1 , 0 ).alignment = WD_ALIGN_PARAGRAPH.LEFT
	table3.cell ( 2 , 0 ).add_paragraph ( 'Range: ' , 'texthead' )
	table3.cell ( 2 , 1 ).add_paragraph ( '2-1000Hz' , 'text' )
	vel = table3.cell ( 1 , 1 ).add_paragraph ( 'Velocity' , 'text' )
	table3.cell ( 1 , 1 ).alignment = WD_ALIGN_PARAGRAPH.LEFT
	vel.add_run ().add_break ()
	vel.add_run ( 'Zone A:        0-2,3 mm/s' )
	vel.add_run ().add_break ()
	vel.add_run ( 'Zone B:        2,3-4,5 mm/s' )
	vel.add_run ().add_break ()
	vel.add_run ( 'Zone C:        4,5-7,1 mm/s' )
	vel.add_run ().add_break ()
	vel.add_run ( 'Zone D:        above 7,1 mm/s' )

	table3.cell ( 1 , 1 ).alignment = WD_ALIGN_PARAGRAPH.LEFT
	set_col_widths ( table3 )
	document.add_paragraph ()

	table4 = document.add_table ( rows=4 , cols=2 )
	# table.style = 'Table Grid'

	for cell in table4.columns [ 0 ].cells:
		cell.width = Inches ( 1.7 )

	for cell in table4.columns [ 1 ].cells:
		cell.width = Inches ( 1.7 )

	table4.cell ( 0 , 0 ).add_paragraph ( 'The criteria(Gear):' , 'textheadU' )
	table4.cell ( 0 , 0 ).alignment = WD_ALIGN_PARAGRAPH.LEFT

	table4.cell ( 0 , 0 ).merge ( table4.cell ( 0 , 1 ) )

	table4.cell ( 1 , 0 ).add_paragraph ( 'Value:' , 'texthead' )
	table4.cell ( 1 , 0 ).alignment = WD_ALIGN_PARAGRAPH.LEFT
	velgear = table4.cell ( 1 , 1 ).add_paragraph ( 'Velocity: ' , 'text' )
	table4.cell ( 2 , 0 ).add_paragraph ( 'Range: ' , 'texthead' )
	table4.cell ( 2 , 1 ).add_paragraph ( '5-1000Hz' , 'text' )
	velgear.add_run ().add_break ()
	velgear.add_run ( 'Zone A:        0-4,5 mm/s' )
	velgear.add_run ().add_break ()
	velgear.add_run ( 'Zone B:        4,5-11,2 mm/s' )
	velgear.add_run ().add_break ()
	velgear.add_run ( 'Zone C:        above 11,2 mm/s' )

	table4.cell ( 1 , 1 ).alignment = WD_ALIGN_PARAGRAPH.LEFT
	set_col_widths ( table4 )

	document.add_paragraph ( 'Legend according to vibration class' , 'listlvl1' )
	limittablelegend = document.add_table ( rows=5 , cols=2 )
	limittablelegend.cell ( 0 , 0 ).paragraphs [ 0 ].add_run (
		'Z. A' ).font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
	limittablelegend.cell ( 0 , 1 ).paragraphs [ 0 ].add_run ( 'Newly commissioned' )

	limittablelegend.cell ( 1 , 0 ).paragraphs [ 0 ].add_run (
		'Z. B' ).font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
	limittablelegend.cell ( 1 , 1 ).paragraphs [ 0 ].add_run ( 'Unrestricted' )

	limittablelegend.cell ( 2 , 0 ).paragraphs [ 0 ].add_run ( 'Z. C' ).font.highlight_color = WD_COLOR_INDEX.YELLOW
	limittablelegend.cell ( 2 , 1 ).paragraphs [ 0 ].add_run ( 'Restricted long-term operation' )

	limittablelegend.cell ( 3 , 0 ).paragraphs [ 0 ].add_run ( 'Z. D' ).font.highlight_color = WD_COLOR_INDEX.RED
	limittablelegend.cell ( 3 , 1 ).paragraphs [ 0 ].add_run ( 'High probability of damage, action required' )

	limittablelegend.cell ( 4 , 0 ).paragraphs [ 0 ].add_run ( 'Z. D' ).font.bold = True
	limittablelegend.cell ( 4 , 1 ).paragraphs [ 0 ].add_run (
		'Vibrations over the limits but actions are not required.' )
	set_col_widths ( limittablelegend )
	for i in range(5):
		limittablelegend.cell(i,0).paragraphs[0].runs[0].font.size=Pt(11)
		limittablelegend.cell(i,1).paragraphs[0].runs[0].font.size=Pt(11)
		limittablelegend.cell(i,0).paragraphs[0].runs[0].font.name='Calibri'
		limittablelegend.cell(i,1).paragraphs[0].runs[0].font.name='Calibri'
	#limittablelegend.paragraphs[0].style=document.styles['text']
	document.add_page_break ()
	l23 = document.add_paragraph ( '2.3.	Measurement equipment' )
	l23.style = document.styles [ 'listlvl2' ]
	l23.alignment = WD_ALIGN_PARAGRAPH.LEFT
	l23.paragraph_format.left_indent = Inches ( 0.25 )

	# document.add_paragraph()

	# t8 = document.add_paragraph('WBK18  8-channel dynamic signal condition modules')
	# t8.style = document.styles['text']
	# t8.alignment = WD_ALIGN_PARAGRAPH.LEFT

	# document.add_paragraph()

	# t10 = document.add_paragraph('eZ – TOMAS (Total Online Monitoring and Analysis Software)')
	# t10.style = document.styles['text']
	# t10.alignment = WD_ALIGN_PARAGRAPH.LEFT

	document.add_paragraph ()

	table5 = document.add_table ( rows=6 , cols=2 )
	# table.style = 'Table Grid'

	for cell in table5.columns [ 0 ].cells:
		cell.width = Inches ( 1.7 )

	for cell in table5.columns [ 1 ].cells:
		cell.width = Inches ( 1.7 )

	table5.cell ( 0 , 0 ).add_paragraph ( 'Device:' , 'text' )
	table5.cell ( 0 , 0 ).alignment = WD_ALIGN_PARAGRAPH.LEFT

	table5.cell ( 0 , 1 ).add_paragraph ( 'Marvib DC750' , 'text' )
	table5.cell ( 0 , 1 ).alignment = WD_ALIGN_PARAGRAPH.LEFT

	table5.cell ( 1 , 0 ).add_paragraph ( 'Serial number :' , 'text' )
	table5.cell ( 1 , 0 ).alignment = WD_ALIGN_PARAGRAPH.LEFT

	table5.cell ( 1 , 1 ).add_paragraph ( 'DC750334' , 'text' )
	table5.cell ( 1 , 1 ).alignment = WD_ALIGN_PARAGRAPH.LEFT

	table5.cell ( 2 , 0 ).add_paragraph ( 'Date of manufacture:' , 'text' )
	table5.cell ( 2 , 0 ).alignment = WD_ALIGN_PARAGRAPH.LEFT

	table5.cell ( 2 , 1 ).add_paragraph ( '15.10.2017' , 'text' )
	table5.cell ( 2 , 1 ).alignment = WD_ALIGN_PARAGRAPH.LEFT

	table5.cell ( 3 , 0 ).add_paragraph ( 'Measuring range:' , 'text' )
	table5.cell ( 3 , 0 ).alignment = WD_ALIGN_PARAGRAPH.LEFT

	table5.cell ( 3 , 1 ).add_paragraph ( '2Hz-30kHz' , 'text' )
	table5.cell ( 3 , 1 ).alignment = WD_ALIGN_PARAGRAPH.LEFT

	table5.cell ( 4 , 0 ).add_paragraph ( 'Indication error:' , 'text' )
	table5.cell ( 4 , 0 ).alignment = WD_ALIGN_PARAGRAPH.LEFT

	table5.cell ( 4 , 1 ).add_paragraph ( '±0,5%' , 'text' )
	table5.cell ( 4 , 1 ).alignment = WD_ALIGN_PARAGRAPH.LEFT

	table5.cell ( 5 , 0 ).add_paragraph ( 'Window applied:' , 'text' )
	table5.cell ( 5 , 0 ).alignment = WD_ALIGN_PARAGRAPH.LEFT

	table5.cell ( 5 , 1 ).add_paragraph ( 'Hanning' , 'text' )
	table5.cell ( 5 , 1 ).alignment = WD_ALIGN_PARAGRAPH.LEFT

	document.add_paragraph ()

	l24 = document.add_paragraph ( '2.4.	Measurement condition' )
	l24.style = document.styles [ 'listlvl2' ]
	l24.alignment = WD_ALIGN_PARAGRAPH.LEFT
	l24.paragraph_format.left_indent = Inches ( 0.25 )

	document.add_paragraph ()

	t11 = document.add_paragraph ( 'Measurement were performed during normal operation of platform.' )
	t11.style = document.styles [ 'text' ]
	t11.alignment = WD_ALIGN_PARAGRAPH.LEFT

	document.add_paragraph ()

	l25 = document.add_paragraph ( '2.5.	Measurement results' )
	l25.style = document.styles [ 'listlvl2' ]
	l25.alignment = WD_ALIGN_PARAGRAPH.LEFT
	l25.paragraph_format.left_indent = Inches ( 0.25 )

	document.add_paragraph ()

	t12 = document.add_paragraph (
		'In below tables are presented RMS velocity values for all measured points of electric motor and gear part. Additionally, FFT graphs for highest values on electric motor and gear part are included.' )
	t12.style = document.styles [ 'text' ]
	t12.alignment = WD_ALIGN_PARAGRAPH.LEFT

	motors = [ 18116 , 18114 , 18112 , 18122 , 18120 , 18118 , 18104 , 18102 , 18100 , 18110 , 18108 , 18106 ]
	thrusters = [ 18117 , 18115 , 18113 , 18123 , 18121 , 18119 , 18105 , 18103 , 18101 , 18111 , 18109 , 18107 ]

	for j in range ( 4 ):
		maxP = 'MAXPOINT'
		progress [ 'value' ] = 50 + j*10
		tk.update_idletasks ()

		document.add_page_break ()
		document.add_paragraph ()
		table6 = document.add_table ( rows=11 , cols=11 )
		resulttable_iloar ( j , host , username , password , table6 , motors [ (((j)*3)) ] , thrusters [ (((j)*3)) ] ,
							motors [ (((j)*3) + 1) ] , thrusters [ (((j)*3) + 1) ] , motors [ (((j)*3) + 2) ] ,
							thrusters [ (((j)*3) + 2) ] , rn_ )

		if j == 0:
			thname = 'SB Thruster AFT'
		elif j == 1:
			thname = 'SB Thruster FWD'
		elif j == 2:
			thname = 'PS Thruster AFT'
		elif j == 3:
			thname = 'PS Thruster FWD'

		chart = document.add_paragraph ()
		r = chart.add_run ()
		maxPoint = createchart ( host , username , password , motors [ (((j)*3)) ] , rn_ , 'RMS' )
		r.add_picture ( 'C:\\overmind\\temp\\tempchart.png' , width=Inches ( 7.0 ) )
		os.remove ( 'C:\\overmind\\temp\\tempchart.png' )
		r = chart.add_run ( ' \n Fig.' + str ( (j*12) + 2 ) + ' FFT velocity graph point ' + str (
			maxPoint ) + ' on ' + thname + '(SB turn) el. motor' )
		# r=chart.add_run(str(motors[(((j)*3))]))
		r.font.name = 'Calibri'
		r.font.size = Pt ( 9 )
		r.font.italic = True
		chart.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

		chart = document.add_paragraph ()
		r = chart.add_run ()
		maxPoint = createchart ( host , username , password , motors [ (((j)*3)) ] , rn_ , 'envelope P-K' )
		r.add_picture ( 'C:\\overmind\\temp\\tempchart.png' , width=Inches ( 7.0 ) )
		os.remove ( 'C:\\overmind\\temp\\tempchart.png' )
		r = chart.add_run ( ' \n Fig.' + str ( (j*12) + 3 ) + ' FFT envelope graph point ' + str (
			maxPoint ) + ' on ' + thname + '(SB turn) el. motor' )
		# r=chart.add_run(str(motors[(((j)*3))]))
		r.font.name = 'Calibri'
		r.font.size = Pt ( 9 )
		r.font.italic = True
		chart.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

		chart = document.add_paragraph ()
		r = chart.add_run ()
		maxPoint = createchart ( host , username , password , thrusters [ (((j)*3)) ] , rn_ , 'RMS' )
		r.add_picture ( 'C:\\overmind\\temp\\tempchart.png' , width=Inches ( 7.0 ) )
		os.remove ( 'C:\\overmind\\temp\\tempchart.png' )
		r = chart.add_run ( ' \n Fig.' + str ( (j*12) + 4 ) + ' FFT velocity graph point ' + str (
			maxPoint ) + ' on ' + thname + '(SB turn) gear' )
		# r=chart.add_run(str(thrusters[(((j)*3))]))
		r.font.name = 'Calibri'
		r.font.size = Pt ( 9 )
		r.font.italic = True
		chart.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

		chart = document.add_paragraph ()
		r = chart.add_run ()
		maxPoint = createchart ( host , username , password , thrusters [ (((j)*3)) ] , rn_ , 'envelope P-K' )
		r.add_picture ( 'C:\\overmind\\temp\\tempchart.png' , width=Inches ( 7.0 ) )
		os.remove ( 'C:\\overmind\\temp\\tempchart.png' )
		r = chart.add_run ( ' \n Fig.' + str ( (j*12) + 5 ) + ' FFT envelope graph point ' + str (
			maxPoint ) + ' on ' + thname + '(SB turn) gear' )
		# r=chart.add_run(str(thrusters[(((j)*3))]))
		r.font.name = 'Calibri'
		r.font.size = Pt ( 9 )
		r.font.italic = True
		chart.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

		chart = document.add_paragraph ()
		r = chart.add_run ()
		maxPoint = createchart ( host , username , password , motors [ (((j)*3) + 1) ] , rn_ , 'RMS' )
		r.add_picture ( 'C:\\overmind\\temp\\tempchart.png' , width=Inches ( 7.0 ) )
		os.remove ( 'C:\\overmind\\temp\\tempchart.png' )
		r = chart.add_run ( ' \n Fig.' + str ( (j*12) + 6 ) + ' FFT velocity graph point ' + str (
			maxPoint ) + ' on ' + thname + '(PS turn) el. motor' )
		# r=chart.add_run(str(motors[(((j)*3)+1)]))
		r.font.name = 'Calibri'
		r.font.size = Pt ( 9 )
		r.font.italic = True
		chart.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

		chart = document.add_paragraph ()
		r = chart.add_run ()
		maxPoint = createchart ( host , username , password , motors [ (((j)*3) + 1) ] , rn_ , 'envelope P-K' )
		r.add_picture ( 'C:\\overmind\\temp\\tempchart.png' , width=Inches ( 7.0 ) )
		os.remove ( 'C:\\overmind\\temp\\tempchart.png' )
		r = chart.add_run ( ' \n Fig.' + str ( (j*12) + 7 ) + ' FFT envelope graph point ' + str (
			maxPoint ) + ' on ' + thname + '(PS turn) el. motor' )
		# r=chart.add_run(str(motors[(((j)*3)+1)]))
		r.font.name = 'Calibri'
		r.font.size = Pt ( 9 )
		r.font.italic = True
		chart.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

		chart = document.add_paragraph ()
		r = chart.add_run ()
		maxPoint = createchart ( host , username , password , thrusters [ (((j)*3) + 1) ] , rn_ , 'RMS' )
		r.add_picture ( 'C:\\overmind\\temp\\tempchart.png' , width=Inches ( 7.0 ) )
		os.remove ( 'C:\\overmind\\temp\\tempchart.png' )
		r = chart.add_run ( ' \n Fig.' + str ( (j*12) + 8 ) + ' FFT velocity graph point ' + str (
			maxPoint ) + ' on ' + thname + '(PS turn) gear' )
		# r=chart.add_run(str(thrusters[(((j)*3)+1)]))
		r.font.name = 'Calibri'
		r.font.size = Pt ( 9 )
		r.font.italic = True
		chart.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

		chart = document.add_paragraph ()
		r = chart.add_run ()
		maxPoint = createchart ( host , username , password , thrusters [ (((j)*3) + 1) ] , rn_ , 'envelope P-K' )
		r.add_picture ( 'C:\\overmind\\temp\\tempchart.png' , width=Inches ( 7.0 ) )
		os.remove ( 'C:\\overmind\\temp\\tempchart.png' )
		r = chart.add_run ( ' \n Fig.' + str ( (j*12) + 9 ) + ' FFT envelope graph point ' + str (
			maxPoint ) + ' on ' + thname + '(PS turn) gear' )
		# r=chart.add_run(str(thrusters[(((j)*3)+1)]))
		r.font.name = 'Calibri'
		r.font.size = Pt ( 9 )
		r.font.italic = True
		chart.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

		chart = document.add_paragraph ()
		r = chart.add_run ()
		maxPoint = createchart ( host , username , password , motors [ (((j)*3) + 2) ] , rn_ , 'RMS' )
		r.add_picture ( 'C:\\overmind\\temp\\tempchart.png' , width=Inches ( 7.0 ) )
		os.remove ( 'C:\\overmind\\temp\\tempchart.png' )
		r = chart.add_run ( ' \n Fig.' + str ( (j*12) + 10 ) + ' FFT velocity graph point ' + str (
			maxPoint ) + ' on ' + thname + '(idle) el. motor' )
		# r=chart.add_run(str(motors[(((j)*3)+2)]))
		r.font.name = 'Calibri'
		r.font.size = Pt ( 9 )
		r.font.italic = True
		chart.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

		chart = document.add_paragraph ()
		r = chart.add_run ()
		maxPoint = createchart ( host , username , password , motors [ (((j)*3) + 2) ] , rn_ , 'envelope P-K' )
		r.add_picture ( 'C:\\overmind\\temp\\tempchart.png' , width=Inches ( 7.0 ) )
		os.remove ( 'C:\\overmind\\temp\\tempchart.png' )
		r = chart.add_run ( ' \n Fig.' + str ( (j*12) + 11 ) + ' FFT envelope graph point ' + str (
			maxPoint ) + ' on ' + thname + '(idle) el. motor' )
		# r=chart.add_run(str(motors[(((j)*3)+2)]))
		r.font.name = 'Calibri'
		r.font.size = Pt ( 9 )
		r.font.italic = True
		chart.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

		chart = document.add_paragraph ()
		r = chart.add_run ()
		maxPoint = createchart ( host , username , password , thrusters [ (((j)*3) + 2) ] , rn_ , 'RMS' )
		r.add_picture ( 'C:\\overmind\\temp\\tempchart.png' , width=Inches ( 7.0 ) )
		os.remove ( 'C:\\overmind\\temp\\tempchart.png' )
		r = chart.add_run ( ' \n Fig.' + str ( (j*12) + 12 ) + ' FFT velocity graph point ' + str (
			maxPoint ) + ' on ' + thname + '(idle) gear' )
		# r=chart.add_run(str(thrusters[(((j)*3)+2)]))
		r.font.name = 'Calibri'
		r.font.size = Pt ( 9 )
		r.font.italic = True
		chart.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

		chart = document.add_paragraph ()
		r = chart.add_run ()
		maxPoint = createchart ( host , username , password , thrusters [ (((j)*3) + 2) ] , rn_ , 'envelope P-K' )
		r.add_picture ( 'C:\\overmind\\temp\\tempchart.png' , width=Inches ( 7.0 ) )
		os.remove ( 'C:\\overmind\\temp\\tempchart.png' )
		r = chart.add_run ( ' \n Fig.' + str ( (j*12) + 13 ) + ' FFT envelope graph point ' + str (
			maxPoint ) + ' on ' + thname + '(idle) gear' )
		# r=chart.add_run(str([(((j)*3)+2)]))
		r.font.name = 'Calibri'
		r.font.size = Pt ( 9 )
		r.font.italic = True
		chart.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

	document.add_page_break ()
	l3 = document.add_paragraph ( '3. Summary' )
	l3.style = document.styles [ 'listlvl1' ]
	l3.alignment = WD_ALIGN_PARAGRAPH.LEFT
	l3.paragraph_format.left_indent = Inches ( 0.25 )
	document.add_paragraph ()
	l31 = document.add_paragraph ( '3.1.	Conclusion // SPRAWDZIC / UZUPEŁNIC PRZED WYSŁANIEM // ' )
	l31.style = document.styles [ 'listlvl2' ]
	l31.alignment = WD_ALIGN_PARAGRAPH.LEFT
	l31.paragraph_format.left_indent = Inches ( 0.25 )
	document.add_paragraph ()
	t13 = document.add_paragraph (
		'The vibration of 6 thruster (electric motors and gear parts) are on very low level. Analysis was performed only on vibration values. Oil analysis is not available.' )
	t13.style = document.styles [ 'text' ]
	t13.alignment = WD_ALIGN_PARAGRAPH.LEFT
	document.add_paragraph ()
	l32 = document.add_paragraph ( '3.2.	Recommendation // SPRAWDZIC / UZUPEŁNIC PRZED WYSŁANIEM // ' )
	l32.style = document.styles [ 'listlvl2' ]
	l32.alignment = WD_ALIGN_PARAGRAPH.LEFT
	l32.paragraph_format.left_indent = Inches ( 0.25 )
	document.add_paragraph ()
	t14 = document.add_paragraph (
		'Next measurements should be done within one month period or earlier if necessary.' )
	t14.style = document.styles [ 'text' ]
	t14.alignment = WD_ALIGN_PARAGRAPH.LEFT
	document.add_paragraph ()
	t15 = document.add_paragraph (
		'This report is prepared in good faith based on measurement diagnostic done on available object and documentation submitted.' )
	t15.style = document.styles [ 'text' ]
	t15.alignment = WD_ALIGN_PARAGRAPH.LEFT

	document.add_paragraph ()
	t15 = document.add_paragraph (
		'Report prepared by:                                                                                                                           Approved by:' )
	t15.style = document.styles [ 'texthead' ]
	t15.alignment = WD_ALIGN_PARAGRAPH.LEFT

	fullname = getname ( host , username , password )
	t16 = document.add_paragraph ( str ( fullname [ 0 ] [ 0 ] ) )
	t16.style = document.styles [ 'text' ]
	t16.alignment = WD_ALIGN_PARAGRAPH.LEFT
	section=document.sections[0]
	footer=section.footer.paragraphs[0]
	footer.add_run('Diagnostic report for platform '+platname+' '+str(rn_)).font.name='Times New Roman'
	#footer.runs[0].font.name='Times New Roman'
	footer.runs[0].font.size=Pt(12)

	mydate = datetime.datetime.now ()
	repdate = str ( mydate.year ) + '_' + str ( mydate.month ) + '_' + str ( mydate.day )

	ini = getini ( host , username , password )

	if id_ == '53':
		document.save ( 'C:\overmind\Reports\Diagnostic report Atlantis Platform ' + str ( rn_ ) + '_' + str (
			ini [ 0 ] [ 0 ] ) + '_' + str ( repdate ) + '.docx' )
	elif id_ == '54':
		document.save ( 'C:\overmind\Reports\Diagnostic report Iolar Platform ' + str ( rn_ ) + '_' + str (
			ini [ 0 ] [ 0 ] ) + '_' + str ( repdate ) + '.docx' )
	elif id_ == '55':
		document.save ( 'C:\overmind\Reports\Diagnostic report Neptuno Platform ' + str ( rn_ ) + '_' + str (
			ini [ 0 ] [ 0 ] ) + '_' + str ( repdate ) + '.docx' )


