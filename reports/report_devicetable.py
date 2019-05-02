import psycopg2
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt ,Cm


def set_col_width ( table ) :
    widths = (Cm ( 5 ) ,Cm ( 7 ))
    for row in table.rows :
        for idx ,width in enumerate ( widths ) :
            row.cells[ idx ].width = width


def devicetable ( document ) :
    rn_ = '1987-2019'
    querry = 'select device, serialno,manudate from equipment where lp = (select equipment from shipsdata where shipid = 84 )'
    conn = psycopg2.connect ( user='postgres' ,dbname='postgres' ,host='localhost' ,password='info' ,port='5432' )
    cur = conn.cursor ()
    cur.execute ( querry )
    result = cur.fetchall ()
    deviceNAME = str ( result[ 0 ][ 0 ] )
    devSN = str ( result[ 0 ][ 1 ] )
    H = document.add_paragraph ( 'Measurement equipment' )
    H1 = document.add_paragraph ( deviceNAME + devSN )
    H.runs[ 0 ].font.name = 'Times New Roman'
    H.runs[ 0 ].font.size = Pt ( 12 )
    H.runs[ 0 ].bold = True

    devtab = document.add_table ( rows=6 ,cols=2 )
    devtab.style = 'Table Grid'
    devtab.cell ( 0 ,0 ).merge ( devtab.cell ( 0 ,1 ) )
    devtab.cell ( 0 ,0 ).text = 'Technical data'
    devtab.cell ( 0 ,0 ).paragraphs[ 0 ].alignment = WD_ALIGN_PARAGRAPH.CENTER
    devtab.cell ( 1 ,0 ).text = 'Maker:'
    devtab.cell ( 1 ,1 ).text = 'Info Marine'
    devtab.cell ( 1 ,1 ).paragraphs[ 0 ].runs[ 0 ].font.name = 'Arial'
    devtab.cell ( 1 ,1 ).paragraphs[ 0 ].runs[ 0 ].font.size = Pt ( 12 )
    devtab.cell ( 2 ,0 ).text = 'Serial number:'
    devtab.cell ( 3 ,0 ).text = 'Date of manufacture:'
    devtab.cell ( 4 ,0 ).text = 'Measuring range:'
    devtab.cell ( 5 ,0 ).text = 'Indication error:'
    devtab.cell ( 4 ,1 ).text = '2Hz-30kHz / RPM = 60-20000'
    devtab.cell ( 5 ,1 ).text = '±0,5%'
    set_col_width ( devtab )

    for i in range ( 6 ) :
        devtab.cell ( i ,0 ).paragraphs[ 0 ].runs[ 0 ].font.name = 'Arial'

        devtab.cell ( i ,0 ).paragraphs[ 0 ].runs[ 0 ].font.size = Pt ( 12 )

    print ( result[ 0 ][ 0 ] )
    print ( result[ 0 ][ 1 ] )
    print ( result[ 0 ][ 2 ] )
    devtab.cell ( 1 ,1 ).text = str ( result[ 0 ][ 0 ] )
    devtab.cell ( 2 ,1 ).text = str ( result[ 0 ][ 1 ] )
    devtab.cell ( 3 ,1 ).text = str ( result[ 0 ][ 2 ] )
    devtab.cell ( 1 ,1 ).paragraphs[ 0 ].runs[ 0 ].font.size = Pt ( 12 )
    devtab.cell ( 1 ,1 ).paragraphs[ 0 ].runs[ 0 ].font.name = 'Arial'

    devtab.cell ( 2 ,1 ).paragraphs[ 0 ].runs[ 0 ].font.size = Pt ( 12 )
    devtab.cell ( 2 ,1 ).paragraphs[ 0 ].runs[ 0 ].font.name = 'Arial'

    devtab.cell ( 3 ,1 ).paragraphs[ 0 ].runs[ 0 ].font.size = Pt ( 12 )
    devtab.cell ( 3 ,1 ).paragraphs[ 0 ].runs[ 0 ].font.name = 'Arial'

    devtab.cell ( 4 ,1 ).paragraphs[ 0 ].runs[ 0 ].font.size = Pt ( 12 )
    devtab.cell ( 4 ,1 ).paragraphs[ 0 ].runs[ 0 ].font.name = 'Arial'

    devtab.cell ( 5 ,1 ).paragraphs[ 0 ].runs[ 0 ].font.size = Pt ( 12 )
    devtab.cell ( 5 ,1 ).paragraphs[ 0 ].runs[ 0 ].font.name = 'Arial'

    devsummary = document.add_paragraph ()
    devsummary.add_run ().add_break ()
    devrun = devsummary.add_run ( 'MarVib is calibrated, certificate for verification - if requied. ' )
    devrun.italic = True
    document.add_paragraph ()
