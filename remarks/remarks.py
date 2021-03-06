import psycopg2
import time
import datetime
import tkinter as tk
import tkinter.ttk as ttk
from tkinter import *
from tkinter import messagebox
import csv
from tkinter.filedialog import askopenfilename
from tkinter import Tk
from tkcalendar import Calendar, DateEntry
import matplotlib.pyplot as plt
import numpy as num
import tkinter as tkk
from tkinter import filedialog
import docx
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Pt, Cm
from docx import Document
import re
class Scrollable(ttk.Frame):
    """
       Make a frame scrollable with scrollbar on the right.
       After adding or removing widgets to the scrollable frame,
       call the update() method to refresh the scrollable area.
    """

    def __init__(self, frame, width=16):

        scrollbar = tk.Scrollbar(frame, width=width)
        scrollbar.pack(side=tk.RIGHT, fill=BOTH, expand=False)

        self.canvas = tk.Canvas(frame, yscrollcommand=scrollbar.set, height = 700)
        self.canvas.pack(side=tk.LEFT, fill=BOTH, expand=True)

        scrollbar.config(command=self.canvas.yview)

        self.canvas.bind('<Configure>', self.__fill_canvas)

        # base class initialization
        tk.Frame.__init__(self, frame)

        # assign this obj (the inner frame) to the windows item of the canvas
        self.windows_item = self.canvas.create_window(0,0, window=self, anchor=tk.NW)


    def __fill_canvas(self, event):
        "Enlarge the windows item to the canvas width"

        canvas_width = event.width
        self.canvas.itemconfig(self.windows_item, width = canvas_width)

    def update(self):
        "Update the canvas and the scrollregion"

        self.update_idletasks()
        self.canvas.config(scrollregion=self.canvas.bbox(self.windows_item))


def q_run(connD, querry):
    username = connD[0]
    password = connD[1]
    host = connD[2]
    kport = "5432"
    kdb = "postgres"
    # cs = ' host="localhost",database="postgres", user= "postgres" , password="info" '
    cs = "dbname=%s user=%s password=%s host=%s port=%s" % (kdb, username, password, host, kport)
    conn = None
    conn = psycopg2.connect(str(cs))
    cur = conn.cursor()
    cur.execute(querry)
    try:
        result = cur.fetchall()
        return result
    except:
        pass
    conn.commit()
    cur.close()

class LogApplication:

    def __init__(self):
        self.root = Tk()
        self.root.title("Log In")
        self.title = tk.Label(self.root, text="Info Datasheet")  # TITLE
        self.title.grid(row=0, column=2)
        self.user_entry_label = tk.Label(self.root, text="Username: ")  # USERNAME LABEL
        self.user_entry_label.grid(row=1, column=1)
        self.user_entry = tk.Entry(self.root, text="Username: ")  # USERNAME ENTRY BOX
        self.user_entry.grid(row=1, column=2)
        self.pass_entry_label = tk.Label(self.root, text="Password: ")  # PASSWORD LABEL
        self.pass_entry_label.grid(row=2, column=1)
        self.pass_entry = tk.Entry(self.root, show="*")  # PASSWORD ENTRY BOX
        self.pass_entry.grid(row=2, column=2)
        try:
            with open('C:\overmind\\temp\log.csv') as csvfile:
                openfile = csv.reader(csvfile, delimiter=' ')
                p = -1
                for lines in openfile:
                    p += 1
                    if p == 0:
                        self.user_entry.insert(0, str(lines[0]))
                    if p == 1:
                        self.pass_entry.insert(0, str(lines[0]))
        except:
            pass
        self.var = IntVar()
        self.checksave = tk.Checkbutton(self.root, text="Remember", variable=self.var)
        self.checksave.grid(row=3, column=2)
        self.sign_in_butt = Button(self.root, text="Sign In", command=lambda ue=self.user_entry, pe=self.pass_entry: self.logging_in(ue, pe))
        self.sign_in_butt.grid(row=5, column=2)
        self.root.mainloop()

    def logging_in(self,user_entry, pass_entry):
        user_get = user_entry.get()  # Retrieve Username
        pass_get = pass_entry.get()  # Retrieve Password
        if bool(self.var.get()) == True:
            # config = Path('C:\overmind\\temp\log.csv')
            with open('C:\overmind\\temp\log.csv', 'w+' ,newline='') as csvfile:
                filewriter = csv.writer(csvfile, quoting=csv.QUOTE_ALL)
                filewriter.writerow([user_get])
                filewriter.writerow([pass_get])
        connD = [user_get, pass_get, '192.168.10.243']


        querry = "SELECT current_user"
        usercheck = ''

        usercheck = q_run(connD, querry)  # PYINSTALLER ma problemy gdzies tu
        if usercheck != '':
            self.root.destroy()
            remarks(connD)

remlist = list()
remlistIN = list()

fedblist = list()
fedblistIN = list()

def createfeedback(rn_,connD):
    def getdatestr():
        querry = """select min(ml.date), max(ml.date)
    	from remarks as rem
    	left join devices as dev on rem.id = dev.id
    	left join (select id,raport_number,date,max(value)
    			   from measurements_low
    			   where type = 'RMS'
    			   group by id,raport_number,date) as ml on rem.id = ml.id and rem.raport_number = ml.raport_number
    	where rem.raport_number = '{}' and rem.sended = True""".format(rn_)
        datemin, datemax = list(q_run(connD, querry))[0]
        if datemin == datemax:
            datestr = str(datemin)
        else:
            datestr = '{} - {}'.format(datemin,datemax)
        return datestr


    document = Document('C:\\overmind\\Data\\base.docx')
    querry = "select main.name, shd.imo, main2.name from main right join shipsdata as shd on main.id = shd.shipid right join main as main2 on main.parent = main2.id where main.id = (select shipid from reports where raport_number = '" + str(
        rn_) + "')"
    result = q_run(connD, querry)

    if len(result) == 0 :
        querry = "select main.name, shd.imo, main2.name from main right join shipsdata as shd on main.id = shd.shipid right join main as main2 on main.parent = main2.id where main.id = (select shipid from harmonogram where report_number = '" + str(
            rn_) + "')"
        result = q_run(connD, querry)


    def headtable(document):
        username = connD[0]
        password = connD[1]
        host = connD[2]


        shipstr = result[0][0]

        headtable = document.add_table(rows=2, cols=3)  # trzeba usunąć enter przed
        headtable.style = 'Table Grid'
        for row in headtable.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        headtable.cell(0, 0).merge(headtable.cell(0, 1).merge(headtable.cell(0, 2)))
        ht = headtable.cell(0, 0).paragraphs[0]
        ht.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r0 = ht.add_run()
        r0.text = 'Work done after vibration diagnostic report '
        r0.font.name = 'Calibri'
        r0.font.size = Pt(12)
        r1 = ht.add_run(rn_)
        r1.bold = True
        r1.font.name = 'Calibri'

        ht = headtable.cell(1, 0).paragraphs[0]
        ht.alignment = WD_ALIGN_PARAGRAPH.LEFT

        r0 = ht.add_run('Project: ')
        r0.font.name = 'Calibri'
        r1 = ht.add_run(shipstr)
        r1.bold = True
        r1.font.name = 'Calibri'

        imostr = str(result[0][1])
        ht = headtable.cell(1, 0).add_paragraph()
        r0 = ht.add_run('IMO no: ')
        r0.font.name = 'Calibri'
        r1 = ht.add_run(imostr)
        r1.font.name = 'Calibri'
        r1.bold = True

        ownerstr = str(result[0][2])
        ht = headtable.cell(1, 0).add_paragraph()
        r0 = ht.add_run('Ordered by: ')
        r0.font.name = 'Calibri'
        r1 = ht.add_run(ownerstr)
        r1.bold = True
        r1.font.name = 'Calibri'
        ht = headtable.cell(1, 1).paragraphs[0]
        r0 = ht.add_run('Date of measurement: ')
        r0.font.name = 'Calibri'

        ht.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ht.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
        ht = headtable.cell(1, 1).add_paragraph()

        ht.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ht.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
        r0 = ht.add_run(getdatestr())
        r0.bold = True
        r0.font.name = 'Calibri'
        ht = headtable.cell(1, 2).paragraphs[0]
        ht.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r0 = ht.add_run('Place of measurement:')
        r0.font.name = 'Calibri'

        ht = headtable.cell(1, 2).add_paragraph()
        ht.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r0 = ht.add_run('During normal operation')
        r0.bold = True
        r0.font.name = 'Calibri'
        section = document.sections[0]
        document.add_paragraph()
    def feedbacktable():

        def countLimit(standard, value):
            limSrt = 'NOPE'
            for limNo in limits:
                if str(limNo[0]) == str(standard):
                    if value <= float(limNo[3]):  # IF LIM1
                        limSrt = str(limNo[2])
                        break
                    else:
                        if value <= float(limNo[5]):  # IF LIM2
                            limSrt = str(limNo[4])
                            break
                        else:
                            if value <= float(limNo[7]):  # IF LIM3
                                limSrt = str(limNo[6])
                                break
                            else:
                                limSrt = str(limNo[8])
                                break
            return limSrt

        querry = """select dev.name, rem.remark,ml.max,ml.date,dev.norm
	from remarks as rem
	left join devices as dev on rem.id = dev.id
	left join (select id,raport_number,date,max(value)
			   from measurements_low
			   where type = 'RMS'
			   group by id,raport_number,date) as ml on rem.id = ml.id and rem.raport_number = ml.raport_number
	where rem.raport_number = '""" + str(rn_) + """' and rem.sended = True"""

        remarklist = q_run(connD,querry)

        querry = """select standard,
								limit_1_value,limit_1_name,
								limit_2_value,limit_2_name,
								limit_3_value,limit_3_name,
								limit_4_value,limit_4_name,
							envflag
						from standards"""
        limits = q_run(connD, querry)


        rws = (len(remarklist)) + 1

        feedbacktable = document.add_table(rows=rws, cols=7)  # trzeba usunąć enter przed
        feedbacktable.style = 'Table Grid'

        for row in feedbacktable.rows:
            widths = (Cm(2.5 ) ,Cm(1.5) ,Cm(2) ,Cm(5) ,Cm( 2),Cm(2.5) ,Cm(2.5))
            for idx, width in enumerate(widths):
                row.cells[idx].width = width
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        ht = feedbacktable.cell(0, 0).paragraphs[0]
        ht.alignment = WD_ALIGN_PARAGRAPH.CENTER

        r0 = ht.add_run('Machine name')
        r0.font.name = 'Calibri'
        r0.font.size = Pt(8)
        r0.font.bold = True

        ht = feedbacktable.cell(0, 1).paragraphs[0]
        ht.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = ht.add_run('ISO' + chr(10) + 'Standard')
        r1.font.name = 'Calibri'
        r1.font.size = Pt(8)
        r1.font.bold = True

        ht = feedbacktable.cell(0, 2).paragraphs[0]
        ht.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = ht.add_run('Measurement date')
        r2.font.name = 'Calibri'
        r2.font.size = Pt(8)
        r2.font.bold = True

        ht = feedbacktable.cell(0, 3).paragraphs[0]
        ht.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r3 = ht.add_run('Remarks and suggestions')
        r3.font.name = 'Calibri'
        r3.font.size = Pt(8)
        r3.font.bold = True

        ht = feedbacktable.cell(0, 4).paragraphs[0]
        ht.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r4 = ht.add_run('Date of' + chr(10) + 'corrective'  + chr(10) + 'action done')
        r4.font.name = 'Calibri'
        r4.font.size = Pt(8)
        r4.font.bold = True

        ht = feedbacktable.cell(0, 5).paragraphs[0]
        ht.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r5 = ht.add_run('Work done on board' + chr(10) + 'vessel' + chr(10) + 'after vibration diagnostic' + chr(10) +'report')
        r5.font.name = 'Calibri'
        r5.font.size = Pt(8)
        r5.font.bold = True

        ht = feedbacktable.cell(0, 6).paragraphs[0]
        ht.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r6 = ht.add_run('Engineers' + chr(10) + 'observations on wear' + chr(10) + 'items of equipment')
        r6.font.name = 'Calibri'
        r6.font.size = Pt(8)
        r6.font.bold = True

        c = -1
        r = 0

        for line in remarklist:
            r +=1
            c =-1
            line = list(line)
            for element in line:
                c += 1

                if c == 0: # machine name
                    P = feedbacktable.cell(r, c).paragraphs[0]
                    R = P.add_run(element)


                elif c == 1:#remarks n sugesstions
                    P = feedbacktable.cell(r, 3).paragraphs[0]
                    P.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    R = P.add_run(element)

                elif c == 2:#ISO
                    P = feedbacktable.cell(r, 1).paragraphs[0]
                    P.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for limitline in limits:
                        if str(line[4]) == str(limitline[0]):
                            limitstr = (countLimit(line[4], line[2]))
                            R = P.add_run(limitstr)
                    if limitstr == 'Cl. A':
                        R.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
                    if limitstr == 'Cl. B':
                        R.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
                    if limitstr == 'Cl. C':
                        R.font.highlight_color = WD_COLOR_INDEX.YELLOW
                    if limitstr == 'Cl. D':
                        R.font.highlight_color = WD_COLOR_INDEX.RED

                elif c == 3: #meas date

                    P = feedbacktable.cell(r, 2).paragraphs[0]
                    P.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    R = P.add_run(str(element))
                else:
                    pass

                R.font.name = 'Calibri'
                R.font.size = Pt(8)
                R.alignment = WD_ALIGN_PARAGRAPH.CENTER


    headtable(document)
    feedbacktable()
    document.add_paragraph()
    document.add_paragraph()
    E = document.add_paragraph()
    R = E.add_run('……………………………………………')
    R.bold = False
    R.font.name = 'Times New Roman'
    R.font.size = Pt(10)
    E.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    E = document.add_paragraph()
    R = E.add_run('(CE in Charge)                 .')
    R.bold = False
    R.font.name = 'Times New Roman'
    R.font.size = Pt(10)
    E.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    f = filedialog.askdirectory()
    document.save( f +'/'+ str(result[0][0]) + ' - feedback for vibration report ' + str(rn_) + '.docx' )

def grabremarks():
    remlistIN.clear()
    class remark:
        def __init__(self):
            self.name = ''
            self.remark = ''
            self.reminder = False
            self.nextdate = ''

    a = filedialog.askopenfilename()
    document = docx.Document(docx=a)
    t = -1
    r = -1
    c = -1
    tableflag = False
    rt = list()
    namecol = list()
    remcol = list()
    for table in document.tables:
        t += 1
        r = -1
        for row in table.rows:
            rem = remark()
            r += 1
            c = -1
            # if tableflag == True:
            for cell in row.cells:
                c += 1
                for paragraph in cell.paragraphs:
                    coord = [t, r, c]
                    if str(paragraph.text) == 'Machine name':
                        tableflag = True
                        namecol.append(c)
                    if str(paragraph.text) == 'Remarks and suggestions':
                        tableflag = True
                        remcol.append(c)
                    if tableflag == True:
                        if t not in rt:
                            rt.append(t)
                    if str(paragraph.text) == 'Technical data':
                        tableflag = False
                    if tableflag == True:
                        if t in rt:
                            if c in namecol:
                                if str(paragraph.text).strip() != '':
                                    rem.name = str(paragraph.text)
                            if c in remcol:
                                if str(paragraph.text).strip() != '' and str(paragraph.text).strip() not in str(
                                        rem.remark).strip():
                                    rem.remark += str(paragraph.text) + chr(10)
            if (str(rem.name)).strip() != '' and (str(rem.remark)).strip() != '':
                if (str(rem.name)).strip() != (str(rem.remark)).strip():
                    if 'up to week' in rem.remark:
                        rem.reminder = True
                        s = str(rem.remark)
                        result = re.search('done up to week (.*).', s)
                        week = int(''.join(filter(str.isdigit, result.group(1))))
                        year = datetime.datetime.now().year
                        atime = time.asctime(time.strptime('{} {} 1'.format(year, week), '%Y %W %w'))
                        atime = datetime.datetime.strptime(atime, '%a %b %d %X %Y').date() - datetime.timedelta(days=1)
                        rem.nextdate = atime
                    remlistIN.append(rem)


    for line2 in remlist:
        for line in remlistIN:

            if (str(line2.name)).strip() == (str(line.name)).strip():
                remarkstr = str(line.remark)
                line2.textfield.delete('1.0', END)
                line2.textfield.insert(INSERT,remarkstr)
                if line.reminder == True:
                    line2.var2.set(1)
                    line2.dateLabel.delete('1.0', END)
                    line2.dateLabel.insert(INSERT, line.nextdate)

        #break

def grabfeedbacks():
    class feedback:
        def __init__(self):
            self.name = ''
            self.feedback = ''
            self.documentdate = ''

    a = filedialog.askopenfilename()
    document = docx.Document(docx=a)
    t = -1
    tableflag = False
    rt = list()
    namecol = list()
    remcol = list()

    fdbtable = document.tables[1]

    r = -1
    fedblistIN.clear()
    for row in fdbtable.rows:
        r += 1
        if r > 0:
            c = -1
            for cell in row.cells:
                c += 1
                paragraph = cell.paragraphs[0]
                if c == 0 :
                    fed = feedback()
                    fed.name = (str(paragraph.text))
                if c == 4:
                    fed.documentdate = (str(paragraph.text))
                if c == 5 :
                    workdone = (str(paragraph.text) )
                if c == 6 :
                    observations = (str(paragraph.text) )
                    fed.feedback = ('Work done: ' + str(workdone) + chr(10) + 'Observation: '+str(observations))
                    fedblistIN.append(fed)



    for line2 in remlist:
        for line in fedblistIN:
            if (str(line2.name)).strip() == (str(line.name)).strip():
                fdbstr = str(line.feedback)
                line2.textfield2.delete('1.0', END)
                line2.textfield2.insert(INSERT, fdbstr)

                line2.FDBdateLabel.delete('1.0', END)
                line2.FDBdateLabel.insert(INSERT, str(line.documentdate))

def remarks(connD):

    querry = "select name,id from main where parent = 1 order by name"
    resultr = q_run(connD, querry)

    def datepick(ob):
        if str(ob.var2.get()) == '0':
            ob.dateLabel.configure(text='')
        if str(ob.var2.get()) == '1':
            def print_sel():
                ob.requestdate =  str(cal.selection_get())
                ob.dateLabel.configure(text=ob.requestdate)
                root.destroy()
            root = Tk()
            root.title("Calendar")
            #s = ttk.Style(root)
            #s.theme_use('clam')
            top = tk.Toplevel(root)

            cal = Calendar(top,
                           font="Arial 14", selectmode='day',
                           cursor="hand1", year=2019, month=4, day=15)
            cal.pack(fill="both", expand=True)
            ttk.Button(top, text="ok", command=print_sel).pack()
            root.mainloop()

    def getships(evt):
        w = evt.widget
        index = int(w.curselection()[0])
        shipname = w.get(index)
        makeships(shipname)

    def makeships(shipname):
        querry = "select name from main where parent =(select id from main where name = '" + str(shipname) +"' limit 1) order by name"
        ships = q_run(connD,querry)
        Shiplistbox.delete(0,'end')

        for line in ships:
            Shiplistbox.insert(END, line[0])

    def getreports(evt):
        w = evt.widget
        index = int(w.curselection()[0])
        shipname = w.get(index)
        makereports(shipname)

    def makereports(shipname):
        querry = "select raport_number from measurements_low where parent = (select id from main where name = '" + str(shipname) +"' limit 1) group by raport_number order by raport_number DESC"
        reports = q_run(connD,querry)
        querry = "(select id from main where name = '" + str(shipname) +"' limit 1)"
        parent = str(q_run(connD,querry)[0][0])
        Reportlistbox.delete(0,'end')

        for line in reports:
            Reportlistbox.insert(END, line[0])

    def getdevices(evt):
        w = evt.widget
        index = int(w.curselection()[0])
        report = w.get(index)
        putdevices(report,None)


    class frame_reminder:
        def __init__(self, measCframe,dev,rn,remtable):
            self.rn = rn
            self.id = dev[1]
            self.name = dev[0]
            self.parent =dev[2]
            self.namedateframe = tk.Frame(measCframe)
            self.namedateframe.pack(side = LEFT)
            self.nameLabel = tk.Text(self.namedateframe,height =1, width=40, borderwidth = 0)
            self.nameLabel.insert(1.0,self.name)
            self.nameLabel.pack(side = TOP,anchor = W)
            self.nameLabel.configure(state = 'disabled')
            self.devdateLabel = tk.Label(self.namedateframe,text = dev[3])
            self.devdateLabel.pack(side = TOP,anchor = W)
            self.remarkframe = tk.Frame(measCframe)


            self.textfield = tk.Text(self.remarkframe , wrap=WORD,width=55, height=5)
            self.remindertextfield = tk.Text(self.remarkframe,wrap=WORD, width=55, height=2)
            self.dateLabel = tk.Text(self.remarkframe , width=15, height=1)

            self.feedbackframe = tk.Frame(measCframe)
            self.textfield2 = tk.Text(self.feedbackframe,wrap=WORD, width=55, height=7)
            self.FDBdateLabel = tk.Text(self.feedbackframe, width=15, height=1)

            try:
                remarktext = self.getrem(str(self.id), str(self.rn),remtable)

                remark__ = remarktext[0]
                sended__ = remarktext[1]
                request_date__ = remarktext[2]
                feedback__ = remarktext[3]
                fdbdocdate__ = remarktext[4]
                remcom__ = remarktext[5]

                if str(remark__) == 'None': remark__ = ''
                self.textfield.insert(INSERT, str(remark__))
                if str(feedback__) == 'None': feedback__ = ''
                self.textfield2.insert(INSERT, str(feedback__))
                if str(request_date__) == 'None': request_date__ = ''
                self.dateLabel.insert(INSERT, str(request_date__))
                if str(fdbdocdate__) == 'None': fdbdocdate__ = ''
                self.FDBdateLabel.insert(INSERT, str(fdbdocdate__))
                if str(remcom__) == 'None': remcom__ = ''
                self.remindertextfield.insert(INSERT, str(remcom__))


            except:
                remark__ = ''
                sended__ = 'None'
                request_date__ = 'None'

            self.remarkframe.pack()
            #self.textfield.pack(side=LEFT)

            self.var = tk.IntVar(value=0)
            self.var2 = tk.IntVar(value=0)
            self.check = ttk.Checkbutton(self.remarkframe, text='Sent', variable=self.var)
            if str(sended__) == 'True':
                self.var.set(1)
            if str(request_date__) == 'None' or  str(request_date__) == '':
                self.var2.set(0)
            else:
                self.var2.set(1)

            self.check2 = ttk.Checkbutton(self.remarkframe, text='Reminder', variable=self.var2)



            self.requestdate = ''
            self.datevar = tk.StringVarVarvalue="1989-06-08"



            measCframe.pack(side=TOP, fill=BOTH, expand=True)
            self.remarkframe.pack(side=LEFT)





            self.check.grid(row=1, column=0, sticky="e")
            self.textfield.grid(row=1, column=1, columnspan=2)


            self.check2.grid(row=2, column=0, sticky="e")
            self.dateLabel.grid(row=3, column=0, sticky="e")
            self.remindertextfield.grid(row=2, column=1,columnspan=2,rowspan=2)


            tk.Label(measCframe,text= 'Feedback: ').pack(side=LEFT)
            self.feedbackframe.pack(side=LEFT)
            self.textfield2.grid(row=2, column = 0)
            self.FDBdateLabel.grid(row=1, column = 0, sticky="w")
        def getrem(self,id,rn,remtable):
            for line in remtable:

                if str(line[0]) == str(id) and str(line[1]) == str(rn):
                    ans = [str(line[2]),str(line[3]),str(line[4]),str(line[5]),str(line[6]),str(line[7])]
                    return ans

                    break

            pass

    def putdevices(report,namefilter):
        def upload():

            querry = """select rem.id, rem.sended, remi.request_date
			from remarks as rem
			left join reminder as remi on rem.id = remi.id and rem.raport_number = remi.raport_number
			where rem.raport_number ='""" + str(report) + "'group by rem.id, rem.sended, remi.request_date"
            remidlist = list(q_run(connD, querry))

            querry = """select id,raport_number from feedbacks where raport_number = '"""  + str(report) + """' group by id,raport_number"""
            fedidlist = list(q_run(connD, querry))
            c=0

            idlist = list(row[0] for row in remidlist)


            iddfblist = list(row[0] for row in fedidlist)



            for line in remlist:

                if str(line.textfield.get("1.0", END)).strip() != '':
                    line.requestdate = (line.dateLabel.get("1.0", END))
                    querry = "select date from measurements_low where id = " + str(
                        line.id) + " and raport_number = '" + str(line.rn) + "' limit 1"
                    try:
                        measdate = str(q_run(connD, querry)[0][0])
                    except:
                        measdate ='1989-08-06'

                    if line.var.get() == 1 : sendflag = 'True'
                    else: sendflag = 'False'
                    if line.var2.get() == 1: remflag = 'True'
                    else: remflag = 'False'



                    if line.id in idlist:
                        querry = "UPDATE REMARKS SET remark ='"+ str((line.textfield.get("1.0", END)).strip()) + \
                                 "',sended = " + str(sendflag) + " where id = " + str(line.id) +  " and raport_number = '" + str(line.rn) + "'"

                        q_run(connD, querry)
                        if remflag == 'True':

                            if (str(line.requestdate)).strip() != '':
                                querry = "SELECT * from reminder where id = " + str(line.id) + \
                                         " and raport_number = '" + str(line.rn) + "'"

                                if len(q_run(connD, querry)) != 0 :
                                    # querry = "UPDATE REMINDER SET request_date = '" + str() + "' " \
                                    # ""\
                                    # "where id = " + str(line.id) + " and raport_number = '" + str(line.rn) + "'"
                                    querry = "UPDATE reminder SET request_date = '{}', remcom ='{}' where id = {} and raport_number = '{}'" \
                                        .format(line.requestdate,line.remindertextfield.get("1.0", END).strip(),line.id,line.rn)

                                else:
                                    querry = "INSERT INTO REMINDER(parent,raport_number,request_date,remcom,id) VALUES (" + str(
                                        line.parent) + ",'" + str(line.rn) + "','" + str(
                                        line.requestdate) + "','" + str(
                                        line.remindertextfield.get("1.0", END)).strip() + "'," + str(
                                        line.id) + ")"

                                q_run(connD, querry)
                            else:
                                messagebox.showinfo("Brak", 'Upload aborted. No request date. Please fill ' + str(
                                    line.name) + ' requestdate')
                                break

                    else:

                        querry = "INSERT INTO REMARKS(id,raport_number,remark,parent,documentdate,sended,reminder) VALUES (" + str(
                            line.id) + ",'" + str(line.rn) + "','" + str(
                            (line.textfield.get("1.0", END)).strip()) + "'," + str(line.parent) + ",'" + str(
                            measdate) +"'," + str(sendflag) + "," + str(remflag) + ")"

                        q_run(connD, querry)
                        if remflag == 'True':
                            if str(line.requestdate) != '':

                                querry = "INSERT INTO REMINDER(parent,raport_number,request_date,remcom,id) VALUES (" + str(
                                    line.parent) + ",'" + str(line.rn) + "','" + str(
                                    line.requestdate) + "','" + str(line.remindertextfield.get("1.0", END)).strip() + "'," + str(
                                    line.id) + ")"

                                q_run(connD, querry)
                            else:
                                messagebox.showinfo("Brak", 'Upload aborted. No request date. Please fill ' + str(line.name) + ' requestdate')
                                break

                    if str(line.textfield2.get("1.0", END)).strip() != '':

                        if line.id in iddfblist:
                            if str((line.FDBdateLabel.get("1.0", END)).strip()) != '':
                                querry = "UPDATE feedbacks SET feedback ='" + str((line.textfield2.get("1.0", END)).strip()) + \
                                         "',documentdate = '" + str((line.FDBdateLabel.get("1.0", END)).strip()) + "' where id = " + str(
                                    line.id) + " and raport_number = '" + str(line.rn) + "'"

                                q_run(connD, querry)
                            else:
                                messagebox.showinfo("Brak", 'Upload aborted. No feedback date. Please fill ' + str(line.name) + ' feedbackdate')
                                break
                        else:
                            if str((line.FDBdateLabel.get("1.0", END)).strip()) != '':
                                querry = "INSERT INTO FEEDBACKS(id,raport_number,feedback,parent,documentdate) VALUES (" + str(
                                    line.id) + ",'" + str(line.rn) + "','" + str(
                                    (line.textfield2.get("1.0", END)).strip()) + "'," + str(line.parent) + ",'" \
                                         + ((line.FDBdateLabel.get("1.0", END)).strip()) + "')"



                            else:
                                messagebox.showinfo("Brak", 'Upload aborted. No feedback date. Please fill ' + str(
                                    line.name) + ' feedbackdate')
                                break
                            q_run(connD, querry)

                    c+=1




            messagebox.showinfo("Finish", 'Upload done: ' + str(c))
        def filter():
            namefilter = filterentry.get()
            putdevices(report, namefilter)




        def noresponse():

            MsgBox = tk.messagebox.askquestion('Mark feedbacks ', 'Mark sent feedbacks as no response?',
                                               icon='question')
            if MsgBox == 'yes':
                try:
                    for line in remlist:
                        if line.var.get() == 1:
                            querry = "SELECT id from feedbacks where id = {} and raport_number = '{}'".format(line.id, report)
                            if len(q_run(connD, querry)) == 0:
                                line.textfield2.delete('1.0', END)
                                line.textfield2.insert(INSERT, 'No response')
                                line.FDBdateLabel.delete('1.0', END)
                                line.FDBdateLabel.insert(INSERT,datetime.datetime.now().date())

                                querry = "INSERT INTO feedbacks(parent,feedback,documentdate,fdbflag,raport_number,id) VALUES"\
                                "({},'No response','{}',8,'{}',{})".format(line.parent,datetime.datetime.now().date(),report,line.id )
                                q_run(connD, querry)

                    querry = "UPDATE harmonogram SET feedbacks_kto = (select ini from users where login = '{}' limit 1),feedbacks = 'True', feedbacks_start = '{}', feedbacks_koniec = '{}' where report_number = '{}'".format(\
                        connD[0],datetime.datetime.now(),datetime.datetime.now(), report)
                    q_run(connD, querry)
                    tk.messagebox.showinfo('Return', 'Upload done',icon='info')
                except:
                    tk.messagebox.showinfo('Error', 'UPLOAD ERROR',icon='error')
            else:
                pass


        querry = """select rem.id, rem.raport_number, rem.remark, rem.sended, remi.request_date, fdb.feedback, fdb.documentdate,remi.remcom
		from remarks as rem
		left join reminder as remi on rem.id = remi.id and rem.raport_number = remi.raport_number
		left join feedbacks as fdb on rem.id = fdb.id and rem.raport_number = fdb.raport_number
		where rem.raport_number = '""" + str(report) + """'
		group by rem.id, rem.raport_number, rem.remark, rem.sended, remi.request_date, fdb.feedback, fdb.documentdate,remi.remcom"""
        remtable = q_run(connD,querry)

        if str(namefilter) == 'None': namefilter = ''
        querry = """
		select dev.name ,ml.id,ml.parent,max(ml.date)
						from (select ml.parent, id, max(value),max(date) as date,raport_number
						from measurements_low as ml 
						where raport_number = '{}'
						group by ml.parent, ml.id,raport_number order by id) as ml 
						
						left join ds_structure as dss on cast(ml.id as text) =dss.id and ml.parent = dss.parent
						left join devices as dev on ml.id =dev.id
						
						where ml.raport_number = '{}' and lower(dev.name) ~'{}'
							group by ml.id,dss.sort,dev.name,ml.parent
							order by dss.sort""".format(report,report,str(namefilter).lower().strip())






        try:
            for widget in MASTERmeasframe.winfo_children():
                widget.destroy()
        except:
            pass

        MASTERmeasframe.pack(side=TOP, anchor=S, fill=BOTH)


        devices = q_run(connD, querry)
        head = tk.Frame(MASTERmeasframe, bd=1, relief=RAISED)
        header = tk.Frame(head, bd=1, relief=RAISED)
        searcher = tk.Frame(head, bd=1, relief=RAISED)
        head.pack()
        searcher.pack(side=LEFT)
        header.pack(side=LEFT)


        body = tk.Frame(MASTERmeasframe, bd=1, relief=RAISED )

        body.pack(fill=BOTH)
        v = StringVar(searcher, value=namefilter)
        filterentry = Entry(searcher, textvariable=v)
        filterentry.pack(side = TOP)


        filterbutton = Button(searcher, text = 'Filter name',command = filter)
        filterbutton.pack(side = TOP)

        UploadButton = Button(header, text='Update remarks & Feedbacks', command=upload)
        UploadButton.pack()
        GetReportButton = Button(header, text='Get remarks from vibration report', command=grabremarks)
        GetReportButton.pack()
        Makefdbrepbut = Button(header, text='Make feedback report', command= lambda rn = str(report), connD = connD : createfeedback(rn,connD))
        Makefdbrepbut.pack()
        Grabfeedback = Button(header, text='Get feedbacks from report', command= lambda rn = str(report), connD = connD : grabfeedbacks())
        Grabfeedback.pack()
        NoResponseBut = Button(header, text='No response', command = noresponse)
        NoResponseBut.pack()
        try:
            querry = "Select send_raport_koniec from harmonogram where report_number = '" + str(report) + "'"
            senddate = q_run(connD, querry)[0][0]
            senddate = str(senddate)[:10]
            if senddate == 'None':senddate = 'no date in harmonogram'
        except:
            senddate = 'no date in harmonogram'

        ttk.Label(header, text="Send date: " + str(senddate)).pack()
        scrollable_body = Scrollable(body, width=16)

        remlist.clear()
        for i in devices:
            measCframe = tk.Frame(scrollable_body,height = 0, bd=1, relief=RAISED)
            X = frame_reminder(measCframe,i,report,remtable)
            remlist.append(X)


        scrollable_body.update()


    remarksWindow = tk.Tk()
    remarksWindow.title("Remarks")
    parent = 1
    MASTERmeasframe = Frame(remarksWindow)
    Ownerlistbox = Listbox(remarksWindow, exportselection=False)
    Ownerlistbox.config(width=0)
    Ownerlistbox.bind('<Double-Button>', getships)
    Shiplistbox = Listbox(remarksWindow, exportselection=False)
    Shiplistbox.config(width=0)
    Shiplistbox.bind('<Double-Button>', getreports)
    Reportlistbox = Listbox(remarksWindow, exportselection=False)
    Reportlistbox.config(width=0)
    Reportlistbox.bind('<Double-Button>', getdevices)





    p=0
    devlist = list()
    for line in resultr:
        Ownerlistbox.insert(END, line[0])
        p+=1



    Ownerlistbox.pack(side=LEFT, fill=BOTH)
    Shiplistbox.pack(side=LEFT, fill=BOTH)
    Reportlistbox.pack(side=LEFT, fill=BOTH)
    remarksWindow.mainloop()


# root = tk.Tk()
# root.title("Login")
# root.geometry("200x120")
# LogApplication(root)  # The frame is inside the widgit
LogApplication()   #TUTAJ JEST APKA DO LOGOWANIE

#remindershow()

#connD=['testuser','info','192.168.10.243']
#remarks(connD)
#
#
#
# x = remarks(connD).frame_reminder(self, measCframe,dev,rn)

